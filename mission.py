import os
import json
from datetime import datetime, timedelta, time
import time as time_module
from itertools import permutations
import requests
import toml
import re
import unicodedata

import streamlit as st
import pandas as pd

# --------------------------
# CONFIG APP (DOIT ÃŠTRE EN PREMIER)
# --------------------------
st.set_page_config(
    page_title="Planificateur de mission terrain", 
    layout="wide",
    page_icon="ðŸ—ºï¸"
)

# Import des modules pour l'export PDF et Word
PDF_AVAILABLE = False
DOCX_AVAILABLE = False

try:
    # VÃ©rifier d'abord si reportlab est installÃ©
    import importlib
    importlib.import_module('reportlab')
    REPORTLAB_INSTALLED = True
except ImportError:
    REPORTLAB_INSTALLED = False

try:
    # VÃ©rifier si python-docx est installÃ©
    importlib.import_module('docx')
    DOCX_INSTALLED = True
except ImportError:
    DOCX_INSTALLED = False

try:
    from pdf_generator import create_pv_pdf, create_word_document, create_mission_pdf, create_docx_document
    # VÃ©rifier si les fonctions sont disponibles
    PDF_AVAILABLE = REPORTLAB_INSTALLED
    DOCX_AVAILABLE = DOCX_INSTALLED
    if not PDF_AVAILABLE:
        st.warning("âš ï¸ Module reportlab non installÃ©. Installez reportlab pour activer l'export PDF.")
    if not DOCX_AVAILABLE:
        st.warning("âš ï¸ Module python-docx non installÃ©. Installez python-docx pour l'export Word.")
except ImportError as e:
    PDF_AVAILABLE = False
    DOCX_AVAILABLE = False
    st.warning(f"âš ï¸ Module PDF/Word non disponible: {e}. Installez: pip install reportlab python-docx")

from geopy.geocoders import Nominatim
from geopy.extra.rate_limiter import RateLimiter

import folium
from streamlit_folium import st_folium

# --------------------------
# AUTHENTIFICATION
# --------------------------
# INITIALISATION DES VARIABLES DE SESSION
# --------------------------

st.title("ðŸ—ºï¸ Planificateur de mission (Moctar)")
st.caption("Optimisation d'itinÃ©raire + planning journalier + carte interactive + Ã©dition de rapport")

# --------------------------
# SIDEBAR: KEYS & OPTIONS
# --------------------------
st.sidebar.header("âš™ï¸ Configuration")

# Chargement sÃ©curisÃ© des clÃ©s API (Secrets â†’ ENV â†’ config.toml)
graphhopper_api_key = (
    st.secrets.get("api_keys", {}).get("graphhopper")
    or os.getenv("GRAPHOPPER_API_KEY")
)
deepseek_api_key = (
    st.secrets.get("api_keys", {}).get("deepseek")
    or os.getenv("DEEPSEEK_API_KEY")
)

if not graphhopper_api_key or not deepseek_api_key:
    try:
        config = toml.load('config.toml')
        graphhopper_api_key = graphhopper_api_key or config.get('api_keys', {}).get('graphhopper', '')
        deepseek_api_key = deepseek_api_key or config.get('api_keys', {}).get('deepseek', '')
        st.sidebar.success("âœ… ClÃ©s API chargÃ©es depuis config.toml")
    except Exception:
        # Pas de config.toml, on garde les valeurs actuelles (Secrets/ENV)
        pass

if graphhopper_api_key:
    st.sidebar.caption("ðŸ”‘ Maps prÃªt")
else:
    st.sidebar.warning("âš ï¸ ClÃ© GraphHopper absente â€” fallback activÃ© si besoin")

if deepseek_api_key:
    st.sidebar.caption("ðŸ¤– Adja prÃªt")
else:
    st.sidebar.warning("âš ï¸ ClÃ© DeepSeek absente â€” fonctionnalitÃ©s IA limitÃ©es")

st.sidebar.subheader("Calcul des distances")
distance_method = st.sidebar.radio(
    "MÃ©thode de calcul",
    [
        "Auto (OSRM â†’ Automatique â†’ Maps â†’ GÃ©omÃ©trique)",
        "Automatique uniquement",
        "OSRM uniquement (rapide)",
        "GÃ©omÃ©trique uniquement",
        "Maps uniquement (prÃ©cis)"
    ],
    index=0
)

use_deepseek_fallback = st.sidebar.checkbox(
    "Utiliser Maps si Automatique Ã©choue", 
    value=True,
    help="Appeler le service de routage Maps si l'estimation automatique Ã©choue"
)

with st.sidebar.expander("Options avancÃ©es"):
    # Charger les paramÃ¨tres par dÃ©faut depuis la configuration (sÃ©curisÃ©)
    try:
        secrets_settings = st.secrets.get("settings", {})
    except Exception:
        secrets_settings = {}
    try:
        local_config = toml.load('config.toml') if os.path.exists('config.toml') else {}
    except Exception:
        local_config = {}

    config_speed = secrets_settings.get("default_speed_kmh")
    config_cache = secrets_settings.get("use_cache")
    config_debug = secrets_settings.get("debug_mode")
    config_osrm = secrets_settings.get("osrm_base_url")

    if config_speed is None:
        config_speed = local_config.get('settings', {}).get('default_speed_kmh', 95)
    if config_cache is None:
        config_cache = local_config.get('settings', {}).get('use_cache', True)
    if config_debug is None:
        config_debug = local_config.get('settings', {}).get('debug_mode', False)
    if config_osrm is None:
        config_osrm = local_config.get('settings', {}).get('osrm_base_url', "https://router.project-osrm.org")
    
    default_speed_kmh = st.number_input(
        "Vitesse moyenne (km/h) pour estimations", 
        min_value=20, max_value=120, value=config_speed
    )
    use_cache = st.checkbox("Utiliser le cache pour gÃ©ocodage", value=config_cache)
    prefer_offline_geocoding = st.checkbox(
        "Prioriser coordonnÃ©es locales pour grandes villes",
        value=True,
        key="prefer_offline_geocoding",
        help="Utiliser des coordonnÃ©es vÃ©rifiÃ©es pour grandes villes du SÃ©nÃ©gal (ex. Dakar, Louga, Touba)."
    )
    debug_mode = st.checkbox("Mode debug (afficher dÃ©tails calculs)", value=config_debug)
    osrm_base_url = st.text_input(
        "OSRM base URL",
        value=config_osrm,
        help="Exemple: http://localhost:5000 ou https://router.project-osrm.org"
    )

# --------------------------
# Ã‰TAT DE SESSION
# --------------------------
if 'planning_results' not in st.session_state:
    st.session_state.planning_results = None

if 'editing_event' not in st.session_state:
    st.session_state.editing_event = None

if 'edit_mode' not in st.session_state:
    st.session_state.edit_mode = False

if 'manual_itinerary' not in st.session_state:
    st.session_state.manual_itinerary = None

# --------------------------
# FONCTIONS CARBURANT ET EMPREINTE CARBONE
# --------------------------

def get_vehicle_types():
    """Retourne les types de vÃ©hicules disponibles avec leurs consommations"""
    return {
        "Station-Wagon": {"consumption": 13.0, "fuel_type": "Essence", "co2_factor": 2.31},
        "Berline": {"consumption": 9.5, "fuel_type": "Essence", "co2_factor": 2.31},
        "SUV": {"consumption": 15.0, "fuel_type": "Essence", "co2_factor": 2.31},
        "4x4": {"consumption": 18.0, "fuel_type": "Diesel", "co2_factor": 2.68},
        "Utilitaire": {"consumption": 12.0, "fuel_type": "Diesel", "co2_factor": 2.68},
        "Minibus": {"consumption": 20.0, "fuel_type": "Diesel", "co2_factor": 2.68},
     }

def calculate_fuel_consumption(total_distance_km, vehicle_type):
    """Calcule la consommation de carburant pour un vÃ©hicule donnÃ©"""
    vehicles = get_vehicle_types()
    if vehicle_type not in vehicles:
        return None
    
    consumption_per_100km = vehicles[vehicle_type]["consumption"]
    fuel_needed = (total_distance_km * consumption_per_100km) / 100
    
    return {
        "fuel_needed_liters": fuel_needed,
        "consumption_per_100km": consumption_per_100km,
        "fuel_type": vehicles[vehicle_type]["fuel_type"]
    }

def calculate_carbon_footprint(fuel_consumption_data, total_distance_km, vehicle_type):
    """Calcule l'empreinte carbone de la mission"""
    vehicles = get_vehicle_types()
    if vehicle_type not in vehicles or not fuel_consumption_data:
        return None
    
    # Facteur d'Ã©mission CO2 (kg CO2 par litre de carburant)
    co2_factor = vehicles[vehicle_type]["co2_factor"]
    
    # Calcul des Ã©missions CO2
    co2_emissions_kg = fuel_consumption_data["fuel_needed_liters"] * co2_factor
    co2_emissions_tons = co2_emissions_kg / 1000
    
    # Ã‰quivalences pour contextualiser
    trees_needed = co2_emissions_kg / 22  # Un arbre absorbe ~22kg CO2/an
    
    return {
        "co2_emissions_kg": co2_emissions_kg,
        "co2_emissions_tons": co2_emissions_tons,
        "trees_equivalent": trees_needed,
        "fuel_type": fuel_consumption_data["fuel_type"],
        "distance_km": total_distance_km
    }

def estimate_fuel_cost(fuel_consumption_data, fuel_price_per_liter=None):
    """Estime le coÃ»t du carburant"""
    if not fuel_consumption_data:
        return None
    
    # Charger les prix depuis la configuration (Streamlit Secrets ou config.toml)
    default_prices = None
    try:
        secrets_settings = st.secrets.get('settings', {})
        fuel_prices = secrets_settings.get('fuel_prices')
    except Exception:
        fuel_prices = None
    if fuel_prices is None:
        try:
            config = toml.load('config.toml') if os.path.exists('config.toml') else {}
            fuel_prices = config.get('settings', {}).get('fuel_prices')
        except Exception:
            fuel_prices = None
    if fuel_prices:
        default_prices = {
            "Essence": fuel_prices.get('essence'),
            "Diesel": fuel_prices.get('diesel')
        }
    else:
        # Prix par dÃ©faut au SÃ©nÃ©gal (en FCFA) si aucune config disponible
        default_prices = {
            "Essence": 1350,  # FCFA par litre
            "Diesel": 1200    # FCFA par litre
        }
    
    fuel_type = fuel_consumption_data["fuel_type"]
    price = fuel_price_per_liter if fuel_price_per_liter else default_prices.get(fuel_type, 1300)
    
    total_cost = fuel_consumption_data["fuel_needed_liters"] * price
    
    return {
        "total_cost_fcfa": total_cost,
        "price_per_liter": price,
        "fuel_type": fuel_type,
        "liters": fuel_consumption_data["fuel_needed_liters"]
    }

# --------------------------
# FONCTIONS RAPPORT IA ADJA
# --------------------------
def collect_mission_data_for_ai():
    """Collecte toutes les donnÃ©es de mission pour l'IA Adja"""
    if not st.session_state.planning_results:
        return None
    
    results = st.session_state.planning_results
    itinerary = st.session_state.manual_itinerary or results['itinerary']
    
    # DonnÃ©es de base
    mission_data = {
        'sites': results['sites_ordered'],
        'stats': results['stats'],
        'itinerary': itinerary,
        'calculation_method': results.get('calculation_method', 'Non spÃ©cifiÃ©'),
        'base_location': results.get('base_location', ''),
        'segments_summary': results.get('segments_summary', [])
    }
    
    # Analyse dÃ©taillÃ©e des activitÃ©s
    activities = {}
    detailed_activities = []
    
    for day, sdt, edt, desc in itinerary:
        activity_type = "Autre"
        if "Visite" in desc or "RÃ©union" in desc:
            activity_type = "Visite/RÃ©union"
        elif "Trajet" in desc or "km" in desc:
            activity_type = "DÃ©placement"
        elif "Pause" in desc or "Repos" in desc:
            activity_type = "Pause"
        elif "NuitÃ©e" in desc:
            activity_type = "HÃ©bergement"
        
        duration_hours = (edt - sdt).total_seconds() / 3600
        
        if activity_type not in activities:
            activities[activity_type] = 0
        activities[activity_type] += duration_hours
        
        # DÃ©tails de chaque activitÃ©
        detailed_activities.append({
            'day': day,
            'start_time': sdt.strftime('%H:%M'),
            'end_time': edt.strftime('%H:%M'),
            'duration': duration_hours,
            'type': activity_type,
            'description': desc
        })
    
    mission_data['activities_breakdown'] = activities
    mission_data['detailed_activities'] = detailed_activities
    
    # Ajouter les donnÃ©es enrichies si disponibles
    if hasattr(st.session_state, 'mission_notes'):
        mission_data['mission_notes'] = st.session_state.mission_notes
    if hasattr(st.session_state, 'activity_details'):
        mission_data['activity_details'] = st.session_state.activity_details
    if hasattr(st.session_state, 'mission_context'):
        mission_data['mission_context'] = st.session_state.mission_context
    
    return mission_data

def collect_construction_report_data():
    """Interface pour collecter des donnÃ©es spÃ©cifiques au procÃ¨s-verbal de chantier"""
    st.markdown("### ðŸ—ï¸ DonnÃ©es pour ProcÃ¨s-Verbal de Chantier")
    
    # Informations gÃ©nÃ©rales du chantier
    col1, col2 = st.columns(2)
    
    with col1:
        project_name = st.text_input(
            "ðŸ—ï¸ Nom du projet/chantier",
            placeholder="Ex: Travaux d'entretien PA DAL zone SUD",
            key="project_name"
        )
        
        report_date = st.date_input(
            "ðŸ“… Date de la visite",
            value=datetime.now().date(),
            key="report_date"
        )
        
        site_location = st.text_input(
            "ðŸ“ Localisation du site",
            placeholder="Ex: VÃ©lingara et Kolda",
            key="site_location"
        )
    
    with col2:
        report_type = st.selectbox(
            "ðŸ“‹ Type de rapport",
            ["ProcÃ¨s-verbal de visite de chantier", "Rapport d'avancement", "Rapport de fin de travaux", "Rapport d'incident"],
            key="construction_report_type"
        )
        
        weather_conditions = st.text_input(
            "ðŸŒ¤ï¸ Conditions mÃ©tÃ©orologiques",
            placeholder="Ex: EnsoleillÃ©, pluvieux, venteux...",
            key="weather_conditions"
        )
    
    # Liste de prÃ©sence
    st.markdown("### ðŸ‘¥ Liste de PrÃ©sence")
    
    if 'attendees' not in st.session_state:
        st.session_state.attendees = []
    
    col_add, col_clear = st.columns([3, 1])
    with col_add:
        new_attendee_name = st.text_input("Nom", key="new_attendee_name")
        new_attendee_structure = st.text_input("Structure/Entreprise", key="new_attendee_structure")
        new_attendee_function = st.text_input("Fonction", key="new_attendee_function")
    
    with col_clear:
        st.write("")  # Espacement
        st.write("")  # Espacement
        if st.button("âž• Ajouter"):
            if new_attendee_name and new_attendee_structure:
                st.session_state.attendees.append({
                    'nom': new_attendee_name,
                    'structure': new_attendee_structure,
                    'fonction': new_attendee_function
                })
                st.rerun()
        
        if st.button("ðŸ—‘ï¸ Vider"):
            st.session_state.attendees = []
            st.rerun()
    
    # Affichage de la liste
    if st.session_state.attendees:
        st.markdown("**Participants enregistrÃ©s :**")
        for i, attendee in enumerate(st.session_state.attendees):
            st.write(f"{i+1}. **{attendee['nom']}** - {attendee['structure']} ({attendee['fonction']})")
    
    # Intervenants dans le projet
    st.markdown("### ðŸ¢ DiffÃ©rents Intervenants dans le Projet")
    
    col1, col2 = st.columns(2)
    with col1:
        master_contractor = st.text_input(
            "ðŸ—ï¸ MaÃ®tre d'ouvrage",
            placeholder="Ex: Sonatel",
            key="master_contractor"
        )
        
        main_contractor = st.text_input(
            "ðŸ”§ Entreprise principale",
            placeholder="Ex: KonÃ© Construction",
            key="main_contractor"
        )
    
    with col2:
        project_manager = st.text_input(
            "ðŸ‘¨â€ðŸ’¼ MaÃ®tre d'Å“uvre",
            placeholder="Ex: Sonatel",
            key="project_manager"
        )
        
        supervisor = st.text_input(
            "ðŸ‘·â€â™‚ï¸ Superviseur/ContrÃ´leur",
            placeholder="Ex: SECK CONS",
            key="supervisor"
        )
    
    # Documents contractuels
    st.markdown("### ðŸ“„ Documents Contractuels")
    
    if 'contract_documents' not in st.session_state:
        st.session_state.contract_documents = []
    
    col_doc1, col_doc2, col_doc3, col_add_doc = st.columns([2, 2, 2, 1])
    
    with col_doc1:
        doc_name = st.text_input("Document", key="doc_name")
    with col_doc2:
        doc_holder = st.text_input("Porteur", key="doc_holder")
    with col_doc3:
        doc_comments = st.text_input("Commentaires", key="doc_comments")
    with col_add_doc:
        st.write("")  # Espacement
        if st.button("âž•", key="add_doc"):
            if doc_name and doc_holder:
                st.session_state.contract_documents.append({
                    'document': doc_name,
                    'porteur': doc_holder,
                    'commentaires': doc_comments
                })
                st.rerun()
    
    if st.session_state.contract_documents:
        st.markdown("**Documents enregistrÃ©s :**")
        for i, doc in enumerate(st.session_state.contract_documents):
            st.write(f"â€¢ **{doc['document']}** - Porteur: {doc['porteur']} - {doc['commentaires']}")
    
    # Respect du planning
    st.markdown("### â° Respect du Planning")
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        start_notification = st.date_input(
            "ðŸ“… Notification dÃ©marrage",
            key="start_notification"
        )
        
        contractual_delay = st.number_input(
            "â±ï¸ DÃ©lai contractuel (jours)",
            min_value=0,
            value=40,
            key="contractual_delay"
        )
    
    with col2:
        remaining_delay = st.number_input(
            "â³ DÃ©lai restant (jours)",
            min_value=0,
            value=0,
            key="remaining_delay"
        )
        
        progress_percentage = st.slider(
            "ðŸ“Š Avancement global (%)",
            min_value=0,
            max_value=100,
            value=50,
            key="progress_percentage"
        )
    
    with col3:
        planning_status = st.selectbox(
            "ðŸ“ˆ Ã‰tat du planning",
            ["En avance", "Dans les temps", "En retard", "Critique"],
            index=2,
            key="planning_status"
        )
    
    # Observations dÃ©taillÃ©es par site
    st.markdown("### ðŸ” Observations DÃ©taillÃ©es par Site")
    
    if st.session_state.planning_results:
        sites = st.session_state.planning_results['sites_ordered']
        
        for i, site in enumerate(sites):
            st.markdown(f"#### ðŸ“ Site de {site['Ville']}")
            
            # Observations par catÃ©gorie
            col1, col2 = st.columns(2)
            
            with col1:
                st.markdown("**ðŸ¢ Agence commerciale :**")
                agency_work = st.text_area(
                    "Travaux rÃ©alisÃ©s",
                    placeholder="Ex: Aucun des travaux prÃ©vus n'a Ã©tÃ© rÃ©alisÃ©...",
                    height=80,
                    key=f"agency_work_{i}"
                )
                
                st.markdown("**ðŸ—ï¸ BÃ¢timent technique :**")
                technical_work = st.text_area(
                    "Ã‰tat des travaux techniques",
                    placeholder="Ex: Travaux de carrelage de faÃ§ade et rÃ©habilitation des toilettes...",
                    height=80,
                    key=f"technical_work_{i}"
                )
            
            with col2:
                st.markdown("**ðŸ  Logement du gardien :**")
                guard_housing = st.text_area(
                    "Ã‰tat du logement",
                    placeholder="Ex: MÃ©canisme de la chasse anglaise installÃ© mais non fonctionnel...",
                    height=80,
                    key=f"guard_housing_{i}"
                )
                
                st.markdown("**ðŸšª FaÃ§ade de l'agence :**")
                facade_work = st.text_area(
                    "Travaux de faÃ§ade",
                    placeholder="Ex: Corriger les portes qui ne se ferment pas...",
                    height=80,
                    key=f"facade_work_{i}"
                )
            
            # Poste de garde
            st.markdown("**ðŸ›¡ï¸ Poste de garde :**")
            guard_post = st.text_area(
                "Ã‰tat du poste de garde",
                placeholder="Ex: Peinture du poste de garde non conforme...",
                height=68,
                key=f"guard_post_{i}"
            )
    
    # Observations gÃ©nÃ©rales et recommandations
    st.markdown("### ðŸ“ Observations GÃ©nÃ©rales et Recommandations")
    
    general_observations = st.text_area(
        "ðŸ” Constat gÃ©nÃ©ral",
        placeholder="Ex: Lors des visites de chantier, plusieurs constats majeurs ont Ã©tÃ© relevÃ©s concernant la qualitÃ© d'exÃ©cution...",
        height=120,
        key="general_observations"
    )
    
    recommendations = st.text_area(
        "ðŸ’¡ Recommandations",
        placeholder="Ex: Il est impÃ©ratif que KONE CONSTRUCTION mette en place un dispositif correctif immÃ©diat...",
        height=120,
        key="construction_recommendations"
    )
    
    # Informations du rapporteur
    st.markdown("### âœï¸ Informations du Rapporteur")
    
    col1, col2 = st.columns(2)
    
    with col1:
        reporter_name = st.text_input(
            "ðŸ‘¤ Nom du rapporteur",
            placeholder="Ex: Moctar TALL",
            key="reporter_name"
        )
        
        report_location = st.text_input(
            "ðŸ“ Lieu de rÃ©daction",
            placeholder="Ex: Dakar",
            key="report_location"
        )
    
    with col2:
        reporter_function = st.text_input(
            "ðŸ’¼ Fonction",
            placeholder="Ex: IngÃ©nieur Projet",
            key="reporter_function"
        )
        
        report_completion_date = st.date_input(
            "ðŸ“… Date de finalisation",
            value=datetime.now().date(),
            key="report_completion_date"
        )
    
    return {
        'project_info': {
            'project_name': project_name,
            'report_date': report_date,
            'site_location': site_location,
            'report_type': report_type,
            'weather_conditions': weather_conditions
        },
        'attendees': st.session_state.attendees,
        'stakeholders': {
            'master_contractor': master_contractor,
            'main_contractor': main_contractor,
            'project_manager': project_manager,
            'supervisor': supervisor
        },
        'contract_documents': st.session_state.contract_documents,
        'planning': {
            'start_notification': start_notification,
            'contractual_delay': contractual_delay,
            'remaining_delay': remaining_delay,
            'progress_percentage': progress_percentage,
            'planning_status': planning_status
        },
        'observations': {
            'general_observations': general_observations,
            'recommendations': recommendations
        },
        'reporter': {
            'reporter_name': reporter_name,
            'reporter_function': reporter_function,
            'report_location': report_location,
            'report_completion_date': report_completion_date
        }
    }

def collect_enhanced_mission_data():
    """Interface pour collecter des donnÃ©es enrichies sur la mission"""
    st.markdown("### ðŸ“ Informations dÃ©taillÃ©es sur la mission")
    
    # Contexte gÃ©nÃ©ral de la mission
    col1, col2 = st.columns(2)
    
    with col1:
        mission_objective = st.text_area(
            "ðŸŽ¯ Objectif principal de la mission",
            placeholder="Ex: Audit des agences rÃ©gionales, formation du personnel, prospection commerciale...",
            height=100,
            key="mission_objective"
        )
        
        mission_participants = st.text_input(
            "ðŸ‘¥ Participants Ã  la mission",
            placeholder="Ex: Jean Dupont (Chef de projet), Marie Martin (Analyste)...",
            key="mission_participants"
        )
    
    with col2:
        mission_budget = st.number_input(
            "ðŸ’° Budget allouÃ© (FCFA)",
            min_value=0,
            value=0,
            step=10000,
            key="mission_budget"
        )
        
        mission_priority = st.selectbox(
            "âš¡ PrioritÃ© de la mission",
            ["Faible", "Normale", "Ã‰levÃ©e", "Critique"],
            index=1,
            key="mission_priority"
        )
    
    # Notes par site/activitÃ©
    st.markdown("### ðŸ“‹ Notes dÃ©taillÃ©es par site")
    
    if st.session_state.planning_results:
        sites = st.session_state.planning_results['sites_ordered']
        
        if 'activity_details' not in st.session_state:
            st.session_state.activity_details = {}
        
        for i, site in enumerate(sites):
            # Utilisation d'un container au lieu d'un expander pour Ã©viter l'imbrication
            st.markdown(f"### ðŸ“ {site['Ville']} - {site['Type']} ({site['ActivitÃ©']})")
            with st.container():
                col_notes, col_details = st.columns(2)
                
                with col_notes:
                    notes = st.text_area(
                        "ðŸ“ Notes et observations",
                        placeholder="DÃ©crivez ce qui s'est passÃ©, les rÃ©sultats obtenus, les difficultÃ©s rencontrÃ©es...",
                        height=120,
                        key=f"notes_{i}"
                    )
                    
                    success_level = st.select_slider(
                        "âœ… Niveau de rÃ©ussite",
                        options=["Ã‰chec", "Partiel", "Satisfaisant", "Excellent"],
                        value="Satisfaisant",
                        key=f"success_{i}"
                    )
                
                with col_details:
                    contacts_met = st.text_input(
                        "ðŸ¤ Personnes rencontrÃ©es",
                        placeholder="Noms et fonctions des contacts",
                        key=f"contacts_{i}"
                    )
                    
                    outcomes = st.text_area(
                        "ðŸŽ¯ RÃ©sultats obtenus",
                        placeholder="Accords signÃ©s, informations collectÃ©es, problÃ¨mes identifiÃ©s...",
                        height=80,
                        key=f"outcomes_{i}"
                    )
                    
                    follow_up = st.text_input(
                        "ðŸ“… Actions de suivi",
                        placeholder="Prochaines Ã©tapes, rendez-vous programmÃ©s...",
                        key=f"follow_up_{i}"
                    )
                
                # Stocker les dÃ©tails
                st.session_state.activity_details[f"site_{i}"] = {
                    'site_name': site['Ville'],
                    'site_type': site['Type'],
                    'activity': site['ActivitÃ©'],
                    'notes': notes,
                    'success_level': success_level,
                    'contacts_met': contacts_met,
                    'outcomes': outcomes,
                    'follow_up': follow_up
                }
    
    # Observations gÃ©nÃ©rales
    st.markdown("### ðŸ” Observations gÃ©nÃ©rales")
    
    col_obs1, col_obs2 = st.columns(2)
    
    with col_obs1:
        challenges = st.text_area(
            "âš ï¸ DifficultÃ©s rencontrÃ©es",
            placeholder="ProblÃ¨mes logistiques, retards, obstacles imprÃ©vus...",
            height=100,
            key="challenges"
        )
        
        lessons_learned = st.text_area(
            "ðŸ“š LeÃ§ons apprises",
            placeholder="Ce qui a bien fonctionnÃ©, ce qu'il faut amÃ©liorer...",
            height=100,
            key="lessons_learned"
        )
    
    with col_obs2:
        recommendations = st.text_area(
            "ðŸ’¡ Recommandations",
            placeholder="Suggestions pour les prochaines missions...",
            height=100,
            key="mission_recommendations"
        )
        
        overall_satisfaction = st.select_slider(
            "ðŸ˜Š Satisfaction globale",
            options=["TrÃ¨s insatisfait", "Insatisfait", "Neutre", "Satisfait", "TrÃ¨s satisfait"],
            value="Satisfait",
            key="overall_satisfaction"
        )
    
    # Stocker le contexte de mission
    st.session_state.mission_context = {
        'objective': mission_objective,
        'participants': mission_participants,
        'budget': mission_budget,
        'priority': mission_priority,
        'challenges': challenges,
        'lessons_learned': lessons_learned,
        'recommendations': recommendations,
        'overall_satisfaction': overall_satisfaction
    }
    
    return True

def ask_interactive_questions():
    """Pose des questions interactives pour orienter le rapport"""
    st.markdown("### ðŸ¤– Questions pour personnaliser votre rapport")
    
    questions_data = {}
    
    # Questions sur le type de rapport souhaitÃ©
    col1, col2 = st.columns(2)
    
    with col1:
        report_focus = st.multiselect(
            "ðŸŽ¯ Sur quoi souhaitez-vous que le rapport se concentre ?",
            ["RÃ©sultats obtenus", "EfficacitÃ© opÃ©rationnelle", "Aspects financiers", 
             "Relations clients", "ProblÃ¨mes identifiÃ©s", "OpportunitÃ©s dÃ©couvertes",
             "Performance de l'Ã©quipe", "Logistique et organisation"],
            default=["RÃ©sultats obtenus", "EfficacitÃ© opÃ©rationnelle"],
            key="report_focus"
        )
        
        target_audience = st.selectbox(
            "ðŸ‘¥ Qui va lire ce rapport ?",
            ["Direction gÃ©nÃ©rale", "Ã‰quipe projet", "Clients", "Partenaires", 
             "Ã‰quipe terrain", "Conseil d'administration"],
            key="target_audience"
        )
    
    with col2:
        report_length = st.selectbox(
            "ðŸ“„ Longueur souhaitÃ©e du rapport",
            ["Court (1-2 pages)", "Moyen (3-5 pages)", "DÃ©taillÃ© (5+ pages)"],
            index=1,
            key="report_length"
        )
        
        include_metrics = st.checkbox(
            "ðŸ“Š Inclure des mÃ©triques et KPIs",
            value=True,
            key="include_metrics"
        )
    
    # Questions spÃ©cifiques selon le contexte
    st.markdown("**Questions spÃ©cifiques :**")
    
    col3, col4 = st.columns(2)
    
    with col3:
        highlight_successes = st.checkbox(
            "ðŸ† Mettre en avant les succÃ¨s",
            value=True,
            key="highlight_successes"
        )
        
        discuss_challenges = st.checkbox(
            "âš ï¸ Discuter des dÃ©fis en dÃ©tail",
            value=True,
            key="discuss_challenges"
        )
        
        future_planning = st.checkbox(
            "ðŸ”® Inclure la planification future",
            value=True,
            key="future_planning"
        )
    
    with col4:
        cost_analysis = st.checkbox(
            "ðŸ’° Analyser les coÃ»ts en dÃ©tail",
            value=False,
            key="cost_analysis"
        )
        
        time_efficiency = st.checkbox(
            "â±ï¸ Analyser l'efficacitÃ© temporelle",
            value=True,
            key="time_efficiency"
        )
        
        stakeholder_feedback = st.checkbox(
            "ðŸ’¬ Inclure les retours des parties prenantes",
            value=False,
            key="stakeholder_feedback"
        )
    
    # Question ouverte pour personnalisation
    specific_request = st.text_area(
        "âœ¨ Y a-t-il des aspects spÃ©cifiques que vous souhaitez voir dans le rapport ?",
        placeholder="Ex: Comparaison avec la mission prÃ©cÃ©dente, focus sur un site particulier, analyse d'un problÃ¨me spÃ©cifique...",
        height=80,
        key="specific_request"
    )
    
    questions_data = {
        'report_focus': report_focus,
        'target_audience': target_audience,
        'report_length': report_length,
        'include_metrics': include_metrics,
        'highlight_successes': highlight_successes,
        'discuss_challenges': discuss_challenges,
        'future_planning': future_planning,
        'cost_analysis': cost_analysis,
        'time_efficiency': time_efficiency,
        'stakeholder_feedback': stakeholder_feedback,
        'specific_request': specific_request
    }
    
    return questions_data

def generate_enhanced_ai_report(mission_data, questions_data, api_key):
    """GÃ©nÃ¨re un rapport de mission amÃ©liorÃ© via l'IA Adja DeepSeek"""
    try:
        # Construction du prompt amÃ©liorÃ©
        prompt = build_enhanced_report_prompt(mission_data, questions_data)
        
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }
        
        # Ajuster max_tokens selon la longueur demandÃ©e
        max_tokens_map = {
            "Court (1-2 pages)": 2000,
            "Moyen (3-5 pages)": 4000,
            "DÃ©taillÃ© (5+ pages)": 6000
        }
        
        max_tokens = max_tokens_map.get(questions_data.get('report_length', 'Moyen (3-5 pages)'), 4000)
        
        data = {
            "model": "deepseek-chat",
            "messages": [{"role": "user", "content": prompt}],
            "temperature": 0.3,
            "max_tokens": max_tokens
        }
        
        response = requests.post(
            "https://api.deepseek.com/chat/completions",
            headers=headers,
            json=data,
            timeout=90
        )
        
        if response.status_code == 200:
            result = response.json()
            content = result.get("choices", [{}])[0].get("message", {}).get("content", "")
            return content
        else:
            st.error(f"Erreur API DeepSeek: {response.status_code}")
            return None
            
    except Exception as e:
        st.error(f"Erreur lors de la gÃ©nÃ©ration: {str(e)}")
        return None

def build_enhanced_report_prompt(mission_data, questions_data):
    """Construit un prompt amÃ©liorÃ© orientÃ© activitÃ©s pour la gÃ©nÃ©ration de rapport"""
    
    stats = mission_data['stats']
    sites = mission_data['sites']
    activities = mission_data['activities_breakdown']
    detailed_activities = mission_data.get('detailed_activities', [])
    mission_context = mission_data.get('mission_context', {})
    activity_details = mission_data.get('activity_details', {})
    
    # Construction des informations dÃ©taillÃ©es sur les activitÃ©s
    activities_info = ""
    if activity_details:
        activities_info = "\nDÃ‰TAILS DES ACTIVITÃ‰S PAR SITE:\n"
        for site_key, details in activity_details.items():
            if details.get('notes') or details.get('outcomes'):
                activities_info += f"\nðŸ“ {details['site_name']} ({details['site_type']}):\n"
                activities_info += f"   - ActivitÃ©: {details['activity']}\n"
                if details.get('notes'):
                    activities_info += f"   - Notes: {details['notes']}\n"
                if details.get('contacts_met'):
                    activities_info += f"   - Contacts: {details['contacts_met']}\n"
                if details.get('outcomes'):
                    activities_info += f"   - RÃ©sultats: {details['outcomes']}\n"
                if details.get('success_level'):
                    activities_info += f"   - Niveau de rÃ©ussite: {details['success_level']}\n"
                if details.get('follow_up'):
                    activities_info += f"   - Suivi: {details['follow_up']}\n"
    
    # Contexte de mission
    context_info = ""
    if mission_context:
        context_info = f"\nCONTEXTE DE LA MISSION:\n"
        if mission_context.get('objective'):
            context_info += f"- Objectif: {mission_context['objective']}\n"
        if mission_context.get('participants'):
            context_info += f"- Participants: {mission_context['participants']}\n"
        if mission_context.get('budget') and mission_context['budget'] > 0:
            context_info += f"- Budget: {mission_context['budget']:,} FCFA\n"
        if mission_context.get('priority'):
            context_info += f"- PrioritÃ©: {mission_context['priority']}\n"
        if mission_context.get('challenges'):
            context_info += f"- DÃ©fis: {mission_context['challenges']}\n"
        if mission_context.get('lessons_learned'):
            context_info += f"- LeÃ§ons apprises: {mission_context['lessons_learned']}\n"
        if mission_context.get('overall_satisfaction'):
            context_info += f"- Satisfaction globale: {mission_context['overall_satisfaction']}\n"
    
    # Focus du rapport selon les rÃ©ponses
    focus_areas = questions_data.get('report_focus', [])
    focus_instruction = ""
    if focus_areas:
        focus_instruction = f"\nLE RAPPORT DOIT SE CONCENTRER PARTICULIÃˆREMENT SUR: {', '.join(focus_areas)}"
    
    # Instructions spÃ©cifiques
    specific_instructions = []
    if questions_data.get('highlight_successes'):
        specific_instructions.append("- Mettre en Ã©vidence les succÃ¨s et rÃ©alisations")
    if questions_data.get('discuss_challenges'):
        specific_instructions.append("- Analyser en dÃ©tail les dÃ©fis rencontrÃ©s")
    if questions_data.get('future_planning'):
        specific_instructions.append("- Inclure des recommandations pour l'avenir")
    if questions_data.get('cost_analysis'):
        specific_instructions.append("- Fournir une analyse dÃ©taillÃ©e des coÃ»ts")
    if questions_data.get('time_efficiency'):
        specific_instructions.append("- Analyser l'efficacitÃ© temporelle de la mission")
    if questions_data.get('stakeholder_feedback'):
        specific_instructions.append("- IntÃ©grer les retours des parties prenantes")
    if questions_data.get('include_metrics'):
        specific_instructions.append("- Inclure des mÃ©triques et indicateurs de performance")
    
    instructions_text = "\n".join(specific_instructions) if specific_instructions else ""
    
    prompt = f"""Tu es un expert en rÃ©daction de rapports de mission professionnels. GÃ©nÃ¨re un rapport dÃ©taillÃ© et orientÃ© ACTIVITÃ‰S (pas trajets) en franÃ§ais.

DONNÃ‰ES DE BASE:
- DurÃ©e totale: {stats['total_days']} jour(s)
- Distance totale: {stats['total_km']:.1f} km
- Temps de visite total: {stats['total_visit_hours']:.1f} heures
- Nombre de sites visitÃ©s: {len([s for s in sites if s.get('Type') != 'Base'])}
- Sites visitÃ©s: {', '.join([s.get('Ville') for s in sites if s.get('Type') != 'Base'])}
- MÃ©thode de calcul: {mission_data['calculation_method']}

RÃ‰PARTITION DES ACTIVITÃ‰S:
{chr(10).join([f"- {act}: {hours:.1f}h" for act, hours in activities.items()])}

{context_info}

{activities_info}

PARAMÃˆTRES DU RAPPORT:
- Public cible: {questions_data.get('target_audience', 'Direction gÃ©nÃ©rale')}
- Longueur: {questions_data.get('report_length', 'Moyen (3-5 pages)')}
{focus_instruction}

INSTRUCTIONS SPÃ‰CIFIQUES:
{instructions_text}

DEMANDE SPÃ‰CIALE:
{questions_data.get('specific_request', 'Aucune demande spÃ©ciale')}

STRUCTURE REQUISE:
1. ðŸ“‹ RÃ‰SUMÃ‰ EXÃ‰CUTIF
2. ðŸŽ¯ OBJECTIFS ET CONTEXTE
3. ðŸ“ DÃ‰ROULEMENT DES ACTIVITÃ‰S (focus principal)
   - DÃ©tail par site avec rÃ©sultats obtenus
   - Personnes rencontrÃ©es et Ã©changes
   - SuccÃ¨s et difficultÃ©s par activitÃ©
4. ðŸ“Š ANALYSE DES RÃ‰SULTATS
   - Objectifs atteints vs prÃ©vus
   - Indicateurs de performance
   - Retour sur investissement
5. ðŸ” OBSERVATIONS ET ENSEIGNEMENTS
6. ðŸ’¡ RECOMMANDATIONS ET ACTIONS DE SUIVI
7. ðŸ“ˆ CONCLUSION ET PERSPECTIVES

IMPORTANT: 
- Concentre-toi sur les ACTIVITÃ‰S et leurs RÃ‰SULTATS, pas sur les trajets
- Utilise les donnÃ©es dÃ©taillÃ©es fournies pour chaque site
- Adopte un ton professionnel adaptÃ© au public cible
- Structure clairement avec des titres et sous-titres
- Inclus des mÃ©triques concrÃ¨tes quand disponibles"""

    return prompt

def build_report_prompt(mission_data, report_type, tone, include_recommendations,
                       include_risks, include_costs, include_timeline, custom_context):
    """Construit le prompt optimisÃ© pour la gÃ©nÃ©ration de rapport"""
    
    stats = mission_data['stats']
    sites = mission_data['sites']
    activities = mission_data['activities_breakdown']
    
    prompt = f"""Tu es un expert en rÃ©daction de rapports de mission professionnels. 

DONNÃ‰ES DE LA MISSION:
- DurÃ©e totale: {stats['total_days']} jour(s)
- Distance totale: {stats['total_km']:.1f} km
- Temps de visite total: {stats['total_visit_hours']:.1f} heures
- Nombre de sites visitÃ©s: {len([s for s in sites if s.get('Type') != 'Base'])}
- Sites visitÃ©s: {', '.join([s.get('Ville') for s in sites if s.get('Type') != 'Base'])}
- MÃ©thode de calcul: {mission_data['calculation_method']}

RÃ‰PARTITION DES ACTIVITÃ‰S:
{chr(10).join([f"- {act}: {hours:.1f}h" for act, hours in activities.items()])}

CONTEXTE SUPPLÃ‰MENTAIRE:
{custom_context if custom_context else "Aucun contexte spÃ©cifique fourni"}

INSTRUCTIONS:
- Type de rapport: {report_type}
- Ton: {tone}
- Inclure recommandations: {'Oui' if include_recommendations else 'Non'}
- Inclure analyse des risques: {'Oui' if include_risks else 'Non'}
- Inclure analyse des coÃ»ts: {'Oui' if include_costs else 'Non'}
- Inclure timeline dÃ©taillÃ©e: {'Oui' if include_timeline else 'Non'}

GÃ©nÃ¨re un rapport complet et structurÃ© en franÃ§ais, avec:
1. RÃ©sumÃ© exÃ©cutif
2. Objectifs et contexte
3. DÃ©roulement de la mission
4. RÃ©sultats et observations
5. Analyse des performances (temps, distances, efficacitÃ©)
{"6. Recommandations pour l'avenir" if include_recommendations else ""}
{"7. Analyse des risques identifiÃ©s" if include_risks else ""}
{"8. Analyse des coÃ»ts et budget" if include_costs else ""}
{"9. Timeline dÃ©taillÃ©e des activitÃ©s" if include_timeline else ""}
10. Conclusion

Utilise un style {tone.lower()} et structure le rapport avec des titres clairs et des sections bien organisÃ©es."""

    return prompt

def generate_pv_report(mission_data, questions_data, deepseek_api_key):
    """GÃ©nÃ¨re un rapport au format procÃ¨s-verbal professionnel avec l'IA Adja DeepSeek"""
    
    if not deepseek_api_key:
        return None, "ClÃ© API DeepSeek manquante"
    
    try:
        # Construction du prompt spÃ©cialisÃ© pour le procÃ¨s-verbal
        prompt = f"""Tu es un expert en rÃ©daction de procÃ¨s-verbaux professionnels pour des projets d'infrastructure. 
GÃ©nÃ¨re un procÃ¨s-verbal de visite de chantier dÃ©taillÃ© et professionnel au format officiel, basÃ© sur les informations suivantes :

INFORMATIONS DE LA MISSION :
- Date : {mission_data.get('date', 'Non spÃ©cifiÃ©e')}
- Lieu/Site : {mission_data.get('location', 'Non spÃ©cifiÃ©')}
- Objectif : {mission_data.get('objective', 'Non spÃ©cifiÃ©')}
- Participants : {', '.join(mission_data.get('participants', []))}
- DurÃ©e : {mission_data.get('duration', 'Non spÃ©cifiÃ©e')}

DÃ‰TAILS SUPPLÃ‰MENTAIRES :
- Contexte : {questions_data.get('context', 'Non spÃ©cifiÃ©')}
- Observations : {questions_data.get('observations', 'Non spÃ©cifiÃ©es')}
- ProblÃ¨mes identifiÃ©s : {questions_data.get('issues', 'Aucun')}
- Actions rÃ©alisÃ©es : {questions_data.get('actions', 'Non spÃ©cifiÃ©es')}
- Recommandations : {questions_data.get('recommendations', 'Aucune')}

STRUCTURE OBLIGATOIRE DU PROCÃˆS-VERBAL (respecter exactement cette numÃ©rotation) :

I. Cadre gÃ©nÃ©ral
   1. Cadre gÃ©nÃ©ral
      - Contexte du projet et objectifs gÃ©nÃ©raux
      - Cadre contractuel et rÃ©glementaire
      - Intervenants principaux du projet

   2. Objet de la mission
      - Motif prÃ©cis de la visite
      - PÃ©rimÃ¨tre d'intervention
      - Objectifs spÃ©cifiques de la mission

II. DÃ©roulement de la mission
   A. SITE DE [NOM DU SITE 1]
      - ReconnaÃ®tre l'Ã©quipe prÃ©sente dans le secteur concernÃ©
      - VÃ©rifier l'avancement des travaux (donner un pourcentage)
      - Faire un bilan, s'enquÃ©rir des Ã©ventuelles difficultÃ©s et contraintes
      - ApprÃ©cier la qualitÃ© des travaux rÃ©alisÃ©s
      - Donner des orientations pour la suite des travaux

   B. SITE DE [NOM DU SITE 2] (si applicable)
      - MÃªmes points que pour le site 1
      - SpÃ©cificitÃ©s du site

III. Bilan et recommandations
   A. Points positifs constatÃ©s
      - Ã‰lÃ©ments satisfaisants observÃ©s
      - Bonnes pratiques identifiÃ©es
      - Respect des dÃ©lais et procÃ©dures

   B. Points d'attention et difficultÃ©s
      - ProblÃ¨mes techniques identifiÃ©s
      - Contraintes rencontrÃ©es
      - Risques potentiels

   C. Recommandations et orientations
      - Actions correctives immÃ©diates
      - Mesures prÃ©ventives
      - Orientations pour la suite du projet

IV. Observations dÃ©taillÃ©es
   - Constats techniques prÃ©cis
   - Mesures et donnÃ©es relevÃ©es
   - Documentation photographique (mentionner si applicable)
   - Respect des normes de sÃ©curitÃ© et environnementales

CONSIGNES DE RÃ‰DACTION STRICTES :
- Style administratif formel et professionnel
- Terminologie technique prÃ©cise du BTP/infrastructure
- Phrases courtes et factuelles
- Ã‰viter absolument les opinions personnelles
- Utiliser le passÃ© composÃ© pour les actions rÃ©alisÃ©es
- Utiliser le prÃ©sent pour les constats
- NumÃ©rotation stricte avec chiffres romains et lettres
- Longueur : 1000-1500 mots minimum
- Inclure des donnÃ©es chiffrÃ©es quand possible (pourcentages, mesures, dÃ©lais)
- Mentionner les normes et rÃ©fÃ©rences techniques applicables

FORMAT DE PRÃ‰SENTATION :
- Titres en majuscules pour les sections principales
- Sous-titres avec numÃ©rotation claire
- Paragraphes structurÃ©s avec puces pour les listes
- Conclusion avec date et lieu de rÃ©daction

Le procÃ¨s-verbal doit Ãªtre conforme aux standards administratifs et prÃªt pour validation hiÃ©rarchique et archivage officiel."""

        # Appel Ã  l'API DeepSeek
        headers = {
            'Authorization': f'Bearer {deepseek_api_key}',
            'Content-Type': 'application/json'
        }
        
        data = {
            'model': 'deepseek-chat',
            'messages': [
                {
                    'role': 'user',
                    'content': prompt
                }
            ],
            'temperature': 0.3,  # Plus faible pour plus de cohÃ©rence
            'max_tokens': 2000
        }
        
        response = requests.post(
            'https://api.deepseek.com/chat/completions',
            headers=headers,
            json=data,
            timeout=30
        )
        
        if response.status_code == 200:
            result = response.json()
            if 'choices' in result and len(result['choices']) > 0:
                pv_content = result['choices'][0]['message']['content']
                return pv_content, None
            else:
                return None, "RÃ©ponse invalide de l'API DeepSeek"
        else:
            return None, f"Erreur API DeepSeek: {response.status_code} - {response.text}"
            
    except requests.exceptions.Timeout:
        return None, "Timeout lors de l'appel Ã  l'API DeepSeek"
    except requests.exceptions.RequestException as e:
        return None, f"Erreur de connexion Ã  l'API DeepSeek: {str(e)}"
    except Exception as e:
        return None, f"Erreur lors de la gÃ©nÃ©ration du PV: {str(e)}"

# --------------------------
# FONCTIONS UTILITAIRES
# --------------------------

def test_graphhopper_connection(api_key):
    """Teste la connexion Ã  GraphHopper"""
    if not api_key:
        return False, "ClÃ© API manquante"
    
    try:
        test_points = [[-17.4441, 14.6928], [-17.2732, 14.7167]]
        url = "https://graphhopper.com/api/1/matrix"
        
        data = {
            "points": test_points,
            "profile": "car",
            "out_arrays": ["times", "distances"]
        }
        
        headers = {"Content-Type": "application/json"}
        params = {"key": api_key}
        
        response = requests.post(url, json=data, params=params, headers=headers, timeout=10)
        
        if response.status_code == 200:
            result = response.json()
            if result and "times" in result and "distances" in result:
                distance_km = result['distances'][0][1] / 1000
                # Les "times" de GraphHopper Matrix sont gÃ©nÃ©ralement en secondes.
                # Si une valeur est trÃ¨s grande (>100000), on suppose des millisecondes.
                time_val = result['times'][0][1]
                time_min = (time_val / 1000 / 60) if time_val > 100000 else (time_val / 60)
                return True, f"Connexion OK - Test: {distance_km:.1f}km en {time_min:.0f}min"
            else:
                return False, "RÃ©ponse invalide de l'API"
        elif response.status_code == 401:
            return False, "ClÃ© API invalide"
        elif response.status_code == 429:
            return False, "Limite de requÃªtes atteinte"
        else:
            return False, f"Erreur HTTP {response.status_code}"
            
    except Exception as e:
        return False, f"Erreur: {str(e)}"

def _get_matrix_ttl_seconds():
    """TTL pour le cache de matrices (par dÃ©faut 24h)."""
    try:
        return int(st.secrets.get("MATRIX_TTL_SECONDS", 24 * 3600))
    except Exception:
        return 24 * 3600

@st.cache_data(ttl=_get_matrix_ttl_seconds(), show_spinner=False)
def improved_graphhopper_duration_matrix(api_key, coords):
    """Calcul de matrice via GraphHopper avec gestion d'erreurs"""
    if not api_key:
        return None, None, "ClÃ© API manquante"
    
    try:
        if len(coords) > 25:
            return None, None, f"Trop de points ({len(coords)}), limite: 25"
        
        # VÃ©rifier que toutes les coordonnÃ©es sont valides
        for i, coord in enumerate(coords):
            if not coord or len(coord) != 2:
                return None, None, f"CoordonnÃ©es invalides pour le point {i+1}"
            lon, lat = coord
            if not (-180 <= lon <= 180) or not (-90 <= lat <= 90):
                return None, None, f"CoordonnÃ©es hors limites pour le point {i+1}: ({lon}, {lat})"
        
        points = [[coord[0], coord[1]] for coord in coords]
        url = "https://graphhopper.com/api/1/matrix"
        data = {
            "points": points,
            "profile": "car",
            "out_arrays": ["times", "distances"]
        }
        headers = {"Content-Type": "application/json"}
        params = {"key": api_key}

        last_error = None
        for attempt in range(3):
            try:
                response = requests.post(url, json=data, params=params, headers=headers, timeout=30)
            except Exception as e:
                last_error = str(e)
                time_module.sleep(1 + attempt)
                continue
            
            if response.status_code == 200:
                result = response.json()
                times = result.get("times")
                distances = result.get("distances")
                if not times or not distances:
                    return None, None, "DonnÃ©es manquantes dans la rÃ©ponse"
                try:
                    flat_times = [t for row in times for t in row]
                    max_time = max(flat_times) if flat_times else 0
                except Exception:
                    max_time = 0
                durations = [[t / 1000.0 for t in row] for row in times] if max_time > 100000 else times
                return durations, distances, "SuccÃ¨s"
            else:
                if response.status_code == 401:
                    return None, None, "ClÃ© API invalide"
                elif response.status_code == 400:
                    try:
                        error_detail = response.json()
                        error_msg = error_detail.get('message', 'RequÃªte invalide')
                        return None, None, f"Erreur HTTP 400: {error_msg}. VÃ©rifiez que toutes les villes sont valides et gÃ©olocalisables."
                    except:
                        return None, None, "Erreur HTTP 400: RequÃªte invalide. VÃ©rifiez que toutes les villes sont valides et gÃ©olocalisables."
                elif response.status_code == 429:
                    last_error = "Limite de requÃªtes atteinte"
                    time_module.sleep(2 + attempt)
                    continue
                elif 500 <= response.status_code < 600:
                    last_error = f"Erreur HTTP {response.status_code}"
                    time_module.sleep(1 + attempt)
                    continue
                else:
                    return None, None, f"Erreur HTTP {response.status_code}"
        return None, None, f"Ã‰chec aprÃ¨s retries: {last_error or 'Erreur inconnue'}"
    except Exception as e:
        return None, None, f"Erreur: {str(e)}"

@st.cache_data(ttl=_get_matrix_ttl_seconds(), show_spinner=False)
def improved_osrm_duration_matrix(base_url, coords):
    """Calcul de matrice via OSRM Table avec gestion d'erreurs et fallback distances.
    Retourne (durations_sec, distances_m, message).
    """
    if not base_url:
        return None, None, "URL de base OSRM manquante"
    try:
        if len(coords) > 100:
            return None, None, f"Trop de points ({len(coords)}), limite recommandÃ©e: 100"
        # PrÃ©parer la chaÃ®ne de coordonnÃ©es pour l'API OSRM
        try:
            coord_str = ';'.join([f"{c[0]},{c[1]}" for c in coords])
        except Exception:
            return None, None, "CoordonnÃ©es invalides"
        url = f"{base_url.rstrip('/')}/table/v1/driving/{coord_str}"
        params = {"annotations": "duration,distance"}
        headers = {"Accept": "application/json"}

        last_error = None
        for attempt in range(3):
            try:
                response = requests.get(url, params=params, headers=headers, timeout=30)
            except Exception as e:
                last_error = str(e)
                time_module.sleep(1 + attempt)
                continue
            if response.status_code == 200:
                result = response.json()
                durations = result.get("durations")
                distances = result.get("distances")
                if durations is None:
                    return None, None, "DonnÃ©es manquantes: durations"
                # OSRM fournit les durÃ©es en secondes; distances en mÃ¨tres si activÃ©es
                if distances is None:
                    # Fallback distances via Haversine (corrigÃ© 1.2) si non fournies
                    n = len(coords)
                    distances = [[0.0]*n for _ in range(n)]
                    for i in range(n):
                        for j in range(n):
                            if i != j:
                                km = haversine(coords[i][0], coords[i][1], coords[j][0], coords[j][1]) * 1.2
                                distances[i][j] = km * 1000.0
                return durations, distances, "SuccÃ¨s"
            else:
                if response.status_code == 429:
                    last_error = "Limite de requÃªtes atteinte (OSRM)"
                    time_module.sleep(2 + attempt)
                    continue
                elif 500 <= response.status_code < 600:
                    last_error = f"Erreur HTTP {response.status_code} (OSRM)"
                    time_module.sleep(1 + attempt)
                    continue
                elif response.status_code == 400:
                    try:
                        err = response.json()
                        msg = err.get('message') or err.get('error') or 'RequÃªte invalide (OSRM)'
                        return None, None, f"Erreur HTTP 400: {msg}"
                    except Exception:
                        return None, None, "Erreur HTTP 400: RequÃªte invalide (OSRM)"
                else:
                    return None, None, f"Erreur HTTP {response.status_code} (OSRM)"
        return None, None, f"Ã‰chec aprÃ¨s retries: {last_error or 'Erreur inconnue'}"
    except Exception as e:
        return None, None, f"Erreur: {str(e)}"

def _get_deepseek_matrix_ttl_seconds():
    """TTL pour le cache de matrices DeepSeek (par dÃ©faut 6h)."""
    try:
        return int(st.secrets.get("DEEPSEEK_MATRIX_TTL_SECONDS", 6 * 3600))
    except Exception:
        return 6 * 3600

@st.cache_data(ttl=_get_deepseek_matrix_ttl_seconds(), show_spinner=False)
def improved_deepseek_estimate_matrix(cities, api_key, debug=False):
    """Estimation via DeepSeek avec distances exactes"""
    if not api_key:
        return None, "DeepSeek non disponible"
    
    try:
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }
        
        prompt = f"""Tu es un expert en transport routier au SÃ©nÃ©gal. Calcule les durÃ©es ET distances de trajet routier entre ces {len(cities)} villes: {', '.join(cities)}

DISTANCES EXACTES PAR ROUTE (Ã€ UTILISER - BIDIRECTIONNELLES):
- Dakar â†” ThiÃ¨s: 70 km (55-65 min)
- Dakar â†” Saint-Louis: 270 km (2h45-3h15)
- Dakar â†” Kaolack: 190 km (2h-2h30)
- ThiÃ¨s â†” Saint-Louis: 200 km (2h-2h30)
- ThiÃ¨s â†” Kaolack: 120 km (1h15-1h30)
- Saint-Louis â†” Kaolack: 240 km (2h30-3h)

IMPORTANT: Les distances sont identiques dans les deux sens (Aâ†’B = Bâ†’A).

RÃ©ponds uniquement en JSON:
{{
  "durations_minutes": [[matrice {len(cities)}x{len(cities)}]],
  "distances_km": [[matrice {len(cities)}x{len(cities)}]]
}}"""

        data = {
            "model": "deepseek-chat",
            "messages": [{"role": "user", "content": prompt}],
            "temperature": 0.1,
            "max_tokens": 2000
        }

        last_error = None
        for attempt in range(3):
            try:
                response = requests.post(
                    "https://api.deepseek.com/chat/completions",
                    headers=headers,
                    json=data,
                    timeout=30
                )
            except Exception as e:
                last_error = str(e)
                time_module.sleep(1 + attempt)
                continue
            
            if response.status_code != 200:
                if response.status_code == 429:
                    last_error = "Limite de requÃªtes atteinte"
                    time_module.sleep(2 + attempt)
                    continue
                elif 500 <= response.status_code < 600:
                    last_error = f"Erreur HTTP {response.status_code}"
                    time_module.sleep(1 + attempt)
                    continue
                else:
                    return None, f"Erreur API: {response.status_code}"

            result = response.json()
            text = result.get("choices", [{}])[0].get("message", {}).get("content", "")
            start = text.find("{")
            end = text.rfind("}") + 1
            if start >= 0 and end > start:
                try:
                    json_str = text[start:end]
                    parsed = json.loads(json_str)
                    minutes_matrix = parsed.get("durations_minutes", [])
                    km_matrix = parsed.get("distances_km", [])
                    seconds_matrix = [[int(m) * 60 for m in row] for row in minutes_matrix]
                    distances_matrix = [[int(km * 1000) for km in row] for row in km_matrix]
                    return (seconds_matrix, distances_matrix), "SuccÃ¨s DeepSeek"
                except Exception as parse_err:
                    return None, f"Format invalide: {parse_err}"
            else:
                last_error = "RÃ©ponse non JSON"
                time_module.sleep(1 + attempt)
                continue
        return None, f"Ã‰chec aprÃ¨s retries: {last_error or 'Erreur inconnue'}"
    
    except Exception as e:
        return None, f"Erreur: {str(e)}"

def build_ics_from_itinerary(itinerary, start_date, mission_title="Mission Terrain"):
    """Construit un fichier ICS Ã  partir du planning."""
    lines = [
        "BEGIN:VCALENDAR",
        "VERSION:2.0",
        "PRODID:-//Mission Planner//EN"
    ]
    now_str = datetime.now().strftime("%Y%m%dT%H%M%S")
    for idx, (day, sdt, edt, desc) in enumerate(itinerary):
        dtstart = sdt.strftime("%Y%m%dT%H%M%S")
        dtend = edt.strftime("%Y%m%dT%H%M%S")
        uid = f"mission-{now_str}-{idx}@planner"
        summary = f"{mission_title} - {desc}"
        lines.extend([
            "BEGIN:VEVENT",
            f"UID:{uid}",
            f"DTSTAMP:{now_str}",
            f"DTSTART:{dtstart}",
            f"DTEND:{dtend}",
            f"SUMMARY:{summary}",
            f"DESCRIPTION:Jour {day}",
            "END:VEVENT"
        ])
    lines.append("END:VCALENDAR")
    return "\r\n".join(lines)

@st.cache_data(show_spinner=False)
def _get_geolocator():
    """Ressource Geopy rÃ©utilisable (Ã©vite recrÃ©ation Ã  chaque appel)."""
    # Utilise un cache ressource pour conserver l'instance et respecter le rate limiter
    @st.cache_resource(show_spinner=False)
    def _cached_geolocator():
        return Nominatim(user_agent="mission-planner-senegal/2.0", timeout=10)
    return _cached_geolocator()

def _get_rate_limited_geocode():
    """Retourne une fonction geocode rate-limitÃ©e avec retries."""
    @st.cache_resource(show_spinner=False)
    def _cached_rate_limiter():
        geolocator = _get_geolocator()
        return RateLimiter(
            geolocator.geocode,
            min_delay_seconds=1,
            max_retries=3,
            error_wait_seconds=2,
            swallow_exceptions=False
        )
    return _cached_rate_limiter()

def _normalize_city_key(name: str) -> str:
    """Normalise un nom de ville pour les correspondances hors-ligne (sans accents/espaces/ponctuations)."""
    if not isinstance(name, str):
        return ""
    s = name.strip().lower()
    s = unicodedata.normalize("NFKD", s)
    s = s.encode("ascii", "ignore").decode("ascii")
    # Unifier les variantes de 'saint', 'ste', 'st'
    s = re.sub(r"\bste\b", "saint", s)
    s = re.sub(r"\bst\b", "saint", s)
    # Supprimer tout sauf alphanumÃ©rique
    s = re.sub(r"[^a-z0-9]", "", s)
    return s

# CoordonnÃ©es approximatives de grandes villes du SÃ©nÃ©gal (lon, lat)
SENEGAL_CITY_COORDS = {
    _normalize_city_key("Dakar"): (-17.4677, 14.7167),
    _normalize_city_key("Pikine"): (-17.3570, 14.7642),
    _normalize_city_key("Touba"): (-15.8833, 14.8667),
    _normalize_city_key("ThiÃ¨s"): (-16.9359, 14.7910),
    _normalize_city_key("Thies"): (-16.9359, 14.7910),  # variante sans accent
    _normalize_city_key("Saint-Louis"): (-16.4896, 16.0179),
    _normalize_city_key("Saint Louis"): (-16.4896, 16.0179),
    _normalize_city_key("St-Louis"): (-16.4896, 16.0179),
    _normalize_city_key("Kaolack"): (-16.0726, 14.1475),
    _normalize_city_key("Ziguinchor"): (-16.2719, 12.5833),
    _normalize_city_key("Louga"): (-16.2167, 15.6167),
    _normalize_city_key("Tambacounda"): (-13.6673, 13.7703),
    _normalize_city_key("Diourbel"): (-16.2348, 14.6550),
    _normalize_city_key("Fatick"): (-16.4150, 14.3390),
    _normalize_city_key("Kolda"): (-14.9500, 12.8833),
    _normalize_city_key("Matam"): (-13.2554, 15.6559),
    _normalize_city_key("Kaffrine"): (-15.5508, 14.1059),
    _normalize_city_key("Kedougou"): (-12.1742, 12.5556),
    _normalize_city_key("KÃ©dougou"): (-12.1742, 12.5556),
    _normalize_city_key("Sedhiou"): (-15.5569, 12.7081),
    _normalize_city_key("SÃ©dhiou"): (-15.5569, 12.7081),
    _normalize_city_key("Rufisque"): (-17.2729, 14.7158),
    _normalize_city_key("Mbour"): (-16.9600, 14.4361),
    _normalize_city_key("Richard-Toll"): (-15.6994, 16.4611),
    _normalize_city_key("Richard Toll"): (-15.6994, 16.4611),
}

def _offline_lookup_city_coords(city: str):
    key = _normalize_city_key(city)
    return SENEGAL_CITY_COORDS.get(key)

def _graphhopper_geocode(city: str):
    """Fallback via GraphHopper Geocoding API si disponible.
    SÃ©lectionne en prioritÃ© les lieux de type city/town/village au SÃ©nÃ©gal.
    """
    try:
        gh_key = globals().get("graphhopper_api_key")
        if not gh_key:
            return None
        url = "https://graphhopper.com/api/1/geocode"
        params = {
            "q": f"{city}, Senegal",
            "locale": "fr",
            "limit": 8,
            "key": gh_key,
        }
        resp = requests.get(url, params=params, timeout=10)
        if resp.status_code != 200:
            return None
        hits = (resp.json().get("hits") or [])
        if not hits:
            return None
        def is_sn(h):
            country = (h.get("country") or h.get("countrycode") or "").lower()
            return country in ("senegal", "sÃ©nÃ©gal", "sn")
        def is_place(h):
            return (h.get("osm_key", "").lower() == "place" and (h.get("osm_value", "").lower() in ("city", "town", "village")))
        chosen = None
        for h in hits:
            if is_sn(h) and is_place(h):
                chosen = h
                break
        if not chosen:
            for h in hits:
                if is_sn(h):
                    chosen = h
                    break
        chosen = chosen or hits[0]
        pt = chosen.get("point") or {}
        lat = pt.get("lat")
        lng = pt.get("lng")
        if lat is None or lng is None:
            return None
        # Valide que le point est dans un bbox raisonnable pour le SÃ©nÃ©gal
        if not (-17.8 <= float(lng) <= -11.0 and 12.0 <= float(lat) <= 16.9):
            return None
        return (float(lng), float(lat))
    except Exception:
        return None

def _geocode_city_senegal_raw(city: str):
    """ImplÃ©mentation brute sans cache (avec ressource + retries). PrivilÃ©gie coordonnÃ©es locales si activÃ©."""
    if not city or not isinstance(city, str) or not city.strip():
        return None

    # Option: prÃ©fÃ©rer coordonnÃ©es locales pour grandes villes (fiabilitÃ©)
    try:
        prefer_offline = st.session_state.get("prefer_offline_geocoding", True)
    except Exception:
        prefer_offline = True
    if prefer_offline:
        offline = _offline_lookup_city_coords(city)
        if offline:
            return offline

    last_error = None
    for attempt in range(3):  # 3 tentatives
        try:
            rate_limited = _get_rate_limited_geocode()
            query = f"{city}, SÃ©nÃ©gal" if "sÃ©nÃ©gal" not in city.lower() else city
            
            # Essai principal: ciblÃ© SÃ©nÃ©gal
            loc = rate_limited(query, language="fr", country_codes="SN")
            
            # Fallback: requÃªte gÃ©nÃ©rale
            if not loc:
                loc = rate_limited(city, language="fr")
            
            if loc:
                lon, lat = (loc.longitude, loc.latitude)
                # VÃ©rifie le type renvoyÃ© par Nominatim si disponible
                try:
                    raw = getattr(loc, "raw", {}) or {}
                    place_type = (raw.get("type") or "").lower()
                    if place_type and place_type not in ("city", "town", "village", "hamlet"):
                        raise ValueError(f"RÃ©sultat non-ville: {place_type}")
                except Exception:
                    pass
                # Valide que les coordonnÃ©es sont plausibles pour le SÃ©nÃ©gal
                if not (-17.8 <= float(lon) <= -11.0 and 12.0 <= float(lat) <= 16.9):
                    raise ValueError("CoordonnÃ©es hors SÃ©nÃ©gal")
                return (lon, lat)
        
        except ConnectionRefusedError as e:
            last_error = f"Connexion refusÃ©e au service de gÃ©ocodage. VÃ©rifiez votre connexion ou l'Ã©tat du service. ({e})"
            time_module.sleep(1 + attempt) # Attente progressive
            continue
        except Exception as e:
            last_error = e
            time_module.sleep(1 + attempt) # Attente progressive
            continue
    # Fallback 1: GraphHopper (si clÃ© dispo)
    gh_coords = _graphhopper_geocode(city)
    if gh_coords:
        st.warning(f"GÃ©ocodage Nominatim indisponible/inexact. Fallback GraphHopper utilisÃ© pour {city}.")
        return gh_coords

    # Fallback 2: Dictionnaire hors-ligne
    offline = _offline_lookup_city_coords(city)
    if offline:
        st.info(f"Mode hors-ligne: coordonnÃ©es vÃ©rifiÃ©es utilisÃ©es pour {city}.")
        return offline

    st.error(f"Erreur de gÃ©ocodage persistante pour {city} aprÃ¨s plusieurs tentatives: {last_error}")
    return None

def _get_geocode_ttl_seconds():
    # TTL configurable via secrets, dÃ©faut 7 jours
    try:
        return int(st.secrets.get("GEOCODE_TTL_SECONDS", 7 * 24 * 3600))
    except Exception:
        return 7 * 24 * 3600

@st.cache_data(ttl=_get_geocode_ttl_seconds(), show_spinner=False)
def _geocode_city_senegal_cached(city: str):
    return _geocode_city_senegal_raw(city)

def geocode_city_senegal(city: str, use_cache: bool = True):
    """GÃ©ocode une ville au SÃ©nÃ©gal, avec cache TTL et ressource partagÃ©e.

    Args:
        city: Nom de la ville
        use_cache: Active ou non le cache des rÃ©sultats
    """
    return _geocode_city_senegal_cached(city) if use_cache else _geocode_city_senegal_raw(city)

def solve_tsp_fixed_start_end(matrix):
    """RÃ©sout le TSP avec dÃ©part et arrivÃ©e fixes"""
    n = len(matrix)
    if n <= 2:
        return list(range(n))
    
    if n > 10:
        st.warning("Plus de 10 sites: heuristique voisin + 2-opt")
        nn_path = solve_tsp_nearest_neighbor(matrix)
        improved_path = two_opt_fixed_start_end(nn_path, matrix)
        return improved_path
    
    nodes = list(range(1, n-1))
    best_order = None
    best_time = float("inf")
    
    for perm in permutations(nodes):
        total_time = matrix[0][perm[0]]
        for i in range(len(perm)-1):
            total_time += matrix[perm[i]][perm[i+1]]
        total_time += matrix[perm[-1]][n-1]
        
        if total_time < best_time:
            best_time = total_time
            best_order = perm
    
    best_path = [0] + list(best_order) + [n-1]
    # Lissage via 2-opt si des incohÃ©rences existent
    try:
        best_path = two_opt_fixed_start_end(best_path, matrix)
    except Exception:
        pass
    return best_path

def solve_tsp_nearest_neighbor(matrix):
    """Heuristique du plus proche voisin"""
    n = len(matrix)
    unvisited = set(range(1, n-1))
    path = [0]
    current = 0
    
    while unvisited:
        nearest = min(unvisited, key=lambda x: matrix[current][x])
        path.append(nearest)
        unvisited.remove(nearest)
        current = nearest
    
    path.append(n-1)
    return path

def path_cost(path, matrix):
    total = 0
    for i in range(len(path)-1):
        total += matrix[path[i]][path[i+1]]
    return total

def two_opt_fixed_start_end(path, matrix):
    """AmÃ©lioration locale 2-opt en conservant dÃ©part (0) et arrivÃ©e (n-1)"""
    if not path or len(path) < 4:
        return path
    improved = True
    while improved:
        improved = False
        for i in range(1, len(path)-2):
            for k in range(i+1, len(path)-1):
                new_path = path[:i] + path[i:k+1][::-1] + path[k+1:]
                if path_cost(new_path, matrix) < path_cost(path, matrix):
                    path = new_path
                    improved = True
                    break
            if improved:
                break
    return path

# OR-Tools integration for advanced optimization (TSP with fixed start/end)
try:
    from ortools.constraint_solver import pywrapcp, routing_enums_pb2
    ORTOOLS_AVAILABLE = True
except Exception:
    ORTOOLS_AVAILABLE = False


def solve_tsp_ortools_fixed_start_end(matrix, service_times=None, time_limit_s=5):
    """Optimise l'ordre via OR-Tools (TSP), dÃ©part 0 et arrivÃ©e n-1 fixÃ©s.
    - Utilise la matrice des durÃ©es (secondes)
    - Peut intÃ©grer les durÃ©es de visite (service_times, en secondes) sur chaque site
    - Retourne un chemin sous forme d'indices: [0, ..., n-1]
    """
    n = len(matrix)
    if n <= 2:
        return list(range(n))

    # Si OR-Tools indisponible, fallback sur l'implÃ©mentation TSP existante
    if not ORTOOLS_AVAILABLE:
        return solve_tsp_fixed_start_end(matrix)

    try:
        manager = pywrapcp.RoutingIndexManager(n, 1, 0, n-1)
        routing = pywrapcp.RoutingModel(manager)

        def time_callback(from_index, to_index):
            from_node = manager.IndexToNode(from_index)
            to_node = manager.IndexToNode(to_index)
            travel = int(matrix[from_node][to_node] or 0)
            service = 0
            if service_times and isinstance(service_times, (list, tuple)):
                if 0 <= from_node < n and from_node not in (0, n - 1):
                    st_value = service_times[from_node]
                    try:
                        service = int(float(st_value) if st_value is not None else 0)
                    except Exception:
                        service = 0
            return travel + service

        transit_cb = routing.RegisterTransitCallback(time_callback)
        routing.SetArcCostEvaluatorOfAllVehicles(transit_cb)

        # Dimension de temps simple (horizon large)
        routing.AddDimension(
            transit_cb,
            0,            # marge
            24 * 3600,    # horizon max (24h)
            True,         # cumul de dÃ©part Ã  0
            "Time"
        )

        search_params = pywrapcp.DefaultRoutingSearchParameters()
        search_params.first_solution_strategy = routing_enums_pb2.FirstSolutionStrategy.PATH_CHEAPEST_ARC
        search_params.local_search_metaheuristic = routing_enums_pb2.LocalSearchMetaheuristic.GUIDED_LOCAL_SEARCH
        search_params.time_limit.seconds = int(time_limit_s)

        solution = routing.SolveWithParameters(search_params)
        if solution:
            index = routing.Start(0)
            path = []
            while not routing.IsEnd(index):
                node = manager.IndexToNode(index)
                path.append(node)
                index = solution.Value(routing.NextVar(index))
            path.append(manager.IndexToNode(index))
            return path
    except Exception:
        pass

    # Fallback en cas d'Ã©chec
    return solve_tsp_fixed_start_end(matrix)

def haversine(lon1, lat1, lon2, lat2):
    """Calcule la distance gÃ©odÃ©sique entre deux points en kilomÃ¨tres"""
    from math import radians, sin, cos, sqrt, atan2
    R = 6371.0
    dlon = radians(lon2 - lon1)
    dlat = radians(lat2 - lat1)
    a = sin(dlat/2)**2 + cos(radians(lat1))*cos(radians(lat2))*sin(dlon/2)**2
    c = 2 * atan2(sqrt(a), sqrt(1-a))
    return R * c

def haversine_fallback_matrix(coords, kmh=95.0):
    """Calcule une matrice basÃ©e sur distances gÃ©odÃ©siques"""
    
    n = len(coords)
    durations = [[0.0]*n for _ in range(n)]
    distances = [[0.0]*n for _ in range(n)]
    
    for i in range(n):
        for j in range(n):
            if i != j:
                km = haversine(coords[i][0], coords[i][1], coords[j][0], coords[j][1])
                # Facteur de correction pour tenir compte des routes rÃ©elles
                km *= 1.2
                hours = km / kmh
                # Retourner les durÃ©es en secondes (cohÃ©rent avec GraphHopper)
                durations[i][j] = hours * 3600
                # Retourner les distances en mÃ¨tres (cohÃ©rent avec GraphHopper)
                distances[i][j] = km * 1000
    
    return durations, distances

def optimize_route_with_ai(sites, coords, base_location=None, api_key=None):
    """
    Optimise l'ordre des sites en utilisant l'IA Adja DeepSeek
    
    Args:
        sites: Liste des sites avec leurs informations
        coords: Liste des coordonnÃ©es correspondantes
        base_location: Point de dÃ©part/arrivÃ©e (optionnel)
        api_key: ClÃ© API DeepSeek
    
    Returns:
        tuple: (ordre_optimal, success, message)
    """
    if not api_key:
        return list(range(len(sites))), False, "ClÃ© API DeepSeek manquante"
    
    try:
        # PrÃ©parer les donnÃ©es des sites pour l'IA Adja
        sites_info = []
        for i, site in enumerate(sites):
            site_data = {
                "index": i,
                "ville": site.get("Ville", f"Site {i}"),
                "type": site.get("Type", "Non spÃ©cifiÃ©"),
                "activite": site.get("ActivitÃ©", "Non spÃ©cifiÃ©"),
                "duree": site.get("DurÃ©e (h)", 1.0),
                "coordonnees": coords[i] if i < len(coords) else None
            }
            sites_info.append(site_data)
        
        # Construire le prompt pour l'IA Adja
        prompt = f"""Tu es un expert en optimisation d'itinÃ©raires au SÃ©nÃ©gal. 

MISSION: Optimise l'ordre de visite des sites suivants pour minimiser le temps de trajet total.

SITES Ã€ VISITER:
"""
        
        for site in sites_info:
            coord_str = f"({site['coordonnees'][0]:.4f}, {site['coordonnees'][1]:.4f})" if site['coordonnees'] else "CoordonnÃ©es inconnues"
            prompt += f"- Site {site['index']}: {site['ville']} - {site['type']} - {site['activite']} ({site['duree']}h) - {coord_str}\n"
        
        if base_location:
            prompt += f"\nPOINT DE DÃ‰PART/ARRIVÃ‰E: {base_location}\n"
        
        prompt += """
CONTRAINTES:
- Minimiser la distance totale de trajet
- Tenir compte de la gÃ©ographie du SÃ©nÃ©gal
- ConsidÃ©rer les types d'activitÃ©s (regrouper les activitÃ©s similaires si logique)
- Optimiser pour un trajet efficace

RÃ‰PONSE ATTENDUE:
Fournis UNIQUEMENT la liste des indices dans l'ordre optimal, sÃ©parÃ©s par des virgules.
Exemple: 0,2,1,3,4

Ne fournis AUCUNE explication, juste la sÃ©quence d'indices."""

        # Appel Ã  l'API DeepSeek
        headers = {
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json"
        }
        
        data = {
            "model": "deepseek-chat",
            "messages": [
                {"role": "user", "content": prompt}
            ],
            "temperature": 0.1,
            "max_tokens": 100
        }
        
        response = requests.post(
            "https://api.deepseek.com/chat/completions",
            headers=headers,
            json=data,
            timeout=30
        )
        
        if response.status_code == 200:
            result = response.json()
            ai_response = result["choices"][0]["message"]["content"].strip()
            
            # Parser la rÃ©ponse de l'IA Adja
            try:
                # Extraire les indices de la rÃ©ponse
                indices_str = ai_response.split('\n')[0].strip()
                indices = [int(x.strip()) for x in indices_str.split(',')]
                
                # VÃ©rifier que tous les indices sont valides
                if len(indices) == len(sites) and set(indices) == set(range(len(sites))):
                    return indices, True, "Optimisation IA Adja rÃ©ussie"
                else:
                    # Fallback: ordre original si la rÃ©ponse IA Adja est invalide
                    return list(range(len(sites))), False, f"RÃ©ponse IA Adja invalide: {ai_response[:100]}..."
                    
            except (ValueError, IndexError) as e:
                return list(range(len(sites))), False, f"Erreur parsing rÃ©ponse IA Adja: {str(e)}"
        
        else:
            return list(range(len(sites))), False, f"Erreur API DeepSeek: {response.status_code}"
            
    except requests.exceptions.Timeout:
        return list(range(len(sites))), False, "Timeout API DeepSeek"
    except requests.exceptions.RequestException as e:
        return list(range(len(sites))), False, f"Erreur rÃ©seau: {str(e)}"
    except Exception as e:
        return list(range(len(sites))), False, f"Erreur inattendue: {str(e)}"

def schedule_itinerary(coords, sites, order, segments_summary,
                       start_date, start_activity_time, end_activity_time,
                       start_travel_time, end_travel_time,
                       use_lunch, lunch_start_time, lunch_end_time,
                       use_prayer, prayer_start_time, prayer_duration_min,
                       max_days=0, tolerance_hours=1.0, base_location=None, 
                       stretch_days=False, end_day_early_threshold=1.5,
                       allow_weekend_travel=True, allow_weekend_activities=True,
                       lunch_duration_min=60):
    """GÃ©nÃ¨re le planning dÃ©taillÃ© avec horaires diffÃ©renciÃ©s pour activitÃ©s et voyages"""
    sites_ordered = [sites[i] for i in order]
    coords_ordered = [coords[i] for i in order]
    
    current_datetime = datetime.combine(start_date, start_travel_time)  # Start with travel time
    day_end_time = datetime.combine(start_date, end_travel_time)  # End with travel time
    day_count = 1
    itinerary = []
    
    # Suivi des pauses par jour pour Ã©viter les doublons
    daily_lunch_added = {}  # {day_count: bool}
    daily_prayer_added = {}  # {day_count: bool}
    
    total_km = 0
    total_visit_hours = 0
    
    for idx, site in enumerate(sites_ordered):
        # Handle travel to this site (except for first site)
        if idx > 0:
            # Weekend skip for travel if disabled
            if not allow_weekend_travel:
                while current_datetime.weekday() >= 5:
                    itinerary.append((day_count, current_datetime, datetime.combine(current_datetime.date(), end_travel_time), "â›±ï¸ Week-end (pas de voyage)"))
                    day_count += 1
                    current_datetime = datetime.combine(start_date + timedelta(days=day_count-1), start_travel_time)
                    day_end_time = datetime.combine(start_date + timedelta(days=day_count-1), end_travel_time)
            seg_idx = idx - 1
            seg_idx = idx - 1
            if seg_idx < len(segments_summary):
                seg = segments_summary[seg_idx]
                travel_sec = seg.get("duration", 0)
                travel_km = seg.get("distance", 0) / 1000.0
                
                # Debug: Afficher les valeurs reÃ§ues
                if debug_mode:
                    st.info(f"ðŸ” Debug Segment {seg_idx}: travel_sec={travel_sec}, travel_km={travel_km:.2f}")
                
                # Si les donnÃ©es sont nulles, utiliser des valeurs par dÃ©faut simples
                if travel_sec <= 0:
                    travel_sec = 3600  # 1 heure par dÃ©faut
                    if debug_mode:
                        st.warning(f"ðŸ” travel_sec Ã©tait â‰¤ 0, fixÃ© Ã  3600s (1h)")
                if travel_km <= 0:
                    travel_km = 50  # 50 km par dÃ©faut
                    if debug_mode:
                        st.warning(f"ðŸ” travel_km Ã©tait â‰¤ 0, fixÃ© Ã  50km")
                
                total_km += travel_km
                
                travel_duration = timedelta(seconds=int(travel_sec))
                travel_end = current_datetime + travel_duration
                
                from_city = sites_ordered[idx-1]['Ville']
                to_city = site['Ville']
                
                # Format travel time for display
                travel_hours = travel_sec / 3600
                if travel_hours >= 1:
                    travel_time_str = f"{travel_hours:.1f}h"
                else:
                    travel_minutes = travel_sec / 60
                    travel_time_str = f"{travel_minutes:.0f}min"
                
                travel_desc = f"ðŸš— {from_city} â†’ {to_city} ({travel_km:.1f} km, {travel_time_str})"
                
                # Check if travel extends beyond travel hours
                travel_end_time = datetime.combine(current_datetime.date(), end_travel_time)
                
                if travel_end > travel_end_time:
                    # Travel extends beyond allowed hours - split across days
                    itinerary.append((day_count, current_datetime, travel_end_time, "ðŸ Fin de journÃ©e"))
                    prev_site = sites_ordered[idx-1]
                    prev_city = prev_site['Ville']
                    prev_overnight_allowed = prev_site.get('PossibilitÃ© de nuitÃ©e', True)
                    if prev_overnight_allowed:
                        itinerary.append((day_count, travel_end_time, travel_end_time, f"ðŸ¨ NuitÃ©e Ã  {prev_city}"))
                    else:
                        # Pas d'hÃ©bergement autorisÃ© Ã  la ville prÃ©cÃ©dente -> avertissement + nuitÃ©e de repli
                        itinerary.append((day_count, travel_end_time, travel_end_time, f"âš ï¸ DÃ©placement nÃ©cessaire - pas d'hÃ©bergement Ã  {prev_city}"))
                        fallback_city = None
                        for j in range(idx, len(sites_ordered)):
                            if sites_ordered[j].get('PossibilitÃ© de nuitÃ©e', True):
                                fallback_city = sites_ordered[j]['Ville']
                                break
                        if not fallback_city and base_location:
                            fallback_city = base_location
                        if fallback_city:
                            itinerary.append((day_count, travel_end_time, travel_end_time, f"ðŸ¨ NuitÃ©e Ã  {fallback_city}"))
                    
                    day_count += 1
                    current_datetime = datetime.combine(start_date + timedelta(days=day_count-1), start_travel_time)
                    day_end_time = datetime.combine(start_date + timedelta(days=day_count-1), end_travel_time)
                    travel_end = current_datetime + travel_duration
                
                # Handle lunch break during travel
                lunch_window_start = datetime.combine(current_datetime.date(), lunch_start_time) if use_lunch else None
                lunch_window_end = datetime.combine(current_datetime.date(), lunch_end_time) if use_lunch else None
                
                travel_added = False
                
                if use_lunch and lunch_window_start and lunch_window_end and not daily_lunch_added.get(day_count, False):
                    if current_datetime < lunch_window_end and travel_end > lunch_window_start:
                        # Si l'arrivÃ©e se situe dans la fenÃªtre de dÃ©jeuner, placer la pause Ã  l'arrivÃ©e
                        if lunch_window_start <= travel_end <= lunch_window_end:
                            # Ajouter le trajet en une seule fois jusqu'Ã  l'arrivÃ©e
                            itinerary.append((day_count, current_datetime, travel_end, travel_desc))
                            travel_added = True
                            
                            # Placer le dÃ©jeuner immÃ©diatement Ã  l'arrivÃ©e
                            lunch_time = max(travel_end, lunch_window_start)
                            lunch_end_time_actual = min(lunch_time + timedelta(minutes=lunch_duration_min), lunch_window_end)
                            desc_text = f"ðŸ½ï¸ DÃ©jeuner (â‰¤{lunch_duration_min} min)"
                            if use_prayer and prayer_start_time and not daily_prayer_added.get(day_count, False):
                                prayer_window_start = datetime.combine(lunch_time.date(), prayer_start_time)
                                prayer_window_end = prayer_window_start + timedelta(hours=2)
                                if lunch_time < prayer_window_end and lunch_end_time_actual > prayer_window_start:
                                    desc_text = f"ðŸ½ï¸ DÃ©jeuner (â‰¤{lunch_duration_min} min) + ðŸ™ PriÃ¨re (â‰¤{prayer_duration_min} min)"
                                    daily_prayer_added[day_count] = True
                            itinerary.append((day_count, lunch_time, lunch_end_time_actual, desc_text))
                            daily_lunch_added[day_count] = True
                            
                            # Mettre Ã  jour l'heure courante Ã  la fin du dÃ©jeuner
                            current_datetime = lunch_end_time_actual
                            # Le trajet est terminÃ©, Ã©viter tout ajout rÃ©siduel
                            travel_end = current_datetime
                        else:
                            # Sinon, conserver lâ€™ancienne logique (pause pendant le trajet)
                            lunch_time = max(current_datetime, lunch_window_start)
                            lunch_end_time_actual = min(lunch_time + timedelta(minutes=lunch_duration_min), lunch_window_end)
                            
                            # Ajouter la partie de trajet avant la pause si nÃ©cessaire
                            if lunch_time > current_datetime:
                                itinerary.append((day_count, current_datetime, lunch_time, travel_desc))
                                travel_added = True
                            
                            # Ajouter la pause dÃ©jeuner
                            itinerary.append((day_count, lunch_time, lunch_end_time_actual, f"ðŸ½ï¸ DÃ©jeuner (â‰¤{lunch_duration_min} min)"))
                            daily_lunch_added[day_count] = True
                            
                            # Reprendre le trajet aprÃ¨s la pause
                            current_datetime = lunch_end_time_actual
                            remaining_travel = travel_end - lunch_time
                            if remaining_travel.total_seconds() < 0:
                                remaining_travel = timedelta(seconds=0)
                            travel_end = current_datetime + remaining_travel
                
                # Handle prayer break during travel (only if no lunch break)
                elif use_prayer and prayer_start_time and not daily_prayer_added.get(day_count, False):
                    prayer_window_start = datetime.combine(current_datetime.date(), prayer_start_time)
                    prayer_window_end = prayer_window_start + timedelta(hours=2)
                    
                    if current_datetime < prayer_window_end and travel_end > prayer_window_start:
                        prayer_time = max(current_datetime, prayer_window_start)
                        prayer_end_time = prayer_time + timedelta(minutes=prayer_duration_min)
                        
                        if prayer_end_time > prayer_window_end:
                            prayer_end_time = prayer_window_end
                        
                        # Add travel before prayer if needed
                        if prayer_time > current_datetime:
                            itinerary.append((day_count, current_datetime, prayer_time, travel_desc))
                            travel_added = True
                        
                        # Add prayer break
                        itinerary.append((day_count, prayer_time, prayer_end_time, "ðŸ™ PriÃ¨re (â‰¤20 min)"))
                        daily_prayer_added[day_count] = True  # Marquer la priÃ¨re comme ajoutÃ©e pour ce jour
                        current_datetime = prayer_end_time
                        
                        # Recalculate remaining travel time
                        remaining_travel = travel_end - prayer_time
                        travel_end = current_datetime + remaining_travel
                
                # Add remaining travel time (include post-break remaining travel if any)
                if current_datetime < travel_end:
                    itinerary.append((day_count, current_datetime, travel_end, travel_desc))
                
                current_datetime = travel_end
        
        visit_hours = float(site.get("DurÃ©e (h)", 0)) if site.get("DurÃ©e (h)") else 0
        
        if visit_hours > 0:
            # Weekend skip for activities if disabled
            if not allow_weekend_activities:
                while current_datetime.weekday() >= 5:
                    itinerary.append((day_count, current_datetime, datetime.combine(current_datetime.date(), end_activity_time), "â›±ï¸ Week-end (pas d'activitÃ©s)"))
                    day_count += 1
                    current_datetime = datetime.combine(start_date + timedelta(days=day_count-1), start_activity_time)
                    day_end_time = datetime.combine(start_date + timedelta(days=day_count-1), end_travel_time)
            total_visit_hours += visit_hours
            visit_duration = timedelta(hours=visit_hours)
            visit_end = current_datetime + visit_duration
            
            type_site = site.get('Type', 'Site')
            activite = site.get('ActivitÃ©', 'Visite')
            city = site['Ville'].upper()
            
            visit_desc = f"{city} â€“ {activite}"
            if type_site not in ["Base"]:
                visit_desc = f"{city} â€“ Visite {type_site}"
            
            # Check if visit extends beyond activity hours
            activity_end_time = datetime.combine(current_datetime.date(), end_activity_time)
            tolerance_end_time = activity_end_time + timedelta(hours=tolerance_hours)
            
            # VÃ©rifier si l'activitÃ© peut continuer (nouvelle option)
            can_continue = site.get('Peut continuer', False)  # Par dÃ©faut False
            
            # VÃ©rifier si la nuitÃ©e est possible dans cette zone
            overnight_allowed = site.get('PossibilitÃ© de nuitÃ©e', True)  # Par dÃ©faut True
            
            # Handle visit that extends beyond activity hours
            if visit_end > activity_end_time:
                # Si l'activitÃ© se termine dans le seuil de tolÃ©rance, elle peut continuer le mÃªme jour
                if visit_end <= tolerance_end_time and can_continue:
                    # L'activitÃ© continue sur le mÃªme jour malgrÃ© le dÃ©passement
                    pass  # Pas de division, traitement normal
                elif can_continue and overnight_allowed:
                    # L'activitÃ© dÃ©passe le seuil de tolÃ©rance et peut Ãªtre divisÃ©e, ET la nuitÃ©e est autorisÃ©e
                    if current_datetime < activity_end_time:
                        # Add partial visit for current day
                        itinerary.append((day_count, current_datetime, activity_end_time, f"{visit_desc} (Ã  continuer)"))
                    
                    # End current day
                    itinerary.append((day_count, activity_end_time, activity_end_time, "ðŸ Fin de journÃ©e"))
                    if overnight_allowed:
                        itinerary.append((day_count, activity_end_time, activity_end_time, f"ðŸ¨ NuitÃ©e Ã  {city}"))
                    else:
                        itinerary.append((day_count, activity_end_time, activity_end_time, f"âš ï¸ DÃ©placement nÃ©cessaire - pas d'hÃ©bergement Ã  {city}"))
                        # NuitÃ©e de repli vers un site prochain autorisÃ© ou la base
                        fallback_city = None
                        for j in range(idx+1, len(sites_ordered)):
                            if sites_ordered[j].get('PossibilitÃ© de nuitÃ©e', True):
                                fallback_city = sites_ordered[j]['Ville']
                                break
                        if not fallback_city and base_location:
                            fallback_city = base_location
                        if fallback_city:
                            itinerary.append((day_count, activity_end_time, activity_end_time, f"ðŸ¨ NuitÃ©e Ã  {fallback_city}"))
                    
                    # Start next day
                    remaining = visit_end - activity_end_time
                    day_count += 1
                    current_datetime = datetime.combine(start_date + timedelta(days=day_count-1), start_activity_time)
                    day_end_time = datetime.combine(start_date + timedelta(days=day_count-1), end_travel_time)
                    visit_end = current_datetime + remaining
                    visit_desc = f"Suite {visit_desc}"
                elif can_continue and not overnight_allowed:
                    # L'activitÃ© peut continuer mais la nuitÃ©e n'est pas autorisÃ©e - chercher un site proche avec nuitÃ©e
                    # Pour l'instant, on force la fin de l'activitÃ© et on ajoute un avertissement
                    visit_end = activity_end_time
                    if current_datetime < activity_end_time:
                        itinerary.append((day_count, current_datetime, activity_end_time, f"{visit_desc} (interrompu - pas de nuitÃ©e possible)"))
                    
                    # End current day et chercher un hÃ©bergement ailleurs
                    itinerary.append((day_count, activity_end_time, activity_end_time, "ðŸ Fin de journÃ©e"))
                    itinerary.append((day_count, activity_end_time, activity_end_time, f"âš ï¸ DÃ©placement nÃ©cessaire - pas d'hÃ©bergement Ã  {city}"))
                    
                    # Start next day
                    day_count += 1
                    current_datetime = datetime.combine(start_date + timedelta(days=day_count-1), start_activity_time)
                    day_end_time = datetime.combine(start_date + timedelta(days=day_count-1), end_travel_time)
                    # Reprendre l'activitÃ© restante le jour suivant
                    remaining_hours = (visit_duration - (activity_end_time - current_datetime)).total_seconds() / 3600
                    if remaining_hours > 0:
                        visit_end = current_datetime + timedelta(hours=remaining_hours)
                        visit_desc = f"Suite {visit_desc}"
                    else:
                        # L'activitÃ© Ã©tait dÃ©jÃ  terminÃ©e
                        visit_end = current_datetime
                else:
                    # L'activitÃ© ne peut pas continuer - la forcer Ã  se terminer Ã  l'heure limite
                    visit_end = activity_end_time
                    if current_datetime >= activity_end_time:
                        # Si on est dÃ©jÃ  en dehors des heures, terminer la journÃ©e et ajouter la nuitÃ©e (avec fallback si nÃ©cessaire)
                        itinerary.append((day_count, current_datetime, current_datetime, "ðŸ Fin de journÃ©e"))
                        if overnight_allowed:
                            itinerary.append((day_count, current_datetime, current_datetime, f"ðŸ¨ NuitÃ©e Ã  {city}"))
                        else:
                            itinerary.append((day_count, current_datetime, current_datetime, f"âš ï¸ DÃ©placement nÃ©cessaire - pas d'hÃ©bergement Ã  {city}"))
                            # Chercher une nuitÃ©e autorisÃ©e dans les sites suivants ou la base
                            fallback_city = None
                            for j in range(idx+1, len(sites_ordered)):
                                if sites_ordered[j].get('PossibilitÃ© de nuitÃ©e', True):
                                    fallback_city = sites_ordered[j]['Ville']
                                    break
                            if not fallback_city and base_location:
                                fallback_city = base_location
                            if fallback_city:
                                itinerary.append((day_count, current_datetime, current_datetime, f"ðŸ¨ NuitÃ©e Ã  {fallback_city}"))
                        
                        # Reporter au jour suivant
                        day_count += 1
                        current_datetime = datetime.combine(start_date + timedelta(days=day_count-1), start_activity_time)
                        day_end_time = datetime.combine(start_date + timedelta(days=day_count-1), end_travel_time)
                        visit_end = current_datetime + visit_duration
            
            # Handle breaks during visit (only if visit fits in current day)
            if visit_end <= activity_end_time:
                lunch_window_start = datetime.combine(current_datetime.date(), lunch_start_time) if use_lunch else None
                lunch_window_end = datetime.combine(current_datetime.date(), lunch_end_time) if use_lunch else None
                
                prayer_window_start = datetime.combine(current_datetime.date(), prayer_start_time) if use_prayer else None
                prayer_window_end = prayer_window_start + timedelta(hours=2) if use_prayer else None
                
                # Check for lunch break during visit â€” do not split, schedule lunch after visit
                place_lunch_after_visit = False
                if use_lunch and lunch_window_start and lunch_window_end and not daily_lunch_added.get(day_count, False):
                    if current_datetime < lunch_window_end and visit_end > lunch_window_start:
                        place_lunch_after_visit = True
                # If lunch will be placed after the visit, and prayer window overlaps that lunch window,
                # combine prayer with lunch instead of splitting the visit
                combine_prayer_with_lunch = False
                if place_lunch_after_visit and use_prayer and prayer_window_start and prayer_window_end and not daily_prayer_added.get(day_count, False):
                    planned_lunch_start = max(visit_end, lunch_window_start)
                    planned_lunch_end = min(planned_lunch_start + timedelta(minutes=lunch_duration_min), lunch_window_end)
                    if planned_lunch_start < prayer_window_end and planned_lunch_end > prayer_window_start:
                        combine_prayer_with_lunch = True
                
                # Check for prayer break during visit (skip if it will be combined with lunch after visit)
                if use_prayer and prayer_window_start and prayer_window_end and not daily_prayer_added.get(day_count, False) and not combine_prayer_with_lunch:
                    if current_datetime < prayer_window_end and visit_end > prayer_window_start:
                        prayer_time = max(current_datetime, prayer_window_start)
                        prayer_end_time = min(prayer_time + timedelta(minutes=prayer_duration_min), prayer_window_end)
                        
                        # Add visit part before prayer
                        if prayer_time > current_datetime:
                            itinerary.append((day_count, current_datetime, prayer_time, visit_desc))
                        
                        # Add prayer break
                        itinerary.append((day_count, prayer_time, prayer_end_time, "ðŸ™ PriÃ¨re (â‰¤20 min)"))
                        daily_prayer_added[day_count] = True  # Marquer la priÃ¨re comme ajoutÃ©e pour ce jour
                        
                        # Update timing for remaining visit
                        current_datetime = prayer_end_time
                        remaining_visit = visit_end - prayer_time
                        visit_end = current_datetime + remaining_visit
                        visit_desc = f"Suite {visit_desc}" if prayer_time > current_datetime else visit_desc
            
            # Add final visit segment
            if current_datetime < visit_end:
                itinerary.append((day_count, current_datetime, visit_end, visit_desc))
                current_datetime = visit_end
            
            # Place lunch right after the visit if the window overlapped
            if 'place_lunch_after_visit' in locals() and place_lunch_after_visit and not daily_lunch_added.get(day_count, False):
                lunch_time = max(current_datetime, lunch_window_start)
                if lunch_time < lunch_window_end:
                    lunch_end_time_actual = min(lunch_time + timedelta(minutes=lunch_duration_min), lunch_window_end)
                    desc_text = f"ðŸ½ï¸ DÃ©jeuner (â‰¤{lunch_duration_min} min)"
                    if 'combine_prayer_with_lunch' in locals() and combine_prayer_with_lunch and use_prayer and not daily_prayer_added.get(day_count, False):
                        desc_text = f"ðŸ½ï¸ DÃ©jeuner (â‰¤{lunch_duration_min} min) + ðŸ™ PriÃ¨re (â‰¤{prayer_duration_min} min)"
                        daily_prayer_added[day_count] = True
                    itinerary.append((day_count, lunch_time, lunch_end_time_actual, desc_text))
                    daily_lunch_added[day_count] = True
                    current_datetime = lunch_end_time_actual
            
            # Check if we need to end the day early
            time_until_end = (day_end_time - current_datetime).total_seconds() / 3600
            
            # Si on doit Ã©taler, on termine la journÃ©e plus tÃ´t pour rÃ©partir sur plus de jours
            if stretch_days and day_count < max_days and idx < len(sites_ordered) - 1:
                itinerary.append((day_count, current_datetime, current_datetime, f"ðŸ Fin de journÃ©e"))
                # NuitÃ©e conditionnelle selon la possibilitÃ©
                if overnight_allowed:
                    itinerary.append((day_count, current_datetime, current_datetime, f"ðŸ¨ NuitÃ©e Ã  {city}"))
                else:
                    itinerary.append((day_count, current_datetime, current_datetime, f"âš ï¸ DÃ©placement nÃ©cessaire - pas d'hÃ©bergement Ã  {city}"))
                    # Chercher une nuitÃ©e autorisÃ©e dans les sites suivants ou la base
                    fallback_city = None
                    for j in range(idx+1, len(sites_ordered)):
                        if sites_ordered[j].get('PossibilitÃ© de nuitÃ©e', True):
                            fallback_city = sites_ordered[j]['Ville']
                            break
                    if not fallback_city and base_location:
                        fallback_city = base_location
                    if fallback_city:
                        itinerary.append((day_count, current_datetime, current_datetime, f"ðŸ¨ NuitÃ©e Ã  {fallback_city}"))

                # DÃ©marrer le jour suivant
                day_count += 1
                current_datetime = datetime.combine(start_date + timedelta(days=day_count-1), start_activity_time)
                day_end_time = datetime.combine(start_date + timedelta(days=day_count-1), end_travel_time)

            elif time_until_end <= end_day_early_threshold and idx < len(sites_ordered) - 1:
                # End current day and prepare for next day
                itinerary.append((day_count, current_datetime, current_datetime, f"ðŸ Fin de journÃ©e"))
                # NuitÃ©e conditionnelle selon la possibilitÃ©
                if overnight_allowed:
                    itinerary.append((day_count, current_datetime, current_datetime, f"ðŸ¨ NuitÃ©e Ã  {city}"))
                else:
                    itinerary.append((day_count, current_datetime, current_datetime, f"âš ï¸ DÃ©placement nÃ©cessaire - pas d'hÃ©bergement Ã  {city}"))
                    # Chercher une nuitÃ©e autorisÃ©e dans les sites suivants ou la base
                    fallback_city = None
                    for j in range(idx+1, len(sites_ordered)):
                        if sites_ordered[j].get('PossibilitÃ© de nuitÃ©e', True):
                            fallback_city = sites_ordered[j]['Ville']
                            break
                    if not fallback_city and base_location:
                        fallback_city = base_location
                    if fallback_city:
                        itinerary.append((day_count, current_datetime, current_datetime, f"ðŸ¨ NuitÃ©e Ã  {fallback_city}"))
                
                # Start next day
                day_count += 1
                current_datetime = datetime.combine(start_date + timedelta(days=day_count-1), start_activity_time)
                day_end_time = datetime.combine(start_date + timedelta(days=day_count-1), end_travel_time)
    
    # Add final overnight stay for the last day
    if day_count > 0 and sites_ordered:
        last_site = sites_ordered[-1]
        last_city = last_site['Ville']
        if last_site.get('PossibilitÃ© de nuitÃ©e', True):
            itinerary.append((day_count, current_datetime, current_datetime, f"ðŸ¨ NuitÃ©e Ã  {last_city}"))
        else:
            # Fallback to base_location if overnight is not possible at the last site
            if base_location:
                itinerary.append((day_count, current_datetime, current_datetime, f"ðŸ¨ NuitÃ©e Ã  {base_location}"))

    # Add final arrival marker
    if day_count > 0 and sites_ordered:
        last_city = sites_ordered[-1]['Ville'].upper()
        
        itinerary.append((day_count, current_datetime, current_datetime, f"ðŸ“ ArrivÃ©e {last_city} â€“ Fin de mission"))
    
    # Message d'avertissement si le nombre de jours est dÃ©passÃ© (ne devrait plus arriver avec la nouvelle logique)
    if max_days > 0 and day_count > max_days and not stretch_days:
        st.warning(f"âš ï¸ L'itinÃ©raire nÃ©cessite {day_count} jours, mais le maximum Ã©tait fixÃ© Ã  {max_days}. Le planning est compressÃ©.")
    
    stats = {
        "total_days": day_count,
        "total_km": total_km,
        "total_visit_hours": total_visit_hours
    }
    
    return itinerary, sites_ordered, coords_ordered, stats

def build_professional_html(itinerary, start_date, stats, sites_ordered, segments_summary=None, speed_kmh=110, mission_title="Mission Terrain", coords_ordered=None, include_map=False, lunch_start_time=None, lunch_end_time=None, lunch_duration_min=60, prayer_start_time=None, prayer_duration_min=20, include_details=True):
    """GÃ©nÃ¨re un HTML professionnel"""
    def fmt_time(dt):
        return dt.strftime("%Hh%M")
    
    def extract_distance_from_desc(desc, speed_kmh_param):
        import re
        # Chercher d'abord le format avec temps rÃ©el : "(123.4 km, 2h30)"
        m_with_time = re.search(r"\(([\d\.]+)\s*km,\s*([^)]+)\)", desc)
        if m_with_time:
            km = float(m_with_time.group(1))
            time_str = m_with_time.group(2).strip()
            return f"~{int(km)} km / ~{time_str}"
        
        # Fallback : ancien format avec seulement distance "(123.4 km)"
        m = re.search(r"\(([\d\.]+)\s*km\)", desc)
        if m:
            km = float(m.group(1))
            hours = km / speed_kmh_param
            h = int(hours)
            minutes = int((hours - h) * 60)
            if h > 0:
                time_str = f"{h}h{minutes:02d}"
            else:
                time_str = f"0h{minutes:02d}"
            return f"~{int(km)} km / ~{time_str}"
        return "-"

    by_day = {}
    night_locations = {}
    
    for day, sdt, edt, desc in itinerary:
        by_day.setdefault(day, []).append((sdt, edt, desc))
        
        if "NuitÃ©e Ã " in desc or "nuitÃ©e Ã " in desc:
            if " Ã  " in desc:
                parts = desc.split(" Ã  ")
                if len(parts) >= 2:
                    city = parts[1].strip().split("(")[0].strip().split(" ")[0]
                    night_locations[day] = city.upper()
        elif "ðŸ¨" in desc and ("NuitÃ©e" in desc or "nuitÃ©e" in desc):
            # Gestion spÃ©cifique pour les descriptions avec emoji ðŸ¨
            if " Ã  " in desc:
                parts = desc.split(" Ã  ")
                if len(parts) >= 2:
                    city = parts[1].strip().split("(")[0].strip().split(" ")[0]
                    night_locations[day] = city.upper()
        elif "installation" in desc.lower() and "nuitÃ©e" in desc.lower():
            words = desc.split()
            for i, word in enumerate(words):
                if "installation" in word.lower() and i + 1 < len(words):
                    city = words[i + 1].strip().split("(")[0].strip()
                    night_locations[day] = city.upper()
                    break
        elif "Fin de journÃ©e" in desc:
            for _, _, d in reversed(by_day[day]):
                if any(x in d for x in ["VISITE", "Visite", "â€“"]) and "â†’" not in d:
                    if "â€“" in d:
                        city = d.split("â€“")[0].strip()
                        night_locations[day] = city.upper()
                        break
    
    max_day = max(by_day.keys()) if by_day else 1
    if max_day in night_locations:
        last_events = by_day[max_day]
        if any("Fin de mission" in desc for _, _, desc in last_events):
            for _, _, desc in last_events:
                if "ArrivÃ©e" in desc and "Fin de mission" in desc:
                    city = desc.split("ArrivÃ©e")[1].split("â€“")[0].strip()
                    night_locations[max_day] = city

    first_date = start_date
    last_date = start_date + timedelta(days=stats['total_days']-1)
    
    months = ['janvier', 'fÃ©vrier', 'mars', 'avril', 'mai', 'juin', 
              'juillet', 'aoÃ»t', 'septembre', 'octobre', 'novembre', 'dÃ©cembre']
    date_range = f"{first_date.strftime('%d')} â†’ {last_date.strftime('%d')} {months[last_date.month-1]} {last_date.year}"
    
    num_nights = stats['total_days'] - 1 if stats['total_days'] > 1 else 0

    # KPIs et mÃ©ta
    actual_sites_count = len([s for s in sites_ordered if s.get('Type') != 'Base'])
    distance_km = stats.get('total_km', 0)
    total_visit_hours = stats.get('total_visit_hours', 0)
    route_summary = " â†’ ".join([s.get('Ville', '').upper() for s in sites_ordered if s.get('Ville')])
    gen_date_str = datetime.now().strftime("%d/%m/%Y")
    
    html = f"""<!DOCTYPE html>
<html lang="fr">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Planning {mission_title} ({date_range})</title>
    <style>
        body {{ font-family: Tahoma, Calibri, 'Segoe UI', sans-serif; margin: 20px; background-color: #f5f5f5; }}
        .container {{ max-width: 1200px; margin: 0 auto; background: white; border-radius: 10px; padding: 20px; box-shadow: 0 4px 6px rgba(0,0,0,0.1); }}
        h1 {{ text-align: center; color: #2c3e50; margin-bottom: 4px; font-size: 24px; }}
        p.subtitle {{ text-align: center; color: #7f8c8d; margin: 0 0 16px; font-size: 13px; }}
        .kpi-grid {{ display: grid; grid-template-columns: repeat(5, 1fr); gap: 12px; margin: 16px 0 18px; }}
        .kpi {{ background:#f8f9fb; border:1px solid #e6e8eb; border-radius:8px; padding:10px 12px; text-align:center; }}
        .kpi-label {{ color:#6c7a89; font-size:12px; }}
        .kpi-value {{ color:#2c3e50; font-weight:bold; font-size:18px; }}
        table {{ width: 100%; border-collapse: collapse; margin-bottom: 10px; font-size: 14px; }}
        th {{ background-color: #34495e; color: white; padding: 12px 8px; text-align: left; font-weight: bold; }}
        td {{ padding: 10px 8px; border-bottom: 1px solid #ddd; vertical-align: top; }}
        tr:nth-child(even) {{ background-color: #f8f9fa; }}
        tr:hover {{ background-color: #e8f4f8; }}
        .jour {{ font-weight: bold; color: #2980b9; background-color: #ecf0f1 !important; }}
        .horaire {{ font-weight: bold; color: #27ae60; white-space: nowrap; }}
        .activite {{ color: #2c3e50; }}
        .mission {{ background-color: #fff3cd; font-weight: bold; }}
        .route {{ color: #7f8c8d; font-style: italic; }}
        .nuit {{ background-color: #d1ecf1; font-weight: bold; color: #0c5460; text-align: center; }}
        .distance {{ color: #e74c3c; font-weight: bold; white-space: nowrap; }}
        .note {{ font-size: 12px; color: #7f8c8d; margin-top: 8px; }}
        .map-section {{ margin-top: 16px; }}
        .section {{ margin-top: 20px; }}
        .section h2 {{ color:#2c3e50; font-size:18px; margin-bottom:8px; }}
        .section ul {{ margin:8px 0 0 18px; }}
        .signatures {{ display:flex; gap:24px; margin-top:16px; }}
        .signature-box {{ flex:1; border:1px dashed #bdc3c7; border-radius:8px; padding:12px; color:#34495e; }}
        .footer {{ margin-top:14px; font-size:12px; color:#7f8c8d; }}
    </style>
</head>
<body>
<div class="container">
    <h1>ðŸ“‹ {mission_title} â€“ {date_range}</h1>
    <p class="subtitle">{stats['total_days']} jour{'s' if stats['total_days']>1 else ''} / {num_nights} nuitÃ©e{'s' if num_nights>1 else ''} â€¢ Pauses : dÃ©jeuner (13h00â€“14h30 â‰¤ 1h) & priÃ¨re (14h00â€“15h00 â‰¤ 20 min)</p>

    <div class="kpi-grid">
        <div class="kpi"><div class="kpi-label">DurÃ©e</div><div class="kpi-value">{stats['total_days']} j</div></div>
        <div class="kpi"><div class="kpi-label">Distance</div><div class="kpi-value">{distance_km:.1f} km</div></div>
        <div class="kpi"><div class="kpi-label">Sites</div><div class="kpi-value">{actual_sites_count}</div></div>
        <div class="kpi"><div class="kpi-label">Visites</div><div class="kpi-value">{total_visit_hours:.1f} h</div></div>
        <div class="kpi"><div class="kpi-label">NuitÃ©es</div><div class="kpi-value">{num_nights}</div></div>
    </div>

    <table>
        <thead>
            <tr>
                <th style="width: 15%;">JOUR</th>
                <th style="width: 15%;">HORAIRES</th>
                <th style="width: 40%;">ACTIVITÃ‰S</th>
                <th style="width: 15%;">TRANSPORT</th>
                <th style="width: 15%;">NUIT</th>
            </tr>
        </thead>
        <tbody>"""

    for day in sorted(by_day.keys()):
        day_events = by_day[day]
        
        display_events = []
        for sdt, edt, desc in day_events:
            if "NuitÃ©e" not in desc and "Fin de journÃ©e" not in desc:
                display_events.append((sdt, edt, desc))
        
        if not display_events:
            continue
            
        day_count = len(display_events)
        night_location = night_locations.get(day, "")
        
        html += f"""
            <!-- JOUR {day} -->"""
        
        for i, (sdt, edt, desc) in enumerate(display_events):
            if "â†’" in desc and "ðŸš—" in desc:
                activity_class = "route"
                activity_text = desc.replace("ðŸš— ", "ðŸš— ")
                transport_info = extract_distance_from_desc(desc, speed_kmh)
            elif any(word in desc.upper() for word in ["VISITE", "AGENCE", "SITE", "CLIENT"]):
                activity_class = "mission"
                activity_text = desc.replace("ðŸ¢", "").replace("ðŸ‘¥", "").replace("ðŸ“", "").replace("ðŸ ", "").strip()
                transport_info = "-"
            elif "dÃ©jeuner" in desc.lower() and "priÃ¨re" in desc.lower():
                activity_class = "activite"
                activity_text = desc
                transport_info = "-"
            elif "dÃ©jeuner" in desc.lower():
                activity_class = "activite"
                activity_text = desc
                transport_info = "-"
            elif "priÃ¨re" in desc.lower():
                activity_class = "activite"
                activity_text = desc
                transport_info = "-"
            elif "installation" in desc.lower() or "arrivÃ©e" in desc.lower():
                activity_class = "activite"
                activity_text = desc
                transport_info = "-"
            elif "fin" in desc.lower() and "mission" in desc.lower():
                activity_class = "activite"
                activity_text = desc
                transport_info = "-"
            else:
                activity_class = "activite"
                activity_text = desc
                transport_info = "-"
            
            if i == 0:
                html += f"""
            <tr class="jour">
                <td rowspan="{day_count}"><strong>JOUR {day}</strong></td>
                <td class="horaire">{fmt_time(sdt)}â€“{fmt_time(edt)}</td>
                <td class="{activity_class}">{activity_text}</td>
                <td class="distance">{transport_info}</td>
                <td rowspan="{day_count}" class="nuit">{night_location}</td>
            </tr>"""
            else:
                html += f"""
            <tr>
                <td class="horaire">{fmt_time(sdt)}â€“{fmt_time(edt)}</td>
                <td class="{activity_class}">{activity_text}</td>
                <td class="distance">{transport_info}</td>
            </tr>"""

    html += f"""
        </tbody>
    </table>

    <p class="note">â„¹ï¸ Distances/temps indicatifs. Les pauses dÃ©jeuner et priÃ¨re sont flexibles et intÃ©grÃ©es sans bloquer les activitÃ©s.</p>
"""

    if include_details:
        html += f"""
    <div class="section">
        <h2>ðŸ“‹ RÃ©sumÃ© exÃ©cutif</h2>
        <ul>
            <li>ItinÃ©raire: {route_summary}</li>
            <li>Distance: {distance_km:.1f} km; Visites: {total_visit_hours:.1f} h; Sites: {actual_sites_count}</li>
            <li>PÃ©riode: {date_range}</li>
        </ul>
    </div>
    <div class="section signatures">
        <div class="signature-box">PrÃ©parÃ© par: __________________<br/>Fonction: __________________<br/>Date: {gen_date_str}</div>
        <div class="signature-box">ValidÃ© par: __________________<br/>Fonction: __________________<br/>Date: {gen_date_str}</div>
    </div>
    <div class="footer">App dev by Moctar TALL â€¢ Document gÃ©nÃ©rÃ© le {gen_date_str}</div>
"""

    # IntÃ©grer la carte en-dessous du tableau si coords_ordered est fourni (optionnel)
    map_embed_html = ""
    try:
        if include_map and coords_ordered and len(coords_ordered) > 0:
            center_lat = sum(c[1] for c in coords_ordered) / len(coords_ordered)
            center_lon = sum(c[0] for c in coords_ordered) / len(coords_ordered)

            m = folium.Map(location=[center_lat, center_lon], zoom_start=7)

            # Route via OSRM, fallback sur ligne droite
            try:
                coord_str = ";".join([f"{c[0]},{c[1]}" for c in coords_ordered])
                url = f"{osrm_base_url.rstrip('/')}/route/v1/driving/{coord_str}?overview=full&geometries=geojson"
                resp = requests.get(url, timeout=10)
                route_pts = None
                if resp.status_code == 200:
                    data = resp.json()
                    if data.get('routes'):
                        geom = data['routes'][0].get('geometry')
                        if isinstance(geom, dict) and geom.get('coordinates'):
                            route_pts = [[lat, lon] for lon, lat in geom['coordinates']]
                if not route_pts:
                    route_pts = [[c[1], c[0]] for c in coords_ordered]
            except Exception:
                route_pts = [[c[1], c[0]] for c in coords_ordered]
            folium.PolyLine(locations=route_pts, color="blue", weight=3, opacity=0.7).add_to(m)

            # Affichage spÃ©cial si dÃ©part et arrivÃ©e identiques
            n_steps = len(sites_ordered)
            start_end_same = False
            if n_steps >= 2:
                lat0, lon0 = coords_ordered[0][1], coords_ordered[0][0]
                latN, lonN = coords_ordered[-1][1], coords_ordered[-1][0]
                start_end_same = abs(lat0 - latN) < 1e-4 and abs(lon0 - lonN) < 1e-4

            for i, site in enumerate(sites_ordered):
                if i == 0 and start_end_same:
                    bg_color_left = '#2ecc71'  # Vert pour dÃ©part
                    bg_color_right = '#e74c3c'  # Rouge pour arrivÃ©e
                    html_num = f"""
<div style=\"display:flex; align-items:center; gap:4px;\">
  <div style=\"background-color:{bg_color_left}; color:white; border-radius:50%; width:28px; height:28px; text-align:center; font-size:14px; font-weight:bold; line-height:28px; border:2px solid white; box-shadow:0 0 3px rgba(0,0,0,0.5);\">1</div>
  <div style=\"background-color:{bg_color_right}; color:white; border-radius:50%; width:28px; height:28px; text-align:center; font-size:14px; font-weight:bold; line-height:28px; border:2px solid white; box-shadow:0 0 3px rgba(0,0,0,0.5);\">{n_steps}</div>
</div>
"""
                    folium.Marker(
                        location=[coords_ordered[i][1], coords_ordered[i][0]],
                        popup=f"Ã‰tapes 1 et {n_steps}: {site.get('Ville','')}<br>{site.get('Type', '-')}",
                        tooltip=f"Ã‰tapes 1 et {n_steps}: {site.get('Ville','')}",
                        icon=folium.DivIcon(
                            icon_size=(36, 28),
                            icon_anchor=(18, 14),
                            html=html_num
                        )
                    ).add_to(m)
                    continue
                if start_end_same and i == n_steps - 1:
                    continue

                bg_color = '#2ecc71' if i == 0 else '#e74c3c' if i == len(sites_ordered)-1 else '#3498db'
                folium.Marker(
                    location=[coords_ordered[i][1], coords_ordered[i][0]],
                    popup=f"Ã‰tape {i+1}: {site.get('Ville','')}<br>{site.get('Type', '-')}",
                    tooltip=f"Ã‰tape {i+1}: {site.get('Ville','')}",
                    icon=folium.DivIcon(
                        icon_size=(28, 28),
                        icon_anchor=(14, 14),
                        html=f"""
<div style=\"background-color:{bg_color}; color:white; border-radius:50%; width:28px; height:28px; text-align:center; font-size:14px; font-weight:bold; line-height:28px; border:2px solid white; box-shadow:0 0 3px rgba(0,0,0,0.5);\">{i+1}</div>
"""
                    )
                ).add_to(m)

            map_html = m.get_root().render()
            import base64
            map_b64 = base64.b64encode(map_html.encode('utf-8')).decode('ascii')
            map_embed_html = f"""
    <div class=\"map-section\">
        <h2>ðŸ—ºï¸ Carte de l'itinÃ©raire</h2>
        <iframe src=\"data:text/html;base64,{map_b64}\" style=\"width:100%; height:600px; border:none;\"></iframe>
    </div>
"""
    except Exception:
        map_embed_html = ""

    if map_embed_html:
        html += map_embed_html

    html += """
</div>
</body>
</html>"""

    return html

def create_mission_excel(itinerary, start_date, stats, sites_ordered, segments_summary=None, mission_title="Mission Terrain"):
    """
    GÃ©nÃ¨re un fichier Excel professionnel Ã  partir des donnÃ©es de planning
    """
    import io
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils.dataframe import dataframe_to_rows
    
    # CrÃ©er un workbook
    wb = Workbook()
    ws = wb.active
    ws.title = "Planning Mission"
    
    # Styles
    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill(start_color="2E86AB", end_color="2E86AB", fill_type="solid")
    subheader_font = Font(bold=True, color="2E86AB")
    border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    center_alignment = Alignment(horizontal='center', vertical='center')
    
    # En-tÃªte principal
    ws.merge_cells('A1:F1')
    ws['A1'] = mission_title
    ws['A1'].font = Font(bold=True, size=16, color="2E86AB")
    ws['A1'].alignment = center_alignment
    
    # Informations gÃ©nÃ©rales
    current_row = 3
    ws[f'A{current_row}'] = f"ðŸ“… PÃ©riode: {start_date.strftime('%d/%m/%Y')} â†’ {(start_date + timedelta(days=len(itinerary)-1)).strftime('%d/%m/%Y')}"
    ws[f'A{current_row}'].font = subheader_font
    current_row += 1
    
    ws[f'A{current_row}'] = f"ðŸƒ {stats['total_days']} jour{'s' if stats['total_days'] > 1 else ''} / 0 nuitÃ©e â€¢ Pauses flexibles : dÃ©jeuner (13h00-14h30 â‰¤ 1h) & priÃ¨re (14h00-15h00 â‰¤ 20 min)"
    current_row += 2
    
    # En-tÃªtes du tableau
    headers = ['JOUR', 'HORAIRES', 'ACTIVITÃ‰S', 'TRANSPORT', 'NUIT']
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=current_row, column=col, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = center_alignment
        cell.border = border
    
    current_row += 1
    
    # DonnÃ©es du planning
    # L'itinÃ©raire est une liste de tuples: (day, start_time, end_time, description)
    current_day = None
    day_start_row = current_row
    
    for event in itinerary:
        day, start_time, end_time, description = event
        
        # Nouvelle journÃ©e
        if day != current_day:
            current_day = day
            day_start_row = current_row
            
            # Colonne JOUR
            ws.cell(row=current_row, column=1, value=f"JOUR {day}")
            ws.cell(row=current_row, column=1).font = Font(bold=True)
            ws.cell(row=current_row, column=1).alignment = center_alignment
            ws.cell(row=current_row, column=1).border = border
        else:
            ws.cell(row=current_row, column=1, value="")
            ws.cell(row=current_row, column=1).border = border
        
        # Colonne HORAIRES
        if isinstance(start_time, str):
            time_str = start_time
        else:
            time_str = f"{start_time.strftime('%Hh%M')}-{end_time.strftime('%Hh%M')}"
        
        ws.cell(row=current_row, column=2, value=time_str)
        ws.cell(row=current_row, column=2).alignment = center_alignment
        ws.cell(row=current_row, column=2).border = border
        
        # Colonne ACTIVITÃ‰S
        ws.cell(row=current_row, column=3, value=description)
        
        # Coloration selon le type d'activitÃ©
        if "ðŸš—" in description or "â†’" in description:
            # Transport
            pass  # Pas de coloration spÃ©ciale
        elif "ðŸ½ï¸" in description or "DÃ©jeuner" in description:
            # DÃ©jeuner
            ws.cell(row=current_row, column=3).fill = PatternFill(start_color="E8F5E8", end_color="E8F5E8", fill_type="solid")
        elif "ðŸ•Œ" in description or "PriÃ¨re" in description:
            # PriÃ¨re
            ws.cell(row=current_row, column=3).fill = PatternFill(start_color="E8F5E8", end_color="E8F5E8", fill_type="solid")
        else:
            # ActivitÃ© normale
            ws.cell(row=current_row, column=3).fill = PatternFill(start_color="FFF2CC", end_color="FFF2CC", fill_type="solid")
        
        ws.cell(row=current_row, column=3).border = border
        
        # Colonne TRANSPORT
        import re
        # Extraire distance et durÃ©e de la description
        if "ðŸš—" in description and "(" in description:
            # Format: "ðŸš— Dakar â†’ Saint-Louis (240.9 km, 0min)"
            match = re.search(r"\(([\d\.]+)\s*km,\s*([^)]+)\)", description)
            if match:
                km = match.group(1)
                duration = match.group(2).strip()
                transport_text = f"~{km} km / ~{duration}"
                ws.cell(row=current_row, column=4, value=transport_text)
                ws.cell(row=current_row, column=4).font = Font(color="D32F2F")
            else:
                ws.cell(row=current_row, column=4, value="-")
        else:
            ws.cell(row=current_row, column=4, value="-")
        
        ws.cell(row=current_row, column=4).alignment = center_alignment
        ws.cell(row=current_row, column=4).border = border
        
        # Colonne NUIT
        ws.cell(row=current_row, column=5, value="")
        ws.cell(row=current_row, column=5).fill = PatternFill(start_color="E3F2FD", end_color="E3F2FD", fill_type="solid")
        ws.cell(row=current_row, column=5).border = border
        
        current_row += 1
    
    # Note en bas
    current_row += 1
    ws[f'A{current_row}'] = "â„¹ï¸ Distances/temps indicatifs. DÃ©jeuner (13h00-14h30, â‰¤1h) et priÃ¨re (14h00-15h00, â‰¤20 min) sont flexibles et intÃ©grÃ©s sans bloquer les activitÃ©s."
    ws[f'A{current_row}'].font = Font(size=9, italic=True)
    ws.merge_cells(f'A{current_row}:E{current_row}')
    
    # Ajuster la largeur des colonnes
    ws.column_dimensions['A'].width = 12
    ws.column_dimensions['B'].width = 15
    ws.column_dimensions['C'].width = 50
    ws.column_dimensions['D'].width = 20
    ws.column_dimensions['E'].width = 12
    
    # Sauvegarder dans un buffer
    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    
    return buffer.getvalue()

# Test de connexion
if st.sidebar.button("ðŸ” Tester connexion Maps"):
    # Animation d'attente amÃ©liorÃ©e
    progress_bar = st.sidebar.progress(0)
    status_text = st.sidebar.empty()
    
    # Ã‰tape 1: Initialisation
    progress_bar.progress(25)
    status_text.text("ðŸ”„ Initialisation du test...")
    
    # Ã‰tape 2: Test de connexion
    progress_bar.progress(75)
    status_text.text("ðŸŒ Test de connexion Maps...")
    
    success, message = test_graphhopper_connection(graphhopper_api_key)
    
    # Ã‰tape 3: Finalisation
    progress_bar.progress(100)
    status_text.text("âœ… Test terminÃ©")
    
    # Nettoyage de l'animation
    progress_bar.empty()
    status_text.empty()
    
    if success:
        st.sidebar.success(f"âœ… {message}")
    else:
        st.sidebar.error(f"âŒ {message}")

# Mention dÃ©veloppeur
st.sidebar.markdown("---")
st.sidebar.caption("ðŸ’» Developed by @Moctar TAll (+221 77 639 96 12)")
st.sidebar.caption("All rights reserved")

# --------------------------
# FORMULAIRE
# --------------------------
st.header("ðŸ“ ParamÃ¨tres de la mission")

# Champ pour le titre de mission personnalisÃ©
st.subheader("ðŸ“ Titre de la mission")
mission_title = st.text_input(
    "Titre personnalisÃ© de votre mission",
    value=st.session_state.get("mission_title", "Mission Terrain"),
    help="Ce titre apparaÃ®tra dans la prÃ©sentation professionnelle et tous les documents gÃ©nÃ©rÃ©s",
    placeholder="Ex: Mission d'inspection technique, Visite commerciale, Audit de site..."
)

st.divider()

tab1, tab2, tab3 = st.tabs(["Sites Ã  visiter", "Dates et Horaires de la mission ", "ParamÃ¨trage des pauses"])

with tab1:
    st.markdown("**Configurez votre mission**")
    
    st.subheader("ðŸ  Point de dÃ©part et d'arrivÃ©e")
    col1, col2 = st.columns(2)
    
    with col1:
        use_base_location = st.checkbox("Utiliser un point de dÃ©part/arrivÃ©e fixe", value=st.session_state.get("use_base_location", True))
    
    with col2:
        if use_base_location:
            base_location = st.text_input("Ville de dÃ©part/arrivÃ©e", value=st.session_state.get("base_location", "Dakar"))
        else:
            base_location = ""
    
    st.divider()
    
    # En-tÃªte optimisÃ© avec informations contextuelles
    col_header1, col_header2 = st.columns([3, 1])
    with col_header1:
        st.subheader("ðŸ“ Sites Ã  visiter")
    with col_header2:
        # Affichage compact du statut et compteur sur la mÃªme ligne
        if 'data_saved' in st.session_state and st.session_state.data_saved:
            col_status, col_count = st.columns([1, 1])
            with col_status:
                st.success("âœ… SauvegardÃ©")
            with col_count:
                st.metric("Sites", len(st.session_state.sites_df) if 'sites_df' in st.session_state else 0)
    
    # Message d'aide contextuel
    if 'sites_df' not in st.session_state or len(st.session_state.sites_df) == 0:
        st.info("ðŸ’¡ **Commencez par ajouter vos sites Ã  visiter** - Utilisez le tableau ci-dessous pour saisir les villes, types d'activitÃ©s et durÃ©es prÃ©vues.")
    
    if 'sites_df' not in st.session_state:
        if use_base_location:
            st.session_state.sites_df = pd.DataFrame([
                {"Ville": "ThiÃ¨s", "Type": "Client", "ActivitÃ©": "RÃ©union commerciale", "DurÃ©e (h)": 2.0, "Peut continuer": False, "PossibilitÃ© de nuitÃ©e": True},
                {"Ville": "Saint-Louis", "Type": "Sites technique", "ActivitÃ©": "Inspection", "DurÃ©e (h)": 3.0, "Peut continuer": False, "PossibilitÃ© de nuitÃ©e": True},
            ])
        else:
            st.session_state.sites_df = pd.DataFrame([
                {"Ville": "Dakar", "Type": "Agence", "ActivitÃ©": "Brief", "DurÃ©e (h)": 0.5, "Peut continuer": False, "PossibilitÃ© de nuitÃ©e": True},
                {"Ville": "ThiÃ¨s", "Type": "Sites technique", "ActivitÃ©": "Visite", "DurÃ©e (h)": 2.0, "Peut continuer": False, "PossibilitÃ© de nuitÃ©e": True},
            ])
    
    # Gestion des types de sites personnalisÃ©s
    if 'custom_site_types' not in st.session_state:
        st.session_state.custom_site_types = []
    
    # Types de base + types personnalisÃ©s
    base_types = ["Agence", "Client", "Sites technique", "Site BTS", "Partenaire", "Autre"]
    all_types = base_types + st.session_state.custom_site_types
    
    # Tableau optimisÃ© avec liste dÃ©roulante et saisie libre
    st.markdown("**ðŸ“‹ Tableau des sites Ã  visiter :**")
    
    # Ajouter une option "Autre (saisir)" pour permettre la saisie libre
    dropdown_options = all_types + ["âœï¸ Autre (saisir)"]
    
    # PrÃ©parer un DataFrame Ã©ditable en y ajoutant une colonne de suppression
    editable_df = st.session_state.sites_df.copy()
    if 'Supprimer' not in editable_df.columns:
        try:
            editable_df['Supprimer'] = False
        except Exception:
            # En cas d'une structure inattendue, garantir l'existence de la colonne
            editable_df = pd.DataFrame(editable_df)
            editable_df['Supprimer'] = False

    sites_df = st.data_editor(
        editable_df, 
        num_rows="dynamic", 
        use_container_width=True,
        key="sites_data_editor",
        height=300,  # Hauteur fixe pour une meilleure lisibilitÃ©
        column_config={
            "Supprimer": st.column_config.CheckboxColumn(
                "ðŸ—‘ï¸",
                default=False,
                help="Cocher pour supprimer cette ligne",
                width=35
            ),
            "Ville": st.column_config.TextColumn(
                "ðŸ™ï¸ Ville", 
                required=True,
                help="Nom de la ville ou localitÃ© Ã  visiter",
                width="medium"
            ),
            "Type": st.column_config.SelectboxColumn(
                "ðŸ¢ Type",
                options=dropdown_options,
                default="Sites technique",
                help="SÃ©lectionnez un type ou choisissez 'Autre (saisir)' pour crÃ©er un nouveau type",
                width="medium"
            ),
            "ActivitÃ©": st.column_config.TextColumn(
                "âš¡ ActivitÃ©", 
                default="Visite",
                help="Nature de l'activitÃ© prÃ©vue",
                width="medium"
            ),
            "DurÃ©e (h)": st.column_config.NumberColumn(
                "â±ï¸ DurÃ©e (h)",
                min_value=0.25,
                max_value=24,
                step=0.25,
                format="%.2f",
                default=1.0,
                help="DurÃ©e estimÃ©e en heures",
                width="small"
            ),
            "Peut continuer": st.column_config.CheckboxColumn(
                "ðŸ”„ Peut continuer",
                default=False,
                help="Cochez si cette activitÃ© peut Ãªtre reportÃ©e au jour suivant si elle dÃ©passe les heures d'activitÃ©",
                width="small"
            ),
            "PossibilitÃ© de nuitÃ©e": st.column_config.CheckboxColumn(
                "ðŸ¨ NuitÃ©e possible",
                default=True,
                help="DÃ©cochez si cette zone ne dispose pas d'hÃ©bergement correct et qu'il faut Ã©viter d'y passer la nuit",
                width="small"
            )
        },
        column_order=["Supprimer", "Ville", "Type", "ActivitÃ©", "DurÃ©e (h)", "Peut continuer", "PossibilitÃ© de nuitÃ©e"]
    )
    
    # Interface pour saisir un nouveau type si "Autre (saisir)" est sÃ©lectionnÃ©
    if sites_df is not None and not sites_df.empty:
        # VÃ©rifier s'il y a des lignes avec "âœï¸ Autre (saisir)"
        custom_rows = sites_df[sites_df['Type'] == "âœï¸ Autre (saisir)"]
        if not custom_rows.empty:
            st.info("ðŸ’¡ **Nouveau type dÃ©tectÃ©** - Veuillez spÃ©cifier le type personnalisÃ© ci-dessous :")
            
            for idx in custom_rows.index:
                col1, col2, col3 = st.columns([2, 3, 1])
                with col1:
                    st.write(f"**Ligne {idx + 1}** - {sites_df.loc[idx, 'Ville']}")
                with col2:
                    new_custom_type = st.text_input(
                        f"Type personnalisÃ© pour la ligne {idx + 1}",
                        placeholder="Ex: Site industriel, Centre de donnÃ©es...",
                        key=f"custom_type_{idx}",
                        label_visibility="collapsed"
                    )
                with col3:
                    if st.button("âœ…", key=f"apply_{idx}", help="Appliquer ce type"):
                        if new_custom_type and new_custom_type.strip():
                            # Ajouter le nouveau type Ã  la liste des types personnalisÃ©s
                            if new_custom_type.strip() not in st.session_state.custom_site_types:
                                st.session_state.custom_site_types.append(new_custom_type.strip())
                            
                            # Mettre Ã  jour la ligne dans le DataFrame
                            sites_df.loc[idx, 'Type'] = new_custom_type.strip()
                            st.session_state.sites_df = sites_df
                            # Pas de rerun automatique pour Ã©viter de ralentir la saisie
    
    # Boutons d'action
    # VÃ©rifier s'il y a des lignes cochÃ©es pour suppression
    has_checked_rows = 'Supprimer' in sites_df.columns and sites_df['Supprimer'].any()
    
    if has_checked_rows:
        col1, col2, col3 = st.columns([2, 1, 2])
        with col1:
            if st.button("ðŸ—‘ï¸ Supprimer lignes cochÃ©es", use_container_width=True):
                # SÃ©curitÃ©: gÃ©rer le cas oÃ¹ la colonne 'Supprimer' serait absente ou non boolÃ©enne
                if 'Supprimer' in sites_df.columns:
                    suppr_series = sites_df['Supprimer'].fillna(False)
                    try:
                        suppr_mask = ~suppr_series.astype(bool)
                    except Exception:
                        # Si conversion Ã©choue, ne supprimer aucune ligne
                        suppr_mask = [True] * len(sites_df)
                    remaining_df = sites_df[suppr_mask].copy()
                    if 'Supprimer' in remaining_df.columns:
                        remaining_df = remaining_df.drop(columns=['Supprimer'])
                else:
                    remaining_df = sites_df.copy()
                st.session_state.sites_df = remaining_df.reset_index(drop=True)
                st.success("Lignes sÃ©lectionnÃ©es supprimÃ©es")
                st.rerun()
    else:
        col1, col2, col3 = st.columns([2, 1, 2])

    with col2:
        if st.button("ðŸ’¾ Enregistrer", use_container_width=True, type="primary"):
            # Nettoyer la colonne de suppression avant sauvegarde
            df_to_save = sites_df.drop(columns=['Supprimer']) if 'Supprimer' in sites_df.columns else sites_df
            st.session_state.sites_df = df_to_save
            st.session_state.data_saved = True  # Marquer comme sauvegardÃ©
            st.rerun()  # RafraÃ®chir pour afficher le statut en haut
    
    # Pas d'enregistrement automatique - seulement lors du clic sur Enregistrer ou Planifier
    # st.session_state.sites_df = sites_df
    
    # Option d'ordre des sites
    order_mode = "ðŸ¤– Automatique (optimisÃ©)"  # Valeur par dÃ©faut pour 0 ou 1 site
    if len(sites_df) > 1:  # Afficher seulement s'il y a plus d'un site
        st.subheader("ðŸ”„ Ordre des visites")
        order_mode = st.radio(
            "Mode d'ordonnancement",
            ["ðŸ¤– Automatique (optimisÃ©)", "âœ‹ Manuel (personnalisÃ©)"],
            horizontal=True,
            help="Automatique: optimise l'ordre pour minimiser les distances. Manuel: vous choisissez l'ordre."
        )
        
        if order_mode == "âœ‹ Manuel (personnalisÃ©)":
            with st.container():
                st.info("ðŸ’¡ **Astuce :** Utilisez les flÃ¨ches pour rÃ©organiser vos sites dans l'ordre de visite souhaitÃ©")
                
                # CrÃ©er une liste ordonnÃ©e des sites pour rÃ©organisation
                if 'manual_order' not in st.session_state or len(st.session_state.manual_order) != len(sites_df):
                    st.session_state.manual_order = list(range(len(sites_df)))
                
                # Interface de rÃ©organisation manuelle amÃ©liorÃ©e
                st.markdown("**ðŸ“‹ Ordre de visite des sites :**")
                
                # Conteneur avec style pour la liste
                with st.container():
                    for i, idx in enumerate(st.session_state.manual_order):
                        if idx < len(sites_df):
                            site = sites_df.iloc[idx]
                            
                            # CrÃ©er une ligne avec un style visuel amÃ©liorÃ©
                            col1, col2, col3, col4, col5 = st.columns([0.8, 2.5, 2, 1, 1])
                            
                            with col1:
                                st.markdown(f"**`{i+1}`**")
                            with col2:
                                st.markdown(f"ðŸ“ **{site['Ville']}**")
                            with col3:
                                st.markdown(f"ðŸ¢ {site['Type']}")
                            with col4:
                                st.markdown(f"â±ï¸ {site['DurÃ©e (h)']}h")
                            with col5:
                                # Boutons de rÃ©organisation dans une ligne
                                subcol1, subcol2 = st.columns(2)
                                with subcol1:
                                    if i > 0:
                                        if st.button("â¬†ï¸", key=f"manual_up_{i}", help="Monter", use_container_width=True):
                                            st.session_state.manual_order[i], st.session_state.manual_order[i-1] = \
                                                st.session_state.manual_order[i-1], st.session_state.manual_order[i]
                                            st.rerun()
                                with subcol2:
                                    if i < len(st.session_state.manual_order) - 1:
                                        if st.button("â¬‡ï¸", key=f"manual_down_{i}", help="Descendre", use_container_width=True):
                                            st.session_state.manual_order[i], st.session_state.manual_order[i+1] = \
                                                st.session_state.manual_order[i+1], st.session_state.manual_order[i]
                                            st.rerun()
                            
                            # SÃ©parateur visuel entre les Ã©lÃ©ments
                            if i < len(st.session_state.manual_order) - 1:
                                st.markdown("---")
                
                # Boutons d'action
                col1, col2, col3 = st.columns([1, 1, 2])
                with col1:
                    if st.button("ðŸ”„ RÃ©initialiser", help="Remettre l'ordre original", use_container_width=True):
                        st.session_state.manual_order = list(range(len(sites_df)))
                        st.rerun()
                with col2:
                    if st.button("ðŸ”€ MÃ©langer", help="Ordre alÃ©atoire", use_container_width=True):
                        import random
                        random.shuffle(st.session_state.manual_order)
                        st.rerun()
        else:
            st.success("ðŸ¤– **Mode automatique activÃ©** - L'ordre des sites sera optimisÃ© automatiquement pour minimiser les temps de trajet")
    else:
        st.info("â„¹ï¸ Ajoutez au moins 1 site pour continuer. L'ordre n'est requis que s'il y a plusieurs sites.")

with tab2:
    col1, col2 = st.columns([1, 2])  # RÃ©duire la largeur de la colonne Dates
    with col1:
        st.subheader("ðŸ“… Dates")
        start_date = st.date_input("Dat                                                                                                                                                                                                                                             e de dÃ©part de la mission", value=st.session_state.get("start_date", datetime.today().date()))
        max_days = st.number_input("Nombre de jours max (Optionel)", min_value=0, value=st.session_state.get("max_days", 0), step=1, help="Laisser zÃ©ro pour le calcul automatique. Agit comme une limite supÃ©rieure.")
        desired_days = st.number_input("Nombre de jours souhaitÃ©s (Optionnel)", min_value=0, value=st.session_state.get("desired_days", 0), step=1, help="Laissez Ã  zÃ©ro pour ignorer. Le planning sera ajustÃ© pour correspondre Ã  ce nombre si possible.")
        
        st.divider()
        
        # Ajouter des informations utiles dans la section Dates
        st.markdown("**ðŸ“Š Informations**")
        if start_date:
            # Jour de la semaine avec date complÃ¨te
            weekdays = ["Lundi", "Mardi", "Mercredi", "Jeudi", "Vendredi", "Samedi", "Dimanche"]
            months = ["janvier", "fÃ©vrier", "mars", "avril", "mai", "juin", 
                     "juillet", "aoÃ»t", "septembre", "octobre", "novembre", "dÃ©cembre"]
            start_weekday = weekdays[start_date.weekday()]
            start_month = months[start_date.month - 1]
            formatted_date = f"{start_weekday.lower()} {start_date.day} {start_month} {start_date.year}"
            st.info(f"ðŸ—“ï¸ Jour de dÃ©but : {formatted_date}")
    
    with col2:
        st.subheader("â° Horaires")
        
        # Horaires d'activitÃ©
        st.markdown("**Horaires d'activitÃ©** (visites, rÃ©unions)")
        col_act1, col_act2 = st.columns(2)
        with col_act1:
            start_activity_time = st.time_input("DÃ©but activitÃ©s", value=st.session_state.get("start_activity_time", time(8, 0)))
        with col_act2:
            end_activity_time = st.time_input("Fin activitÃ©s", value=st.session_state.get("end_activity_time", time(16, 30)))
        
        # Horaires de voyage
        st.markdown("**Horaires de voyage** (trajets)")
        col_travel1, col_travel2 = st.columns(2)
        with col_travel1:
            start_travel_time = st.time_input("DÃ©but voyages", value=st.session_state.get("start_travel_time", time(7, 30)))
        with col_travel2:
            end_travel_time = st.time_input("Fin voyages", value=st.session_state.get("end_travel_time", time(19, 0)))
        
        # Options week-end
        st.markdown("**Options week-end**")
        allow_weekend_travel = st.checkbox(
            "Autoriser les voyages le week-end",
            value=st.session_state.get("allow_weekend_travel", True)
        )
        allow_weekend_activities = st.checkbox(
            "Autoriser les activitÃ©s le week-end",
            value=st.session_state.get("allow_weekend_activities", True)
        )
        st.session_state.allow_weekend_travel = allow_weekend_travel
        st.session_state.allow_weekend_activities = allow_weekend_activities
        
        st.divider()
        
        # Gestion des activitÃ©s longues
        st.markdown("**Gestion des activitÃ©s longues**")
        col_tol1, col_tol2 = st.columns(2)
        with col_tol1:
            tolerance_hours = st.number_input(
                "Seuil de tolÃ©rance (heures)", 
                min_value=0.0, 
                max_value=3.0, 
                value=st.session_state.get("tolerance_hours", 1.0), 
                step=0.25,
                help="ActivitÃ©s se terminant dans ce dÃ©lai aprÃ¨s la fin des heures d'activitÃ© peuvent continuer le mÃªme jour"
            )
        with col_tol2:
            default_can_continue = st.checkbox(
                "Une partie dâ€™une activitÃ© non achevÃ©e Ã  lâ€™heure de la descente pourra Ãªtre poursuivie le lendemain", 
                value=False,
                help="Non poursuite cochÃ©e par dÃ©faut"
            )
        
        # Maintenir la compatibilitÃ© avec l'ancien code
        start_day_time = start_activity_time
        end_day_time = end_activity_time

with tab3:
    st.subheader("ðŸ½ï¸ Pauses flexibles")
    st.info("ðŸ’¡ Les pauses s'insÃ¨rent automatiquement pendant les trajets ou visites qui chevauchent les fenÃªtres dÃ©finies")
    
    col1, col2 = st.columns(2)
    with col1:
        use_lunch = st.checkbox("Pause dÃ©jeuner", value=st.session_state.get("use_lunch", True))
        if use_lunch:
            st.markdown("**FenÃªtre de dÃ©jeuner**")
            lunch_start_time = st.time_input("DÃ©but fenÃªtre", value=st.session_state.get("lunch_start_time", time(12, 30)), key="lunch_start")
            lunch_end_time = st.time_input("Fin fenÃªtre", value=st.session_state.get("lunch_end_time", time(15, 0)), key="lunch_end")
            lunch_duration_min = st.number_input(
                "DurÃ©e pause (min)",
                min_value=5,
                max_value=180,
                step=5,
                value=st.session_state.get("lunch_duration_min", 60),
                key="lunch_duration"
            )
    
    with col2:
        use_prayer = st.checkbox("Pause priÃ¨re", value=st.session_state.get("use_prayer", False))
        if use_prayer:
            st.markdown("**FenÃªtre de priÃ¨re**")
            prayer_start_time = st.time_input("DÃ©but fenÃªtre", value=st.session_state.get("prayer_start_time", time(13, 0)), key="prayer_start")
            prayer_duration_min = st.number_input("DurÃ©e pause (min)", min_value=5, max_value=60, value=st.session_state.get("prayer_duration_min", 20) or 20, key="prayer_duration")

    st.divider()
    st.subheader("ðŸ“¦ Import/Export JSON")
    with st.expander("Sauvegarde et reprise (JSON)"):
        col_export, col_import = st.columns(2)
        with col_export:
            mission_config = {
                "mission_title": mission_title,
                "use_base_location": use_base_location,
                "base_location": base_location,
                "sites": (st.session_state.sites_df.to_dict(orient="records") if "sites_df" in st.session_state else []),
                "start_date": start_date.strftime("%Y-%m-%d") if isinstance(start_date, datetime) else str(start_date),
                "max_days": int(max_days),
                "start_activity_time": (start_activity_time.strftime("%H:%M") if start_activity_time else None),
                "end_activity_time": (end_activity_time.strftime("%H:%M") if end_activity_time else None),
                "start_travel_time": (start_travel_time.strftime("%H:%M") if start_travel_time else None),
                "end_travel_time": (end_travel_time.strftime("%H:%M") if end_travel_time else None),
                "tolerance_hours": float(tolerance_hours),
                "use_lunch": bool(use_lunch),
                "lunch_start_time": (lunch_start_time.strftime("%H:%M") if use_lunch and lunch_start_time else None),
                "lunch_end_time": (lunch_end_time.strftime("%H:%M") if use_lunch and lunch_end_time else None),
                "lunch_duration_min": (int(lunch_duration_min) if use_lunch else None),
                "use_prayer": bool(use_prayer),
                "prayer_start_time": (prayer_start_time.strftime("%H:%M") if use_prayer and prayer_start_time else None),
                "prayer_duration_min": (int(prayer_duration_min) if use_prayer and prayer_duration_min is not None else None),
                "distance_method": distance_method,
            }
            json_str = json.dumps(mission_config, ensure_ascii=False, indent=2)
            st.download_button(
                label="ðŸ’¾ Exporter JSON",
                data=json_str,
                file_name=f"mission_config_{datetime.now().strftime('%Y%m%d_%H%M')}.json",
                mime="application/json",
                use_container_width=True,
            )
        with col_import:
            uploaded_file = st.file_uploader("Importer configuration JSON", type=["json"], help="Chargez un fichier exportÃ© prÃ©cÃ©demment pour reprendre la mission")
            if uploaded_file is not None:
                try:
                    imported = json.loads(uploaded_file.getvalue().decode("utf-8"))
                    def parse_time(val, fallback):
                        try:
                            if isinstance(val, str) and ":" in val:
                                hh, mm = val.split(":")
                                return time(int(hh), int(mm))
                        except Exception:
                            pass
                        return fallback

                    st.session_state.mission_title = imported.get("mission_title", mission_title)
                    st.session_state.use_base_location = imported.get("use_base_location", use_base_location)
                    st.session_state.base_location = imported.get("base_location", base_location)

                    # Sites
                    sites_records = imported.get("sites", [])
                    if isinstance(sites_records, list):
                        st.session_state.sites_df = pd.DataFrame(sites_records)

                    # Dates et horaires
                    sd = imported.get("start_date")
                    try:
                        if sd:
                            st.session_state.start_date = datetime.strptime(sd, "%Y-%m-%d").date()
                    except Exception:
                        st.session_state.start_date = st.session_state.get("start_date", datetime.today().date())

                    st.session_state.max_days = imported.get("max_days", max_days)
                    st.session_state.start_activity_time = parse_time(imported.get("start_activity_time"), start_activity_time)
                    st.session_state.end_activity_time = parse_time(imported.get("end_activity_time"), end_activity_time)
                    st.session_state.start_travel_time = parse_time(imported.get("start_travel_time"), start_travel_time)
                    st.session_state.end_travel_time = parse_time(imported.get("end_travel_time"), end_travel_time)
                    st.session_state.tolerance_hours = imported.get("tolerance_hours", tolerance_hours)

                    # Pauses
                    st.session_state.use_lunch = imported.get("use_lunch", use_lunch)
                    st.session_state.lunch_start_time = parse_time(imported.get("lunch_start_time"), lunch_start_time)
                    st.session_state.lunch_end_time = parse_time(imported.get("lunch_end_time"), lunch_end_time)
                    st.session_state.lunch_duration_min = imported.get("lunch_duration_min", st.session_state.get("lunch_duration_min", 60))
                    st.session_state.use_prayer = imported.get("use_prayer", use_prayer)
                    st.session_state.prayer_start_time = parse_time(imported.get("prayer_start_time"), prayer_start_time)
                    val_prayer_dur = imported.get("prayer_duration_min", None)
                    try:
                        st.session_state.prayer_duration_min = int(val_prayer_dur) if val_prayer_dur is not None else st.session_state.get("prayer_duration_min", 20)
                    except Exception:
                        st.session_state.prayer_duration_min = st.session_state.get("prayer_duration_min", 20)

                    st.success("âœ… Configuration importÃ©e. Les paramÃ¨tres et sites ont Ã©tÃ© mis Ã  jour.")
                except Exception as e:
                    st.error(f"âŒ Import JSON invalide: {e}")

# --------------------------
# PLANIFICATION
# --------------------------

col1, col2, col3 = st.columns([1, 2, 1])
with col2:
    plan_button = st.button("ðŸš€ Planifier la mission", type="primary", use_container_width=True)

if plan_button:
    # Sauvegarde automatique des donnÃ©es avant planification
    st.session_state.sites_df = sites_df
    # Persistance des paramÃ¨tres saisis
    st.session_state.mission_title = mission_title
    st.session_state.use_base_location = use_base_location
    st.session_state.base_location = base_location
    st.session_state.start_date = start_date
    st.session_state.max_days = max_days
    st.session_state.start_activity_time = start_activity_time
    st.session_state.end_activity_time = end_activity_time
    st.session_state.start_travel_time = start_travel_time
    st.session_state.end_travel_time = end_travel_time
    st.session_state.tolerance_hours = tolerance_hours
    st.session_state.use_lunch = use_lunch
    st.session_state.lunch_start_time = lunch_start_time if use_lunch else None
    st.session_state.lunch_end_time = lunch_end_time if use_lunch else None
    st.session_state.lunch_duration_min = lunch_duration_min if use_lunch else None
    st.session_state.use_prayer = use_prayer
    st.session_state.prayer_start_time = prayer_start_time if use_prayer else None
    st.session_state.prayer_duration_min = prayer_duration_min if use_prayer else None

    # Validations basiques avant planification
    if sites_df is None or sites_df.empty:
        st.error("âŒ Ajoutez au moins un site avant de planifier.")
        st.stop()
    issues = []
    for i, row in sites_df.iterrows():
        city = str(row.get("Ville", "")).strip()
        dur = row.get("DurÃ©e (h)", None)
        if not city:
            issues.append(f"Ligne {i + 1}: Ville manquante ou vide")
        try:
            val = float(dur) if dur is not None else 0
        except Exception:
            val = 0
        if val <= 0:
            issues.append(f"Ligne {i + 1}: DurÃ©e (h) doit Ãªtre > 0")
    if use_base_location and not str(base_location).strip():
        issues.append("Point de dÃ©part/arrivÃ©e activÃ© mais ville non renseignÃ©e")
    if issues:
        st.error("âš ï¸ Veuillez corriger ces points avant la planification:")
        for msg in issues[:10]:
            st.write(f"- {msg}")
        st.stop()
    
    # Animation CSS moderne pour l'attente
    st.markdown("""
    <style>
    .planning-container {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        border-radius: 15px;
        padding: 30px;
        margin: 20px 0;
        text-align: center;
        color: white;
        box-shadow: 0 10px 30px rgba(0,0,0,0.2);
    }
    
    .spinner-icon {
        font-size: 3em;
        animation: spin 2s linear infinite;
        margin-bottom: 20px;
        display: inline-block;
    }
    
    @keyframes spin {
        0% { transform: rotate(0deg); }
        100% { transform: rotate(360deg); }
    }
    
    .pulse-text {
        animation: pulse 1.5s ease-in-out infinite alternate;
        font-size: 1.2em;
        font-weight: bold;
        margin: 10px 0;
    }
    
    @keyframes pulse {
        0% { opacity: 0.6; }
        100% { opacity: 1; }
    }
    
    .progress-enhanced {
        background: linear-gradient(90deg, #667eea 0%, #764ba2 100%);
        border-radius: 15px;
        overflow: hidden;
        margin: 20px 0;
        height: 8px;
        position: relative;
        box-shadow: 0 4px 15px rgba(0,0,0,0.2);
    }
    
    .progress-enhanced::before {
        content: '';
        position: absolute;
        top: 0;
        left: -100%;
        width: 100%;
        height: 100%;
        background: linear-gradient(90deg, transparent, rgba(255,255,255,0.4), transparent);
        animation: shimmer 2s infinite;
    }
    
    @keyframes shimmer {
        0% { left: -100%; }
        100% { left: 100%; }
    }
    
    .step-indicator {
        display: flex;
        justify-content: space-between;
        margin: 20px 0;
        font-size: 0.9em;
        position: relative;
    }
    
    .step-indicator::before {
        content: '';
        position: absolute;
        top: 50%;
        left: 0;
        right: 0;
        height: 2px;
        background: linear-gradient(90deg, rgba(255,255,255,0.3) 0%, rgba(255,255,255,0.6) 50%, rgba(255,255,255,0.3) 100%);
        z-index: 1;
        transform: translateY(-50%);
    }
    
    .step {
        padding: 8px 15px;
        border-radius: 20px;
        background: rgba(255,255,255,0.15);
        transition: all 0.4s cubic-bezier(0.4, 0, 0.2, 1);
        position: relative;
        z-index: 2;
        border: 2px solid transparent;
        backdrop-filter: blur(10px);
        box-shadow: 0 4px 15px rgba(0,0,0,0.1);
    }
    
    .step.active {
        background: linear-gradient(135deg, #4facfe 0%, #00f2fe 100%);
        transform: scale(1.15);
        color: white;
        border: 2px solid rgba(255,255,255,0.5);
        box-shadow: 0 8px 25px rgba(79, 172, 254, 0.4);
        animation: glow 2s ease-in-out infinite alternate;
    }
    
    .step.completed {
        background: linear-gradient(135deg, #56ab2f 0%, #a8e6cf 100%);
        color: white;
        border: 2px solid rgba(255,255,255,0.3);
        box-shadow: 0 4px 15px rgba(86, 171, 47, 0.3);
    }
    
    @keyframes glow {
        0% { box-shadow: 0 8px 25px rgba(79, 172, 254, 0.4); }
        100% { box-shadow: 0 12px 35px rgba(79, 172, 254, 0.6); }
    }
    </style>
    """, unsafe_allow_html=True)
    
    # Container d'animation
    animation_container = st.empty()
    
    with animation_container.container():
        st.markdown("""
        <div class="planning-container">
            <div class="spinner-icon">ðŸ—ºï¸</div>
            <div class="pulse-text">Planification intelligente en cours...</div>
            <div class="step-indicator">
                <span class="step active" id="step-1">ðŸ“ GÃ©ocodage</span>
                <span class="step" id="step-2">ðŸ—ºï¸ Distances</span>
                <span class="step" id="step-3">ðŸ”„ Optimisation</span>
                <span class="step" id="step-4">ðŸ›£ï¸ ItinÃ©raire</span>
                <span class="step" id="step-5">ðŸ“… Planning</span>
            </div>
        </div>
        """, unsafe_allow_html=True)
    
    rows = sites_df.replace({pd.NA: None}).to_dict(orient="records")
    sites = [r for r in rows if r.get("Ville") and str(r["Ville"]).strip()]
    
    if use_base_location and base_location and base_location.strip():
        base_site = {"Ville": base_location.strip(), "Type": "Base", "ActivitÃ©": "DÃ©part", "DurÃ©e (h)": 0}
        return_site = {"Ville": base_location.strip(), "Type": "Base", "ActivitÃ©": "Retour", "DurÃ©e (h)": 0}
        all_sites = [base_site] + sites + [return_site]
        
        if len(sites) < 1:
            st.error("âŒ Ajoutez au moins 1 site Ã  visiter")
            st.stop()
    else:
        all_sites = sites
        # Autoriser la planification avec un seul site (sans base)
        if len(all_sites) < 1:
            st.error("âŒ Ajoutez au moins 1 site Ã  visiter")
            st.stop()
        first_site = all_sites[0].copy()
        first_site["ActivitÃ©"] = "Retour"
        all_sites = all_sites + [first_site]
    
    progress = st.progress(0)
    status = st.empty()
    
    # Fonction pour mettre Ã  jour l'animation avec JavaScript
    def update_animation_step(step_number, icon, message, completed_steps=None):
        if completed_steps is None:
            completed_steps = []
        
        def get_step_class(step_num):
            if step_num in completed_steps:
                return 'completed'
            elif step_num == step_number:
                return 'active'
            else:
                return ''
        
        animation_container.markdown(f"""
        <div class="planning-container">
            <div class="spinner-icon">{icon}</div>
            <div class="pulse-text">{message}</div>
            <div class="step-indicator">
                <span class="step {get_step_class(1)}" id="step-1">ðŸ“ GÃ©ocodage</span>
                <span class="step {get_step_class(2)}" id="step-2">ðŸ—ºï¸ Distances</span>
                <span class="step {get_step_class(3)}" id="step-3">ðŸ”„ Optimisation</span>
                <span class="step {get_step_class(4)}" id="step-4">ðŸ›£ï¸ ItinÃ©raire</span>
                <span class="step {get_step_class(5)}" id="step-5">ðŸ“… Planning</span>
            </div>
            <div class="progress-enhanced"></div>
        </div>
        """, unsafe_allow_html=True)
    
    # Messages dynamiques pour chaque Ã©tape
    geocoding_messages = [
        "ðŸ” Recherche des coordonnÃ©es GPS...",
        "ðŸ“ GÃ©olocalisation des sites en cours...",
        "ðŸŒ Validation des adresses...",
        "âœ… GÃ©ocodage terminÃ© avec succÃ¨s!"
    ]
    
    # Ã‰tape 1: GÃ©ocodage
    update_animation_step(1, "ðŸ“", geocoding_messages[0], [])
    status.text("ðŸ“ GÃ©ocodage...")
    coords = []
    failed = []
    
    for i, s in enumerate(all_sites):
        progress.progress((i+1) / (len(all_sites) * 4))
        # Message dynamique pendant le gÃ©ocodage
        if i < len(geocoding_messages) - 1:
            update_animation_step(1, "ðŸ“", geocoding_messages[min(i, len(geocoding_messages)-2)], [])
        city_val = str(s.get("Ville", "")).strip()
        if s.get("Type") == "Base" and city_val.lower() == "dakar":
            coord = (-17.470602, 14.711404)
        else:
            coord = geocode_city_senegal(city_val, use_cache)
        if not coord:
            failed.append(s["Ville"])
        else:
            coords.append(coord)
    
    update_animation_step(1, "âœ…", geocoding_messages[-1], [1])
    
    if failed:
        st.error(f"âŒ Villes introuvables: {', '.join(failed)}")
        st.stop()
    
    # Ã‰tape 2: Calcul des distances
    distance_messages = [
        "ðŸ—ºï¸ Connexion aux services de cartographie...",
        "ðŸ“ Calcul des distances entre les sites...",
        "â±ï¸ Estimation des temps de trajet...",
        "âœ… Matrice de distances calculÃ©e!"
    ]
    
    update_animation_step(2, "ðŸ—ºï¸", distance_messages[0], [1])
    status.text("ðŸ—ºï¸ Calcul des distances...")
    progress.progress(0.4)
    
    durations_sec = None
    distances_m = None
    calculation_method = ""
    city_list = [s["Ville"] for s in all_sites]
    
    if distance_method.startswith("Maps uniquement"):
        update_animation_step(2, "ðŸ—ºï¸", distance_messages[1], [1])
        durations_sec, distances_m, error_msg = improved_graphhopper_duration_matrix(graphhopper_api_key, coords)
        calculation_method = "Maps"
        if durations_sec is None:
            st.error(f"âŒ {error_msg}")
            st.stop()
        else:
            if debug_mode:
                st.info(f"ðŸ” Debug Maps: {len(durations_sec)} x {len(durations_sec[0]) if durations_sec else 0} matrice de durÃ©es reÃ§ue")
                if durations_sec and len(durations_sec) > 0:
                    sample = durations_sec[0][1] if len(durations_sec[0]) > 1 else 0
                    st.info(f"ðŸ” Debug Maps: Exemple durÃ©e [0][1] = {sample} secondes ({sample/3600:.2f}h)")

    elif distance_method == "OSRM uniquement (rapide)":
        update_animation_step(2, "ðŸ—ºï¸", distance_messages[1], [1])
        durations_sec, distances_m, error_msg = improved_osrm_duration_matrix(osrm_base_url, coords)
        calculation_method = "OSRM"
        if durations_sec is None:
            st.error(f"âŒ {error_msg}")
            st.stop()
        else:
            if debug_mode:
                st.info(f"ðŸ” Debug OSRM: {len(durations_sec)} x {len(durations_sec[0]) if durations_sec else 0} matrice de durÃ©es reÃ§ue")
                if durations_sec and len(durations_sec) > 0:
                    sample = durations_sec[0][1] if len(durations_sec[0]) > 1 else 0
                    st.info(f"ðŸ” Debug OSRM: Exemple durÃ©e [0][1] = {sample} secondes ({sample/3600:.2f}h)")

    elif distance_method == "Automatique uniquement":
        result, error_msg = improved_deepseek_estimate_matrix(city_list, deepseek_api_key, debug_mode)
        if result:
            durations_sec, distances_m = result
            calculation_method = "Automatique"
            st.info(f"ðŸ“Š MÃ©thode: {calculation_method}")
        else:
            st.error(f"âŒ {error_msg}")
            st.stop()

    elif distance_method == "GÃ©omÃ©trique uniquement":
        durations_sec, distances_m = haversine_fallback_matrix(coords, default_speed_kmh)
        calculation_method = f"GÃ©omÃ©trique ({default_speed_kmh} km/h)"
        st.warning(f"ðŸ“Š MÃ©thode: {calculation_method}")

    else:
        # Mode Auto: dynamique
        # - Site unique â†’ OSRM â†’ Automatique â†’ Maps â†’ GÃ©omÃ©trique
        # - Plusieurs sites â†’ OSRM â†’ Automatique â†’ Maps â†’ GÃ©omÃ©trique
        single_site = len(sites) == 1

        if single_site:
            update_animation_step(2, "ðŸ—ºï¸", distance_messages[1], [1])
            # 1) OSRM en premier
            durations_sec, distances_m, error_msg = improved_osrm_duration_matrix(osrm_base_url, coords)
            if durations_sec is not None:
                calculation_method = "OSRM"
            else:
                # 2) Automatique (DeepSeek)
                if deepseek_api_key:
                    result, _ = improved_deepseek_estimate_matrix(city_list, deepseek_api_key, debug_mode)
                    if result:
                        durations_sec, distances_m = result
                        calculation_method = "Automatique"
                    else:
                        # 3) Maps (GraphHopper), si option activÃ©e
                        if use_deepseek_fallback and graphhopper_api_key:
                            durations_sec, distances_m, error_msg = improved_graphhopper_duration_matrix(graphhopper_api_key, coords)
                            if durations_sec is not None:
                                calculation_method = "Maps"
                            else:
                                durations_sec, distances_m = haversine_fallback_matrix(coords, default_speed_kmh)
                                calculation_method = f"GÃ©omÃ©trique ({default_speed_kmh} km/h)"
                        else:
                            durations_sec, distances_m = haversine_fallback_matrix(coords, default_speed_kmh)
                            calculation_method = f"GÃ©omÃ©trique ({default_speed_kmh} km/h)"
                else:
                    # Pas de clÃ© DeepSeek, tenter Maps si autorisÃ© puis gÃ©omÃ©trique
                    if use_deepseek_fallback and graphhopper_api_key:
                        durations_sec, distances_m, error_msg = improved_graphhopper_duration_matrix(graphhopper_api_key, coords)
                        if durations_sec is not None:
                            calculation_method = "Maps"
                        else:
                            durations_sec, distances_m = haversine_fallback_matrix(coords, default_speed_kmh)
                            calculation_method = f"GÃ©omÃ©trique ({default_speed_kmh} km/h)"
                    else:
                        durations_sec, distances_m = haversine_fallback_matrix(coords, default_speed_kmh)
                        calculation_method = f"GÃ©omÃ©trique ({default_speed_kmh} km/h)"
        else:
            # OSRM â†’ Automatique â†’ Maps (si activÃ©) â†’ GÃ©omÃ©trique
            durations_sec, distances_m, error_msg = improved_osrm_duration_matrix(osrm_base_url, coords)
            if durations_sec is not None:
                calculation_method = "OSRM"
            else:
                result, error_msg = improved_deepseek_estimate_matrix(city_list, deepseek_api_key, debug_mode)
                if result:
                    durations_sec, distances_m = result
                    calculation_method = "Automatique"
                else:
                    if use_deepseek_fallback and graphhopper_api_key:
                        durations_sec, distances_m, error_msg = improved_graphhopper_duration_matrix(graphhopper_api_key, coords)
                        if durations_sec is not None:
                            calculation_method = "Maps"
                        else:
                            durations_sec, distances_m = haversine_fallback_matrix(coords, default_speed_kmh)
                            calculation_method = f"GÃ©omÃ©trique ({default_speed_kmh} km/h)"
                    else:
                        durations_sec, distances_m = haversine_fallback_matrix(coords, default_speed_kmh)
                        calculation_method = f"GÃ©omÃ©trique ({default_speed_kmh} km/h)"

        method_color = "success" if ("Maps" in calculation_method or "OSRM" in calculation_method) else "info" if "Automatique" in calculation_method else "warning"
        getattr(st, method_color)(f"ðŸ“Š MÃ©thode: {calculation_method}")
    
    # Ã‰tape 3: Optimisation (commune Ã  tous les modes)
    update_animation_step(3, "ðŸ”„", "Optimisation de l'itinÃ©raire...", [1, 2])
    status.text("ðŸ”„ Optimisation de l'ordre des sites...")
    progress.progress(0.6)
    
    # DÃ©terminer l'ordre des sites selon le mode choisi
    if order_mode == "âœ‹ Manuel (personnalisÃ©)":
        # Utiliser l'ordre manuel dÃ©fini par l'utilisateur
        if use_base_location and base_location and base_location.strip():
            # Avec base: [base] + sites_manuels + [base]
            manual_sites_order = [0]  # Base de dÃ©part
            for manual_idx in st.session_state.manual_order:
                if manual_idx < len(sites):
                    manual_sites_order.append(manual_idx + 1)  # +1 car base est Ã  l'index 0
            manual_sites_order.append(len(all_sites) - 1)  # Base de retour
            order = manual_sites_order
        else:
            # Sans base: sites_manuels + [premier_site]
            manual_sites_order = []
            for manual_idx in st.session_state.manual_order:
                if manual_idx < len(sites):
                    manual_sites_order.append(manual_idx)
            manual_sites_order.append(len(all_sites) - 1)  # Site de retour
            order = manual_sites_order
        
        st.success("âœ… Ordre manuel appliquÃ©")
    else:
        # Utiliser l'optimisation IA Adja au lieu du TSP traditionnel
        if len(coords) >= 3:
            # Essayer d'abord l'optimisation IA Adja
            ai_order, ai_success, ai_message = optimize_route_with_ai(
                all_sites, coords, 
                base_location if use_base_location else None, 
                deepseek_api_key
            )
            
            if ai_success:
                order = ai_order
                st.success(f"âœ… Ordre optimisÃ© par IA Adja: {ai_message}")
            else:
                # Fallback vers TSP si l'IA Adja Ã©choue
                order = solve_tsp_fixed_start_end(durations_sec)
                st.warning(f"âš ï¸ IA Adja Ã©chouÃ©e ({ai_message}), utilisation TSP classique")
        else:
            order = list(range(len(coords)))
            st.success("âœ… Ordre sÃ©quentiel (moins de 3 sites)")
            
        if debug_mode and durations_sec:
            # Calculer coÃ»t total pour transparence
            total_cost = sum(durations_sec[order[i]][order[i+1]] for i in range(len(order)-1))
            st.info(f"ðŸ” Debug Optimisation: ordre={order} | coÃ»t total={total_cost/3600:.2f}h")
        
    status.text("ðŸ›£ï¸ Calcul de l'itinÃ©raire dÃ©taillÃ©...")
    # Ã‰tape 4: GÃ©nÃ©ration de l'itinÃ©raire
    update_animation_step(4, "ðŸ›£ï¸", "GÃ©nÃ©ration de l'itinÃ©raire dÃ©taillÃ©...", [1, 2, 3])
    progress.progress(0.8)
    
    segments = []
    zero_segments_indices = []
    
    for i in range(len(order)-1):
        from_idx = order[i]
        to_idx = order[i+1]
        
        if from_idx < len(durations_sec) and to_idx < len(durations_sec[0]):
            duration = durations_sec[from_idx][to_idx]
            distance = distances_m[from_idx][to_idx] if distances_m else 0
            segment_method = "Matrix"
            
            # Si la distance/durÃ©e est nulle, recalculer via OSRM/Maps avec cache, puis fallback gÃ©omÃ©trique
            if duration == 0 or distance == 0:
                # Cache par segment
                if 'segment_route_cache' not in st.session_state:
                    st.session_state.segment_route_cache = {}
                segment_cache = st.session_state.segment_route_cache
                SEGMENT_CACHE_TTL = int(st.session_state.get('segment_cache_ttl_sec', 43200))
                
                coord_from = coords[from_idx]
                coord_to = coords[to_idx]
                key = (coord_from[0], coord_from[1], coord_to[0], coord_to[1])
                now_ts = datetime.now().timestamp()
                
                # 1) Cache
                cached = segment_cache.get(key)
                if cached and (now_ts - cached.get('ts', 0)) < SEGMENT_CACHE_TTL:
                    distance = int(cached.get('distance', 0))
                    duration = int(cached.get('duration', 0))
                    segment_method = cached.get('method', 'Matrix')
                else:
                    # 2) OSRM route
                    try:
                        coord_str = f"{coord_from[0]},{coord_from[1]};{coord_to[0]},{coord_to[1]}"
                        url = f"{osrm_base_url.rstrip('/')}/route/v1/driving/{coord_str}?overview=false"
                        resp = requests.get(url, timeout=10)
                        if resp.status_code == 200:
                            data = resp.json()
                            if data.get('routes'):
                                r0 = data['routes'][0]
                                d_m = int(r0.get('distance', 0))
                                t_s = int(r0.get('duration', 0))
                                if d_m > 0 and t_s > 0:
                                    distance = d_m
                                    duration = t_s
                                    segment_method = "OSRM"
                    except Exception:
                        pass
                    
                    # 3) GraphHopper route
                    if (duration == 0 or distance == 0) and graphhopper_api_key:
                        try:
                            gh_url = "https://graphhopper.com/api/1/route"
                            qp = f"point={coord_from[1]},{coord_from[0]}&point={coord_to[1]},{coord_to[0]}&vehicle=car&locale=fr&points_encoded=false&calc_points=false&key={graphhopper_api_key}"
                            gh_resp = requests.get(f"{gh_url}?{qp}", timeout=10)
                            if gh_resp.status_code == 200:
                                gh_data = gh_resp.json()
                                paths = gh_data.get('paths')
                                if paths:
                                    p0 = paths[0]
                                    d_m = int(p0.get('distance', 0))
                                    t_ms = int(p0.get('time', 0))
                                    t_s = int(t_ms / 1000)
                                    if d_m > 0 and t_s > 0:
                                        distance = d_m
                                        duration = t_s
                                        segment_method = "Maps"
                        except Exception:
                            pass
                    
                    # 4) Fallback gÃ©omÃ©trique
                    if duration == 0 or distance == 0:
                        geometric_km = haversine(coord_from[0], coord_from[1], coord_to[0], coord_to[1]) * 1.2
                        if distance == 0:
                            distance = int(geometric_km * 1000)
                        if duration == 0:
                            distance_for_time_calc = distance / 1000 if distance > 0 else geometric_km
                            geometric_hours = distance_for_time_calc / default_speed_kmh
                            duration = int(geometric_hours * 3600)
                        segment_method = "Geo"
                    
                    # Mettre Ã  jour le cache
                    segment_cache[key] = {
                        "distance": int(distance),
                        "duration": int(duration),
                        "method": segment_method,
                        "ts": now_ts
                    }
                
                zero_segments_indices.append(i)
                if debug_mode:
                    st.info(f"ðŸ” Segment {i} recalculÃ© via {segment_method}: {distance/1000:.1f}km, {duration/3600:.2f}h")
            
            # Debug: Afficher les valeurs des segments
            if debug_mode:
                st.info(f"ðŸ” Debug Segment {i}: de {from_idx} vers {to_idx} = {duration}s ({duration/3600:.2f}h), {distance/1000:.1f}km")
            
            segments.append({
                "distance": distance,
                "duration": duration,
                "method": segment_method
            })
        else:
            segments.append({"distance": 0, "duration": 0})
    
    if not segments:
        st.error("âŒ AUCUN segment crÃ©Ã©!")
        st.stop()
    
    # RÃ©sumÃ© des mÃ©thodes de recalcul des segments
    if zero_segments_indices:
        method_summary = {"OSRM": 0, "Maps": 0, "Geo": 0}
        for idx in zero_segments_indices:
            m = segments[idx].get("method", "")
            if m in method_summary:
                method_summary[m] += 1
        summary_str = " | ".join([
            f"OSRM: {method_summary['OSRM']}",
            f"Maps: {method_summary['Maps']}",
            f"GÃ©o: {method_summary['Geo']}"
        ])
        st.success(f"âœ… {len(zero_segments_indices)} segment(s) recalculÃ©(s) ({summary_str})")
    
    # VÃ©rifier s'il reste des segments Ã  zÃ©ro aprÃ¨s le recalcul
    remaining_zero_segments = [i for i, s in enumerate(segments) if s['duration'] == 0 or s['distance'] == 0]
    if remaining_zero_segments:
        st.warning(f"âš ï¸ {len(remaining_zero_segments)} segment(s) avec valeurs manquantes aprÃ¨s recalcul")
    
    status.text("ðŸ“… GÃ©nÃ©ration du planning dÃ©taillÃ©...")
    # Ã‰tape 5: GÃ©nÃ©ration du planning
    update_animation_step(5, "ðŸ“…", "Finalisation du planning...", [1, 2, 3, 4])
    progress.progress(0.9)

    # Calcul prÃ©alable: nombre de jours optimal (dry-run, sans Ã©tirement)
    _, _, _, dry_stats = schedule_itinerary(
        coords=coords,
        sites=all_sites,
        order=order,
        segments_summary=segments,
        start_date=start_date,
        start_activity_time=start_activity_time,
        end_activity_time=end_activity_time,
        start_travel_time=start_travel_time,
        end_travel_time=end_travel_time,
        use_lunch=use_lunch,
        lunch_start_time=lunch_start_time if use_lunch else time(12,30),
        lunch_end_time=lunch_end_time if use_lunch else time(14,0),
        use_prayer=use_prayer,
        prayer_start_time=prayer_start_time if use_prayer else time(14,0),
        prayer_duration_min=prayer_duration_min if use_prayer else 20,
        lunch_duration_min=st.session_state.get("lunch_duration_min", 60),
        max_days=0,
        tolerance_hours=tolerance_hours,
        base_location=base_location,
        allow_weekend_travel=allow_weekend_travel,
        allow_weekend_activities=allow_weekend_activities
    )

    optimal_days = int(dry_stats.get('total_days', 1))
    user_max = int(max_days) if isinstance(max_days, (int, float)) else 0
    user_desired = int(desired_days) if isinstance(desired_days, (int, float)) else 0

    # Logique de dÃ©cision pour les jours effectifs
    if user_desired > 0:
        if user_max > 0 and user_desired > user_max:
            st.warning(f"Le nombre de jours souhaitÃ©s ({user_desired}) dÃ©passe le maximum autorisÃ© ({user_max}). Utilisation du maximum.")
            effective_max_days = user_max
        else:
            effective_max_days = user_desired

        if effective_max_days < optimal_days:
            # Cas compression: on tente de tenir en moins de jours (journÃ©es plus chargÃ©es)
            stretch_days_flag = False
            st.warning(f"âš ï¸ Objectif ({effective_max_days} jours) < optimal ({optimal_days}). Compression: journÃ©es potentiellement plus chargÃ©es.")
        elif effective_max_days > optimal_days:
            # Cas Ã©talement: on rÃ©partit sur plus de jours, fin de journÃ©e plus tÃ´t
            stretch_days_flag = True
            st.success(f"âœ… Planning Ã©talÃ© sur {effective_max_days} jours (optimal: {optimal_days}). JournÃ©es plus lÃ©gÃ¨res.")
        else:
            # Ã‰gal Ã  l'optimal
            stretch_days_flag = False
            st.info(f"ðŸŸ° Planning sur {effective_max_days} jours, Ã©gal Ã  l'optimal.")

    elif user_max > 0:
        if user_max < optimal_days:
            effective_max_days = user_max
            stretch_days_flag = True
            st.warning(f"âš ï¸ Objectif ({user_max} jours) < optimal ({optimal_days}). Compression en {user_max} jours avec journÃ©es Ã©tirÃ©es.")
        else:
            effective_max_days = user_max
            stretch_days_flag = False
            st.success(f"âœ… Le planning tient en {optimal_days} jours (objectif max: {user_max} jours).")
    else:
        effective_max_days = optimal_days
        stretch_days_flag = False
        st.info(f"ðŸ§® Jours optimaux calculÃ©s automatiquement: {optimal_days} jour(s).")

    # Planification finale avec paramÃ¨tres effectifs
    itinerary, sites_ordered, coords_ordered, stats = schedule_itinerary(
        coords=coords,
        sites=all_sites,
        order=order,
        segments_summary=segments,
        start_date=start_date,
        start_activity_time=start_activity_time,
        end_activity_time=end_activity_time,
        start_travel_time=start_travel_time,
        end_travel_time=end_travel_time,
        use_lunch=use_lunch,
        lunch_start_time=lunch_start_time if use_lunch else time(12,30),
        lunch_end_time=lunch_end_time if use_lunch else time(14,0),
        use_prayer=use_prayer,
        prayer_start_time=prayer_start_time if use_prayer else time(14,0),
        prayer_duration_min=prayer_duration_min if use_prayer else 20,
        lunch_duration_min=st.session_state.get("lunch_duration_min", 60),
        max_days=effective_max_days,
        tolerance_hours=tolerance_hours,
        base_location=base_location,
        stretch_days=stretch_days_flag,
        allow_weekend_travel=allow_weekend_travel,
        allow_weekend_activities=allow_weekend_activities
    )

    if stretch_days_flag and stats.get('total_days', 0) > effective_max_days:
        st.error(f"âŒ Impossible de tenir en {effective_max_days} jour(s). Besoin de {stats.get('total_days')} jours mÃªme en Ã©tirant les journÃ©es.")
    
    progress.progress(1.0)
    status.text("âœ… TerminÃ©!")
    
    st.session_state.planning_results = {
        'itinerary': itinerary,
        'sites_ordered': sites_ordered,
        'coords_ordered': coords_ordered,
        'route_polyline': None,
        'stats': stats,
        'start_date': start_date,
        'calculation_method': calculation_method,
        'segments_summary': segments,
        'original_order': order.copy(),  # Sauvegarder l'ordre original
        'durations_matrix': durations_sec,
        'distances_matrix': distances_m,
        'all_coords': coords,
        'base_location': base_location
    }
    st.session_state.manual_itinerary = None
    st.session_state.edit_mode = False

    # Nettoyer l'animation et la barre de progression pour Ã©viter le spinner persistant
    try:
        progress.empty()
    except Exception:
        pass
    try:
        animation_container.empty()
    except Exception:
        pass

# --------------------------
# AFFICHAGE RÃ‰SULTATS
# --------------------------
if st.session_state.planning_results:
    results = st.session_state.planning_results
    itinerary = st.session_state.manual_itinerary if st.session_state.manual_itinerary else results['itinerary']
    sites_ordered = results['sites_ordered']
    coords_ordered = results['coords_ordered']
    stats = results['stats']
    start_date = results['start_date']
    calculation_method = results.get('calculation_method', 'Inconnu')
    segments_summary = results.get('segments_summary', [])
    
    st.header("ðŸ“Š RÃ©sumÃ© de la mission")
    
    method_color = "success" if "Maps" in calculation_method else "info" if "Automatique" in calculation_method else "warning"
    # Message supprimÃ© pour allÃ©ger l'UI
# st.caption("ðŸ“Š Distances calculÃ©es")
    
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("DurÃ©e totale", f"{stats['total_days']} jour(s)")
    with col2:
        st.metric("Distance totale", f"{stats['total_km']:.1f} km")
    with col3:
        # Compter seulement les vrais sites (exclure les sites de type "Base")
        actual_sites_count = len([site for site in sites_ordered if site.get('Type') != 'Base'])
        st.metric("Sites visitÃ©s", f"{actual_sites_count}")
    with col4:
        st.metric("Temps de visite", f"{stats['total_visit_hours']:.1f} h")
    
    tab_planning, tab_map, tab_fuel, tab_edit, tab_manual, tab_report, tab_export = st.tabs(["ðŸ“… Planning", "ðŸ—ºï¸ Carte", "â›½ Carburant", "âœï¸ Ã‰diter", "ðŸ”„ Modifier ordre", "ðŸ“‹ Rapport", "ðŸ’¾ Export"])
    
    with tab_planning:
        st.subheader("Planning dÃ©taillÃ©")
        
        view_mode = st.radio(
            "Mode d'affichage",
            ["ðŸ“‹ Vue interactive", "ðŸŽ¨ PrÃ©sentation professionnelle"],
            horizontal=True,
            index=1
        )
        
        if view_mode == "ðŸŽ¨ PrÃ©sentation professionnelle":
            include_map_prof = st.checkbox("Inclure la carte", value=st.session_state.get("include_map_prof_html", False))
            st.session_state.include_map_prof_html = include_map_prof

            include_prof_details = st.checkbox(
                "Inclure section rÃ©sumÃ©",
                value=st.session_state.get("include_prof_details", False)
            )
            st.session_state.include_prof_details = include_prof_details

            html_str = build_professional_html(
                itinerary,
                start_date,
                stats,
                sites_ordered,
                segments_summary,
                default_speed_kmh,
                mission_title,
                coords_ordered,
                include_map=include_map_prof,
                lunch_start_time=st.session_state.get("lunch_start_time"),
                lunch_end_time=st.session_state.get("lunch_end_time"),
                lunch_duration_min=st.session_state.get("lunch_duration_min", 60),
                prayer_start_time=st.session_state.get("prayer_start_time"),
                prayer_duration_min=st.session_state.get("prayer_duration_min", 20),
                include_details=include_prof_details
            )
            st.components.v1.html(html_str, height=1100, scrolling=True)
            
            col_html, col_pdf = st.columns(2)
            
            with col_html:
                st.download_button(
                    label="ðŸ“¥ TÃ©lÃ©charger HTML",
                    data=html_str,
                    file_name=f"mission_{datetime.now().strftime('%Y%m%d_%H%M')}.html",
                    mime="text/html"
                )
            
            with col_pdf:
                try:
                    excel_data = create_mission_excel(
                        itinerary=itinerary,
                        start_date=start_date,
                        stats=stats,
                        sites_ordered=sites_ordered,
                        segments_summary=segments_summary,
                        mission_title=mission_title
                    )
                    st.download_button(
                        label="ðŸ“Š TÃ©lÃ©charger Excel",
                        data=excel_data,
                        file_name=f"mission_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                    )
                except Exception as e:
                    st.error(f"âŒ Erreur lors de la gÃ©nÃ©ration du fichier Excel: {str(e)}")
        
        else:
            total_days = max(ev[0] for ev in itinerary) if itinerary else 1
            
            if total_days > 1:
                selected_day = st.selectbox(
                    "Jour",
                    options=range(1, total_days + 1),
                    format_func=lambda x: f"Jour {x} - {(start_date + timedelta(days=x-1)).strftime('%d/%m/%Y')}"
                )
            else:
                selected_day = 1
            
            day_events = [ev for ev in itinerary if ev[0] == selected_day]
            
            if day_events:
                date_str = (start_date + timedelta(days=selected_day-1)).strftime("%A %d %B %Y")
                st.info(f"**{date_str}**")
                
                for day, sdt, edt, desc in day_events:
                    col1, col2 = st.columns([1, 3])
                    with col1:
                        st.write(f"**{sdt.strftime('%H:%M')} - {edt.strftime('%H:%M')}**")
                    with col2:
                        if "â†’" in desc:
                            st.write(f"ðŸš— {desc}")
                        elif "Visite" in desc or "Site" in desc or "Client" in desc:
                            st.success(desc)
                        elif "Pause" in desc or "DÃ©jeuner" in desc or "PriÃ¨re" in desc:
                            st.info(desc)
                        elif "NuitÃ©e" in desc:
                            st.warning(desc)
                        else:
                            st.write(desc)
    
    with tab_fuel:
        st.subheader("â›½ Module Carburant")
        
        if st.session_state.planning_results is None:
            st.warning("âš ï¸ Veuillez d'abord gÃ©nÃ©rer un planning dans l'onglet 'Planning' pour calculer la consommation de carburant.")
        else:
            # RÃ©cupÃ©rer la distance totale du planning
            stats = st.session_state.planning_results.get('stats', {})
            total_distance_km = stats.get('total_km', 0)
            
            if total_distance_km > 0:
                st.info(f"ðŸ“ **Distance totale de la mission :** {total_distance_km:.1f} km")
                
                # SÃ©lection du type de vÃ©hicule
                st.subheader("ðŸš— SÃ©lection du vÃ©hicule")
                
                vehicle_types = get_vehicle_types()
                vehicle_names = list(vehicle_types.keys())
                
                # Station-Wagon par dÃ©faut
                default_index = vehicle_names.index("Station-Wagon") if "Station-Wagon" in vehicle_names else 0
                
                selected_vehicle = st.selectbox(
                    "Type de vÃ©hicule",
                    options=vehicle_names,
                    index=default_index,
                    help="SÃ©lectionnez le type de vÃ©hicule pour calculer la consommation"
                )
                
                # Affichage des caractÃ©ristiques du vÃ©hicule sÃ©lectionnÃ©
                vehicle_info = vehicle_types[selected_vehicle]
                
                col1, col2 = st.columns(2)
                with col1:
                    st.metric("Consommation", f"{vehicle_info['consumption']} L/100km")
                with col2:
                    st.metric("Facteur COâ‚‚", f"{vehicle_info['co2_factor']} kg COâ‚‚/L")
                
                st.divider()
                
                # Calculs de consommation et d'empreinte carbone
                fuel_data = calculate_fuel_consumption(total_distance_km, selected_vehicle)
                carbon_data = calculate_carbon_footprint(fuel_data, total_distance_km, selected_vehicle)
                cost_data = estimate_fuel_cost(fuel_data)
                
                # Affichage des rÃ©sultats
                st.subheader("ðŸ“Š RÃ©sultats des calculs")
                
                # MÃ©triques principales
                col1, col2 = st.columns(2)
                
                with col1:
                    st.metric(
                        "ðŸ›¢ï¸ Carburant nÃ©cessaire",
                        f"{fuel_data['fuel_needed_liters']:.1f} L",
                        help="QuantitÃ© de carburant nÃ©cessaire pour la mission"
                    )
                
                with col2:
                    st.metric(
                        "ðŸŒ COâ‚‚ Ã©mis",
                        f"{carbon_data['co2_emissions_kg']:.1f} kg",
                        help="Ã‰missions de COâ‚‚ pour la mission"
                    )
                
                st.divider()
                
                # DÃ©tails de l'empreinte carbone
                st.subheader("ðŸŒ± Empreinte carbone dÃ©taillÃ©e")
                
                col1, col2 = st.columns(2)
                
                with col1:
                    st.write("**Ã‰missions COâ‚‚ :**")
                    st.write(f"â€¢ En kilogrammes : **{carbon_data['co2_emissions_kg']:.2f} kg**")
                    st.write(f"â€¢ En tonnes : **{carbon_data['co2_emissions_tons']:.3f} tonnes**")
                
                with col2:
                    st.write("**Ã‰quivalence environnementale :**")
                    st.write(f"â€¢ Arbres Ã  planter pour compenser : **{carbon_data['trees_equivalent']:.0f} arbres**")
                    st.write("â€¢ *(1 arbre absorbe ~22 kg COâ‚‚/an)*")

                # Message d'engagement environnemental
                try:
                    trees_to_plant = int(carbon_data['trees_equivalent'] + 0.9999)  # Arrondi Ã  l'entier supÃ©rieur
                except Exception:
                    trees_to_plant = int(round(carbon_data.get('trees_equivalent', 0)))

                st.warning(
                    f"ðŸŒ¿ Pour un engagement en faveur de l'environnement, engagez-vous Ã  planter au moins "
                    f"**{trees_to_plant} arbre(s)** lors de votre mission."
                )

                st.info(
                    "Conseils Ã©co-responsables: privilÃ©giez l'Ã©co-conduite et le covoiturage lors des missions, maintenez une pression "
                    "des pneus optimale, limitez la climatisation et optimisez vos trajets "
                    "pour rÃ©duire les kilomÃ¨tres Ã  vide."
                )
                
                st.divider()
                
                # Section demande de carburant
                st.subheader("ðŸ“‹ Demande de vÃ©hicule ou de carburant")
                
                col_btn1, col_btn2 = st.columns(2)
                with col_btn1:
                    if st.button("ðŸ“ GÃ©nÃ©rer demande de carburant", type="primary", use_container_width=True):
                        st.session_state.show_fuel_request_modal = True
                
                with col_btn2:
                    st.info("ðŸ’¡ GÃ©nÃ¨re un document Word")
                
                # Modal pour la demande de carburant
                if st.session_state.get('show_fuel_request_modal', False):
                    with st.container():
                        st.markdown("---")
                        st.subheader("ðŸ“‹ Informations pour la demande de carburant")
                        st.info("ðŸ’¡ Remplissez les informations manquantes pour gÃ©nÃ©rer le document")
                        
                        # Afficher les informations de la mission si disponibles
                        if st.session_state.planning_results:
                            stats = st.session_state.planning_results['stats']
                            with st.expander("ðŸ“Š Informations de la mission", expanded=True):
                                col_info1, col_info2, col_info3 = st.columns(3)
                                with col_info1:
                                    st.metric("ðŸ—“ï¸ DurÃ©e", f"{stats['total_days']} jour(s)")
                                with col_info2:
                                    # Utiliser le nombre de sites configurÃ©s par l'utilisateur
                                    nb_sites = len(st.session_state.sites_df) if 'sites_df' in st.session_state else 0
                                    st.metric("ðŸ“ Sites Ã  visiter", f"{nb_sites}")
                                with col_info3:
                                    st.metric("ðŸ›£ï¸ Distance totale", f"{stats.get('total_km', 0):.1f} km")
                        
                        # Informations du demandeur
                        col_req1, col_req2 = st.columns(2)
                        
                        with col_req1:
                            st.markdown("**ðŸ‘¤ Informations du demandeur**")
                            demandeur_nom = st.text_input("Nom et qualitÃ© du demandeur", 
                                                        value="",
                                                        placeholder="Ex: Moctar TALL Responsable de projets",
                                                        key="fuel_req_nom")
                            demandeur_dir = st.text_input("Direction/DÃ©partement", 
                                                        value="",
                                                        placeholder="Ex: DAL/GPR",
                                                        key="fuel_req_dir")
                            demandeur_cr = st.text_input("NÂ° C.R.", 
                                                       value="",
                                                       placeholder="Ex: L2100",
                                                       key="fuel_req_cr")
                            demandeur_tel = st.text_input("NÂ° TÃ©lÃ©phone", 
                                                        value="",
                                                        placeholder="Ex: 77 639 96 12",
                                                        key="fuel_req_tel")
                        
                        with col_req2:
                            st.markdown("**ðŸ“‹ DÃ©tails de la mission**")
                            motif_demande = st.text_area("Motif de la demande", 
                                                       value=mission_title,
                                                       key="fuel_req_motif",
                                                       height=100)
                            
                            col_nb_pers, col_carburant = st.columns(2)
                            with col_nb_pers:
                                nb_personnes = st.number_input("Nombre de personnes", 
                                                             min_value=1, max_value=20, value=2,
                                                             key="fuel_req_nb_pers")
                            
                            with col_carburant:
                                # Utiliser le mÃªme calcul que dans les rÃ©sultats des calculs
                                default_fuel = 50
                                if st.session_state.planning_results:
                                    distance = st.session_state.planning_results['stats'].get('total_km', 0)
                                    if distance > 0:
                                        # Utiliser le mÃªme calcul que dans la section "Carburant nÃ©cessaire"
                                        # Par dÃ©faut, utiliser Station-Wagon (8.5 L/100km)
                                        fuel_data = calculate_fuel_consumption(distance, "Station-Wagon")
                                        default_fuel = int(fuel_data['fuel_needed_liters'])
                                    else:
                                        default_fuel = 50
                                
                                quantite_carburant = st.number_input("QuantitÃ© de carburant (litres)", 
                                                                   min_value=0, max_value=1000, 
                                                                   value=default_fuel,
                                                                   key="fuel_req_quantity",
                                                                   help="QuantitÃ© calculÃ©e automatiquement selon la distance et le type de vÃ©hicule (Station-Wagon par dÃ©faut)")
                        
                        # Dates automatiquement rÃ©cupÃ©rÃ©es du planning
                        if st.session_state.planning_results:
                            # RÃ©cupÃ©rer les dates du planning
                            planning_start_date = st.session_state.planning_results['start_date']
                            itinerary = st.session_state.manual_itinerary or st.session_state.planning_results['itinerary']
                            stats = st.session_state.planning_results['stats']
                            
                            # Calculer la date de retour (date de dÃ©but + nombre de jours - 1)
                            planning_end_date = planning_start_date + timedelta(days=stats['total_days'] - 1)
                            
                            # Afficher les dates rÃ©cupÃ©rÃ©es du planning
                            col_date1, col_date2 = st.columns(2)
                            with col_date1:
                                st.markdown("**ðŸ“… Date de dÃ©part**")
                                st.info(f"ðŸ—“ï¸ {planning_start_date.strftime('%d/%m/%Y')}")
                                date_depart = planning_start_date
                            with col_date2:
                                st.markdown("**ðŸ“… Date de retour**")
                                st.info(f"ðŸ—“ï¸ {planning_end_date.strftime('%d/%m/%Y')}")
                                date_retour = planning_end_date
                        else:
                            # Si pas de planning, utiliser les champs manuels
                            col_date1, col_date2 = st.columns(2)
                            with col_date1:
                                date_depart = st.date_input("Date de dÃ©part prÃ©vue", 
                                                          value=datetime.now().date(),
                                                          key="fuel_req_date_dep")
                            with col_date2:
                                date_retour = st.date_input("Date de retour prÃ©vue", 
                                                          value=datetime.now().date(),
                                                          key="fuel_req_date_ret")
                        
                        # Boutons d'action
                        col_action1, col_action2, col_action3 = st.columns(3)
                        
                        with col_action1:
                            if st.button("ðŸ“„ GÃ©nÃ©rer document Word", type="primary", use_container_width=True):
                                # GÃ©nÃ©rer le document Word
                                try:
                                    from docx import Document
                                    from docx.shared import Inches, Pt, Cm
                                    from docx.enum.text import WD_ALIGN_PARAGRAPH
                                    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
                                    from docx.oxml.shared import OxmlElement, qn
                                    from docx.oxml.ns import nsdecls
                                    from docx.oxml import parse_xml
                                    import io
                                    
                                    # CrÃ©er le document
                                    doc = Document()
                                    
                                    # DÃ©finir les marges
                                    sections = doc.sections
                                    for section in sections:
                                        section.top_margin = Cm(2)
                                        section.bottom_margin = Cm(2)
                                        section.left_margin = Cm(2)
                                        section.right_margin = Cm(2)
                                    
                                    # En-tÃªte principal
                                    header_para = doc.add_paragraph()
                                    header_run = header_para.add_run('DEMANDE DE CARBURANT')
                                    header_run.font.name = 'Tahoma'
                                    header_run.font.size = Pt(14)
                                    header_run.bold = True
                                    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    
                                    # Sous-titre
                                    subtitle_para = doc.add_paragraph()
                                    subtitle_run = subtitle_para.add_run('A remplir et Ã  dÃ©poser Ã  la DAL/GPR')
                                    subtitle_run.font.name = 'Tahoma'
                                    subtitle_run.font.size = Pt(11)
                                    subtitle_run.italic = True
                                    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    
                                    # Espace
                                    doc.add_paragraph()
                                    
                                    # NumÃ©ro de demande avec encadrement
                                    num_table = doc.add_table(rows=1, cols=1)
                                    num_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                                    num_cell = num_table.cell(0, 0)
                                    num_cell.width = Cm(6)
                                    num_para = num_cell.paragraphs[0]
                                    num_run = num_para.add_run('NÂ°')
                                    num_run.font.name = 'Tahoma'
                                    num_run.font.size = Pt(11)
                                    num_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    
                                    # Bordures pour le numÃ©ro
                                    def set_cell_border(cell, **kwargs):
                                        tc = cell._tc
                                        tcPr = tc.get_or_add_tcPr()
                                        tcBorders = tcPr.first_child_found_in("w:tcBorders")
                                        if tcBorders is None:
                                            tcBorders = OxmlElement('w:tcBorders')
                                            tcPr.append(tcBorders)
                                        
                                        for edge in ('top', 'left', 'bottom', 'right'):
                                            edge_data = kwargs.get(edge)
                                            if edge_data:
                                                tag = 'w:{}'.format(edge)
                                                element = tcBorders.find(qn(tag))
                                                if element is None:
                                                    element = OxmlElement(tag)
                                                    tcBorders.append(element)
                                                for key, value in edge_data.items():
                                                    element.set(qn('w:{}'.format(key)), str(value))
                                    
                                    border_kwargs = {
                                        'top': {'sz': 12, 'val': 'single', 'color': '000000'},
                                        'bottom': {'sz': 12, 'val': 'single', 'color': '000000'},
                                        'left': {'sz': 12, 'val': 'single', 'color': '000000'},
                                        'right': {'sz': 12, 'val': 'single', 'color': '000000'}
                                    }
                                    set_cell_border(num_cell, **border_kwargs)
                                    
                                    # Espace
                                    doc.add_paragraph()
                                    
                                    # Tableau principal des informations
                                    main_table = doc.add_table(rows=2, cols=2)
                                    main_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                                    
                                    # DÃ©finir les largeurs des colonnes
                                    main_table.columns[0].width = Cm(8)
                                    main_table.columns[1].width = Cm(8)
                                    
                                    # PremiÃ¨re ligne - Nom du demandeur
                                    cell_demandeur = main_table.cell(0, 0)
                                    cell_demandeur.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                                    para_demandeur = cell_demandeur.paragraphs[0]
                                    run_title = para_demandeur.add_run('Nom et qualitÃ© du demandeur\n')
                                    run_title.font.name = 'Tahoma'
                                    run_title.font.size = Pt(11)
                                    run_title.bold = True
                                    
                                    run_name = para_demandeur.add_run(f'{demandeur_nom}\n\n')
                                    run_name.font.name = 'Tahoma'
                                    run_name.font.size = Pt(11)
                                    
                                    run_details = para_demandeur.add_run(f'DIR. /DEP. : {demandeur_dir}\nNÂ° C.R.     : {demandeur_cr}\nNÂ° TÃ©l.     : {demandeur_tel}')
                                    run_details.font.name = 'Tahoma'
                                    run_details.font.size = Pt(11)
                                    
                                    # PremiÃ¨re ligne - Motif de la demande
                                    cell_motif = main_table.cell(0, 1)
                                    cell_motif.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                                    para_motif = cell_motif.paragraphs[0]
                                    run_motif_title = para_motif.add_run('Motif de la demande\n\n')
                                    run_motif_title.font.name = 'Tahoma'
                                    run_motif_title.font.size = Pt(11)
                                    run_motif_title.bold = True
                                    
                                    run_motif_content = para_motif.add_run(motif_demande)
                                    run_motif_content.font.name = 'Tahoma'
                                    run_motif_content.font.size = Pt(11)
                                    
                                    # DeuxiÃ¨me ligne - Dates
                                    cell_dates = main_table.cell(1, 0)
                                    cell_dates.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                                    para_dates = cell_dates.paragraphs[0]
                                    
                                    run_depart = para_dates.add_run(f'DÃ©part prÃ©vu : {date_depart.strftime("%d/%m/%Y")}\n\n')
                                    run_depart.font.name = 'Tahoma'
                                    run_depart.font.size = Pt(11)
                                    
                                    run_retour = para_dates.add_run(f'Retour prÃ©vu : {date_retour.strftime("%d/%m/%Y")}')
                                    run_retour.font.name = 'Tahoma'
                                    run_retour.font.size = Pt(11)
                                    
                                    # DeuxiÃ¨me ligne - Nombre de personnes et quantitÃ© de carburant
                                    cell_nb = main_table.cell(1, 1)
                                    cell_nb.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                                    para_nb = cell_nb.paragraphs[0]
                                    
                                    run_nb = para_nb.add_run(f'Nombre de personnes : {nb_personnes:02d}\n\n')
                                    run_nb.font.name = 'Tahoma'
                                    run_nb.font.size = Pt(11)
                                    
                                    # Nouveau champ pour la quantitÃ© de carburant
                                    run_carburant_title = para_nb.add_run('QuantitÃ© de carburant demandÃ©e :\n\n')
                                    run_carburant_title.font.name = 'Tahoma'
                                    run_carburant_title.font.size = Pt(11)
                                    run_carburant_title.bold = True
                                    
                                    run_carburant_value = para_nb.add_run(f'{quantite_carburant} litres')
                                    run_carburant_value.font.name = 'Tahoma'
                                    run_carburant_value.font.size = Pt(11)
                                    
                                    # Appliquer les bordures au tableau principal
                                    for row in main_table.rows:
                                        for cell in row.cells:
                                            set_cell_border(cell, **border_kwargs)
                                    
                                    # Espace
                                    doc.add_paragraph()
                                    
                                    # Tableau itinÃ©raire
                                    itinerary_table = doc.add_table(rows=1, cols=2)
                                    itinerary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                                    itinerary_table.columns[0].width = Cm(12)
                                    itinerary_table.columns[1].width = Cm(4)
                                    
                                    # En-tÃªtes du tableau itinÃ©raire
                                    hdr_cells = itinerary_table.rows[0].cells
                                    
                                    hdr_para1 = hdr_cells[0].paragraphs[0]
                                    hdr_run1 = hdr_para1.add_run('ItinÃ©raire Ã  suivre')
                                    hdr_run1.font.name = 'Tahoma'
                                    hdr_run1.font.size = Pt(11)
                                    hdr_run1.bold = True
                                    hdr_para1.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    
                                    hdr_para2 = hdr_cells[1].paragraphs[0]
                                    hdr_run2 = hdr_para2.add_run('KM')
                                    hdr_run2.font.name = 'Tahoma'
                                    hdr_run2.font.size = Pt(11)
                                    hdr_run2.bold = True
                                    hdr_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    
                                    # Ajouter les sites de la mission ou lignes vides
                                    if st.session_state.planning_results:
                                        sites = st.session_state.planning_results.get('sites_ordered', [])
                                        segments = st.session_state.planning_results.get('segments_summary', [])
                                        base_location = st.session_state.planning_results.get('base_location', '')
                                        
                                        # Commencer Ã  partir du deuxiÃ¨me site pour Ã©viter d'afficher juste "Dakar"
                                        for i in range(1, len(sites)):
                                            row_cells = itinerary_table.add_row().cells
                                            
                                            para_site = row_cells[0].paragraphs[0]
                                            prev_site = sites[i-1]
                                            current_site = sites[i]
                                            site_text = f"{prev_site['Ville']} â†’ {current_site['Ville']}"
                                            
                                            run_site = para_site.add_run(site_text)
                                            run_site.font.name = 'Tahoma'
                                            run_site.font.size = Pt(11)
                                            
                                            para_km = row_cells[1].paragraphs[0]
                                            # Utiliser l'index i-1 pour les segments car on commence Ã  i=1
                                            if (i-1) < len(segments):
                                                distance_km = segments[i-1]['distance'] / 1000
                                                km_text = f"{distance_km:.1f}"
                                            else:
                                                km_text = "___"
                                            
                                            run_km = para_km.add_run(km_text)
                                            run_km.font.name = 'Tahoma'
                                            run_km.font.size = Pt(11)
                                            para_km.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                        
                                        # Distance totale
                                        total_row = itinerary_table.add_row().cells
                                        para_total = total_row[0].paragraphs[0]
                                        run_total = para_total.add_run('Distance totale :')
                                        run_total.font.name = 'Tahoma'
                                        run_total.font.size = Pt(11)
                                        run_total.bold = True
                                        
                                        para_total_km = total_row[1].paragraphs[0]
                                        run_total_km = para_total_km.add_run(f"{stats.get('total_km', 0):.1f}")
                                        run_total_km.font.name = 'Tahoma'
                                        run_total_km.font.size = Pt(11)
                                        run_total_km.bold = True
                                        para_total_km.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    else:
                                        # Lignes vides si pas de planning
                                        for _ in range(8):
                                            row_cells = itinerary_table.add_row().cells
                                            para_empty = row_cells[0].paragraphs[0]
                                            run_empty = para_empty.add_run("")
                                            run_empty.font.name = 'Tahoma'
                                            run_empty.font.size = Pt(11)
                                            
                                            para_empty_km = row_cells[1].paragraphs[0]
                                            run_empty_km = para_empty_km.add_run("")
                                            run_empty_km.font.name = 'Tahoma'
                                            run_empty_km.font.size = Pt(11)
                                        
                                        # Distance totale vide
                                        total_row = itinerary_table.add_row().cells
                                        para_total = total_row[0].paragraphs[0]
                                        run_total = para_total.add_run('Distance totale :')
                                        run_total.font.name = 'Tahoma'
                                        run_total.font.size = Pt(11)
                                        run_total.bold = True
                                        
                                        para_total_km = total_row[1].paragraphs[0]
                                        run_total_km = para_total_km.add_run("")
                                        run_total_km.font.name = 'Tahoma'
                                        run_total_km.font.size = Pt(11)
                                    
                                    # Appliquer les bordures au tableau itinÃ©raire
                                    for row in itinerary_table.rows:
                                        for cell in row.cells:
                                            set_cell_border(cell, **border_kwargs)
                                    
                                    # Espace rÃ©duit
                                    doc.add_paragraph()
                                    
                                    # Date
                                    date_para = doc.add_paragraph()
                                    date_run = date_para.add_run(f'Date : Le {datetime.now().strftime("%d/%m/%Y")}')
                                    date_run.font.name = 'Tahoma'
                                    date_run.font.size = Pt(11)
                                    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                                    
                                    # Tableau des signatures
                                    sig_table = doc.add_table(rows=1, cols=3)
                                    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                                    
                                    # DÃ©finir les largeurs des colonnes de signature
                                    sig_table.columns[0].width = Cm(5.3)
                                    sig_table.columns[1].width = Cm(5.3)
                                    sig_table.columns[2].width = Cm(5.3)
                                    
                                    # Contenu des cellules de signature
                                    sig_cells = sig_table.rows[0].cells
                                    
                                    # PremiÃ¨re signature
                                    para_sig1 = sig_cells[0].paragraphs[0]
                                    run_sig1_title = para_sig1.add_run('Signature et cachet\n')
                                    run_sig1_title.font.name = 'Tahoma'
                                    run_sig1_title.font.size = Pt(11)
                                    run_sig1_title.bold = True
                                    para_sig1.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    
                                    run_sig1_subtitle = para_sig1.add_run('Chef de service Demandeur')
                                    run_sig1_subtitle.font.name = 'Tahoma'
                                    run_sig1_subtitle.font.size = Pt(11)
                                    
                                    # DeuxiÃ¨me signature
                                    para_sig2 = sig_cells[1].paragraphs[0]
                                    run_sig2 = para_sig2.add_run('Responsable POOL')
                                    run_sig2.font.name = 'Tahoma'
                                    run_sig2.font.size = Pt(11)
                                    run_sig2.bold = True
                                    para_sig2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    
                                    # TroisiÃ¨me signature
                                    para_sig3 = sig_cells[2].paragraphs[0]
                                    run_sig3 = para_sig3.add_run('DAL/GPR')
                                    run_sig3.font.name = 'Tahoma'
                                    run_sig3.font.size = Pt(11)
                                    run_sig3.bold = True
                                    para_sig3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    
                                    # DÃ©finir la hauteur des cellules de signature
                                    for cell in sig_cells:
                                        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                                        # Ajouter de l'espace pour les signatures
                                        for _ in range(4):
                                            cell.add_paragraph()
                                    
                                    # Appliquer les bordures au tableau de signatures
                                    for row in sig_table.rows:
                                        for cell in row.cells:
                                            set_cell_border(cell, **border_kwargs)
                                    
                                    # Sauvegarder dans un buffer
                                    buffer = io.BytesIO()
                                    doc.save(buffer)
                                    buffer.seek(0)
                                    
                                    # Bouton de tÃ©lÃ©chargement
                                    st.success("âœ… Document gÃ©nÃ©rÃ© avec succÃ¨s!")
                                    st.download_button(
                                        label="ðŸ“¥ TÃ©lÃ©charger la demande de carburant (Word)",
                                        data=buffer.getvalue(),
                                        file_name=f"Demande_carburant_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                                    )
                                    
                                except ImportError:
                                    st.error("âŒ Le module python-docx n'est pas installÃ©. Veuillez l'installer avec: pip install python-docx")
                                except Exception as e:
                                    st.error(f"âŒ Erreur lors de la gÃ©nÃ©ration du document: {str(e)}")
                        
                        with col_action2:
                            if st.button("âŒ Annuler", use_container_width=True):
                                st.session_state.show_fuel_request_modal = False
                                st.rerun()


                
            else:
                st.error("âŒ Aucune distance calculÃ©e. VÃ©rifiez votre planning.")
    
    with tab_edit:
        st.subheader("âœï¸ Ã‰dition manuelle du planning")
        
        st.info("ðŸ’¡ Modifiez les horaires, ajoutez ou supprimez des Ã©vÃ©nements. Les modifications sont automatiquement sauvegardÃ©es.")
        
        # Initialiser manual_itinerary si nÃ©cessaire
        if st.session_state.manual_itinerary is None:
            st.session_state.manual_itinerary = list(itinerary)
        
        # SÃ©lection du jour
        total_days = max(ev[0] for ev in st.session_state.manual_itinerary) if st.session_state.manual_itinerary else 1
        
        selected_edit_day = st.selectbox(
            "SÃ©lectionnez le jour Ã  Ã©diter",
            options=range(1, total_days + 1),
            format_func=lambda x: f"Jour {x} - {(start_date + timedelta(days=x-1)).strftime('%d/%m/%Y')}",
            key="edit_day_select"
        )
        
        # Filtrer les Ã©vÃ©nements du jour
        day_events_edit = [(i, ev) for i, ev in enumerate(st.session_state.manual_itinerary) if ev[0] == selected_edit_day]
        
        st.markdown("---")
        
        # Afficher chaque Ã©vÃ©nement avec possibilitÃ© d'Ã©dition
        for idx, (global_idx, (day, sdt, edt, desc)) in enumerate(day_events_edit):
            with st.expander(f"**Ã‰vÃ©nement {idx+1}** : {desc[:50]}...", expanded=False):
                col1, col2 = st.columns(2)
                
                with col1:
                    new_start = st.time_input(
                        "Heure de dÃ©but",
                        value=sdt.time(),
                        key=f"start_{global_idx}"
                    )
                
                with col2:
                    new_end = st.time_input(
                        "Heure de fin",
                        value=edt.time(),
                        key=f"end_{global_idx}"
                    )
                
                new_desc = st.text_area(
                    "Description",
                    value=desc,
                    height=100,
                    key=f"desc_{global_idx}"
                )

                # Saisie manuelle des distances pour les trajets
                override_km = None
                override_h = None
                override_m = None
                if "â†’" in desc:
                    import re
                    km_val = 0.0
                    h_val = 0
                    m_val = 0
                    # Essayer d'extraire (123.4 km, 2h30) ou (123.4 km)
                    m_with_time = re.search(r"\((['\d\.']+)\s*km,\s*([^)]+)\)", desc)
                    if m_with_time:
                        try:
                            km_val = float(m_with_time.group(1))
                            time_str = m_with_time.group(2).strip()
                            m_time = re.match(r"(\d+)\s*h\s*(\d{1,2})", time_str)
                            if m_time:
                                h_val = int(m_time.group(1))
                                m_val = int(m_time.group(2))
                        except Exception:
                            pass
                    else:
                        m_km = re.search(r"\(([\d\.]+)\s*km\)", desc)
                        if m_km:
                            try:
                                km_val = float(m_km.group(1))
                            except Exception:
                                pass
                    st.markdown("**Distance du trajet (si diffÃ©rente)**")
                    col_d1, col_d2, col_d3 = st.columns([1, 1, 1])
                    with col_d1:
                        override_km = st.number_input("Distance (km)", min_value=0.0, value=float(km_val), key=f"km_{global_idx}")
                    with col_d2:
                        override_h = st.number_input("Heures", min_value=0, value=int(h_val), key=f"kh_{global_idx}")
                    with col_d3:
                        override_m = st.number_input("Minutes", min_value=0, max_value=59, value=int(m_val), key=f"kmn_{global_idx}")
                
                col_btn1, col_btn2, col_btn3 = st.columns(3)
                
                with col_btn1:
                    if st.button("ðŸ’¾ Sauvegarder", key=f"save_{global_idx}", use_container_width=True):
                        new_sdt = datetime.combine(sdt.date(), new_start)
                        new_edt = datetime.combine(edt.date(), new_end)
                        # Appliquer Ã©ventuelles distances sur la description (pour les trajets)
                        updated_desc = new_desc
                        if override_km is not None:
                            import re
                            # Retirer ancien motif distance s'il existe
                            updated_desc = re.sub(r"\s*\([\d\.\s]+km(?:,\s*[^)]*)?\)\s*$", "", updated_desc).strip()
                            if override_km > 0:
                                if (override_h or 0) > 0 or (override_m or 0) > 0:
                                    updated_desc = f"{updated_desc} ({override_km} km, {int(override_h or 0)}h{int(override_m or 0):02d})"
                                else:
                                    updated_desc = f"{updated_desc} ({override_km} km)"
                        st.session_state.manual_itinerary[global_idx] = (day, new_sdt, new_edt, updated_desc)
                        st.success("Modifications sauvegardÃ©es!")
                        st.rerun()
                
                with col_btn2:
                    if st.button("ðŸ—‘ï¸ Supprimer", key=f"delete_{global_idx}", use_container_width=True):
                        st.session_state.manual_itinerary.pop(global_idx)
                        st.success("Ã‰vÃ©nement supprimÃ©!")
                        st.rerun()
                
                with col_btn3:
                    if st.button("â†•ï¸ DÃ©placer", key=f"move_{global_idx}", use_container_width=True):
                        st.session_state.editing_event = global_idx
        
        # Ajouter un nouvel Ã©vÃ©nement
        st.markdown("---")
        st.subheader("âž• Ajouter un Ã©vÃ©nement")
        
        with st.form("add_event_form"):
            col1, col2, col3 = st.columns(3)
            
            with col1:
                new_event_start = st.time_input("DÃ©but", value=time(8, 0))
            
            with col2:
                new_event_end = st.time_input("Fin", value=time(9, 0))
            
            with col3:
                event_type = st.selectbox(
                    "Type",
                    ["Visite", "Trajet", "Pause", "Autre"]
                )
            
            new_event_desc = st.text_input("Description", value="Nouvel Ã©vÃ©nement")
            
            if st.form_submit_button("Ajouter l'Ã©vÃ©nement"):
                event_date = start_date + timedelta(days=selected_edit_day-1)
                new_sdt = datetime.combine(event_date, new_event_start)
                new_edt = datetime.combine(event_date, new_event_end)
                
                prefix = ""
                if event_type == "Trajet":
                    prefix = "ðŸš— "
                elif event_type == "Pause":
                    prefix = "â¸ï¸ "
                elif event_type == "Visite":
                    prefix = ""
                
                new_event = (selected_edit_day, new_sdt, new_edt, f"{prefix}{new_event_desc}")
                st.session_state.manual_itinerary.append(new_event)
                st.session_state.manual_itinerary.sort(key=lambda x: (x[0], x[1]))
                st.success("Ã‰vÃ©nement ajoutÃ©!")
                st.rerun()
        
        # Boutons d'action globaux
        st.markdown("---")
        col_reset, col_recalc = st.columns(2)
        
        with col_reset:
            if st.button("ðŸ”„ RÃ©initialiser les modifications", use_container_width=True):
                st.session_state.manual_itinerary = None
                st.success("Planning rÃ©initialisÃ©!")
                st.rerun()
        
        with col_recalc:
            if st.button("ðŸ”¢ Recalculer les statistiques", use_container_width=True):
                # Recalculer les stats basÃ©es sur manual_itinerary
                total_km = 0
                total_visit_hours = 0
                
                for day, sdt, edt, desc in st.session_state.manual_itinerary:
                    import re
                    # Compter les kilomÃ¨tres UNIQUEMENT pour les trajets (flÃ¨che ou emoji voiture)
                    if ("â†’" in desc) or ("ðŸš—" in desc):
                        # Nouveau format avec temps rÃ©el: "(... km, HhMM)"
                        m_with_time = re.search(r"\(([\d\.]+)\s*km,\s*([^)]+)\)", desc)
                        if m_with_time:
                            try:
                                total_km += float(m_with_time.group(1))
                            except Exception:
                                pass
                        else:
                            # Ancien format: "(... km)"
                            m = re.search(r"\(([\d\.]+)\s*km\)", desc)
                            if m:
                                try:
                                    total_km += float(m.group(1))
                                except Exception:
                                    pass
                    
                    # Cumuler les heures de visite (Ã©viter de compter les trajets)
                    if any(x in desc for x in ["Visite", "â€“"]) and "â†’" not in desc:
                        duration = (edt - sdt).total_seconds() / 3600
                        total_visit_hours += duration
                
                stats['total_km'] = total_km
                stats['total_visit_hours'] = total_visit_hours
                
                st.success("Statistiques recalculÃ©es!")
                st.rerun()
    
    with tab_manual:
        st.subheader("ðŸ”„ Modification manuelle de l'ordre des sites")
        
        st.info("ðŸ’¡ RÃ©organisez l'ordre des sites en les faisant glisser. L'itinÃ©raire sera automatiquement recalculÃ©.")
        
        # VÃ©rifier que nous avons les donnÃ©es nÃ©cessaires
        if 'original_order' not in results or 'durations_matrix' not in results:
            st.warning("âš ï¸ DonnÃ©es insuffisantes pour la modification manuelle. Veuillez relancer le calcul.")
        else:
            # RÃ©cupÃ©rer les donnÃ©es
            original_order = results['original_order']
            durations_matrix = results['durations_matrix']
            distances_matrix = results['distances_matrix']
            all_coords = results['all_coords']
            
            # Ã‰diteur de table (optionnel)
            with st.expander("ðŸ§® Mode tableau (numÃ©ro d'ordre)", expanded=False):
                import pandas as pd
                rows = []
                for i, site_idx in enumerate(st.session_state.get('manual_order', original_order)):
                    if isinstance(site_idx, int) and 0 <= site_idx < len(sites_ordered):
                        s = sites_ordered[site_idx]
                        rows.append({
                            "Index": site_idx,
                            "Ville": s['Ville'],
                            "Type": s.get('Type', 'Site'),
                            "ActivitÃ©": s.get('ActivitÃ©', 'ActivitÃ©'),
                            "Ordre": i+1
                        })
                df_order = pd.DataFrame(rows)
                edited_df = st.data_editor(
                    df_order,
                    column_config={
                        "Ordre": st.column_config.NumberColumn("Ordre", min_value=1, max_value=len(rows), step=1),
                        "Index": st.column_config.TextColumn("Index", disabled=True),
                        "Ville": st.column_config.TextColumn("Ville", disabled=True),
                        "Type": st.column_config.TextColumn("Type", disabled=True),
                        "ActivitÃ©": st.column_config.TextColumn("ActivitÃ©", disabled=True),
                    },
                    column_order=["Ville", "Type", "ActivitÃ©", "Ordre"],
                    use_container_width=True,
                    num_rows="fixed"
                )
                colA, colB = st.columns([2,1])
                with colA:
                    if st.button("âœ… Appliquer l'ordre (table)", type="primary", use_container_width=True):
                        try:
                            ords = edited_df["Ordre"].tolist()
                            if sorted(ords) != list(range(1, len(rows)+1)):
                                st.error("Veuillez attribuer des numÃ©ros d'ordre uniques de 1 Ã  N.")
                            else:
                                new_order = [int(row["Index"]) for _, row in edited_df.sort_values("Ordre").iterrows()]
                                st.session_state.manual_order = new_order
                                st.success("Ordre mis Ã  jour depuis la table!")
                                st.rerun()
                        except Exception as e:
                            st.error(f"Erreur lors de l'application: {str(e)[:120]}...")
                with colB:
                    if st.button("â†©ï¸ Restaurer", use_container_width=True):
                        st.session_state.manual_order = original_order.copy()
                        st.success("Ordre original restaurÃ©!")
                        st.rerun()
            
            # CrÃ©er une liste des sites avec leur ordre actuel
            if 'manual_order' not in st.session_state:
                st.session_state.manual_order = original_order.copy()
            
            # Afficher l'ordre actuel des sites
            st.markdown("**Ordre actuel des sites :**")
            st.info(f"ðŸ“Š **{len(st.session_state.manual_order)} sites** dans l'ordre actuel")
            
            # Interface pour rÃ©organiser les sites avec conteneur scrollable
            with st.container():
                # Utiliser des boutons pour dÃ©placer les sites
                for i, site_idx in enumerate(st.session_state.manual_order):
                    # VÃ©rifier que l'index est valide et de type entier
                    if isinstance(site_idx, int) and 0 <= site_idx < len(sites_ordered):
                        site = sites_ordered[site_idx]
                        
                        col1, col2, col3, col4 = st.columns([3, 1, 1, 1])
                        
                        with col1:
                            st.write(f"**{i+1}.** {site['Ville']} - {site.get('Type', 'Site')} - {site.get('ActivitÃ©', 'ActivitÃ©')}")
                        
                        with col2:
                            if i > 0 and st.button("â¬†ï¸", key=f"enhanced_up_{i}", help="Monter"):
                                # Ã‰changer avec l'Ã©lÃ©ment prÃ©cÃ©dent
                                st.session_state.manual_order[i], st.session_state.manual_order[i-1] = \
                                    st.session_state.manual_order[i-1], st.session_state.manual_order[i]
                                st.rerun()
                        
                        with col3:
                            if i < len(st.session_state.manual_order) - 1 and st.button("â¬‡ï¸", key=f"enhanced_down_{i}", help="Descendre"):
                                # Ã‰changer avec l'Ã©lÃ©ment suivant
                                st.session_state.manual_order[i], st.session_state.manual_order[i+1] = \
                                    st.session_state.manual_order[i+1], st.session_state.manual_order[i]
                                st.rerun()
                        
                        with col4:
                            if i != 0 and i != len(st.session_state.manual_order) - 1:  # Ne pas permettre de supprimer le dÃ©part et l'arrivÃ©e
                                if st.button("ðŸ—‘ï¸", key=f"remove_{i}", help="Supprimer"):
                                    st.session_state.manual_order.pop(i)
                                    st.rerun()
                    else:
                        # Index invalide - nettoyer
                        st.warning(f"âš ï¸ Index invalide dÃ©tectÃ© ({site_idx}), nettoyage en cours...")
                        st.session_state.manual_order = [idx for idx in st.session_state.manual_order if idx < len(sites_ordered)]
                        st.rerun()
            
            st.markdown("---")
            
            # Boutons d'action
            col1, col2, col3, col4 = st.columns(4)
            
            with col1:
                if st.button("ðŸ”„ Recalculer l'itinÃ©raire", use_container_width=True):
                    # Recalculer l'itinÃ©raire avec le nouvel ordre
                    new_order = st.session_state.manual_order
                    
                    # Recalculer les segments
                    new_segments = []
                    for i in range(len(new_order)-1):
                        from_idx = new_order[i]
                        to_idx = new_order[i+1]
                        
                        if from_idx < len(durations_matrix) and to_idx < len(durations_matrix[0]):
                            duration = durations_matrix[from_idx][to_idx]
                            distance = distances_matrix[from_idx][to_idx] if distances_matrix else 0
                            
                            new_segments.append({
                                "distance": distance,
                                "duration": duration
                            })
                        else:
                            new_segments.append({"distance": 0, "duration": 0})
                    
                    # Recalculer l'itinÃ©raire complet
                    new_sites = [sites_ordered[i] for i in new_order]
                    new_coords = [coords_ordered[i] for i in new_order]
                    # Calcul du nombre de jours optimal (dry-run)
                    _, _, _, new_dry_stats = schedule_itinerary(
                        coords=new_coords,
                        sites=new_sites,
                        order=list(range(len(new_order))),
                        segments_summary=new_segments,
                        start_date=start_date,
                        start_activity_time=time(8, 0),
                        end_activity_time=time(17, 0),
                        start_travel_time=time(7, 0),
                        end_travel_time=time(19, 0),
                        use_lunch=True,
                        lunch_start_time=time(12, 30),
                        lunch_end_time=time(14, 0),
                        use_prayer=False,
                        prayer_start_time=time(14, 0),
                        prayer_duration_min=20,
                        lunch_duration_min=st.session_state.get('lunch_duration_min', 60),
                        max_days=0,
                        tolerance_hours=1.0,
                        base_location=results.get('base_location', ''),
                        allow_weekend_travel=st.session_state.get('allow_weekend_travel', True),
                        allow_weekend_activities=st.session_state.get('allow_weekend_activities', True)
                    )

                    new_optimal_days = int(new_dry_stats.get('total_days', 1))
                    user_max_recalc = int(st.session_state.get('max_days', 0))
                    user_desired_recalc = int(st.session_state.get('desired_days', 0))

                    # Logique de dÃ©cision pour les jours effectifs lors du recalcul
                    if user_desired_recalc > 0:
                        if user_max_recalc > 0 and user_desired_recalc > user_max_recalc:
                            st.warning(f"Souhait ({user_desired_recalc} jours) > max ({user_max_recalc}). Utilisation du max.")
                            new_effective_days = user_max_recalc
                        else:
                            new_effective_days = user_desired_recalc
                        
                        if new_effective_days < new_optimal_days:
                            new_stretch = True
                            st.warning(f"âš ï¸ Objectif ({new_effective_days}) < optimal ({new_optimal_days}). Compression.")
                        else:
                            new_stretch = False
                            st.success(f"âœ… AjustÃ© Ã  {new_effective_days} jours.")

                    elif user_max_recalc > 0:
                        if user_max_recalc < new_optimal_days:
                            new_effective_days = user_max_recalc
                            new_stretch = True
                            st.warning(f"âš ï¸ Objectif ({user_max_recalc}) < optimal ({new_optimal_days}). Compression.")
                        else:
                            new_effective_days = user_max_recalc
                            new_stretch = False
                            st.success(f"âœ… Tient en {new_optimal_days} jours (max: {user_max_recalc}).")
                    else:
                        new_effective_days = new_optimal_days
                        new_stretch = False
                        st.info(f"ðŸ§® Jours optimaux recalculÃ©s: {new_optimal_days}.")

                    # Planification finale recalculÃ©e
                    new_itinerary, new_sites_ordered, new_coords_ordered, new_stats = schedule_itinerary(
                        coords=new_coords,
                        sites=new_sites,
                        order=list(range(len(new_order))),  # Ordre sÃ©quentiel car sites dÃ©jÃ  rÃ©organisÃ©s
                        segments_summary=new_segments,
                        start_date=start_date,
                        start_activity_time=time(8, 0),  # Utiliser les valeurs par dÃ©faut ou rÃ©cupÃ©rer depuis session_state
                        end_activity_time=time(17, 0),
                        start_travel_time=time(7, 0),
                        end_travel_time=time(19, 0),
                        use_lunch=True,
                        lunch_start_time=time(12, 30),
                        lunch_end_time=time(14, 0),
                        use_prayer=False,
                        prayer_start_time=time(14, 0),
                        prayer_duration_min=20,
                        lunch_duration_min=st.session_state.get('lunch_duration_min', 60),
                        max_days=new_effective_days,
                        tolerance_hours=1.0,
                        base_location=results.get('base_location', ''),
                        stretch_days=new_stretch,
                        allow_weekend_travel=st.session_state.get('allow_weekend_travel', True),
                        allow_weekend_activities=st.session_state.get('allow_weekend_activities', True)
                    )

                    if new_stretch and new_stats.get('total_days', 0) > new_effective_days:
                        st.error(f"âŒ Impossible de tenir en {new_effective_days} jour(s). Besoin de {new_stats.get('total_days')} jours mÃªme en Ã©tirant les journÃ©es.")
                    
                    # Mettre Ã  jour les rÃ©sultats
                    st.session_state.manual_itinerary = new_itinerary
                    st.session_state.planning_results.update({
                        'sites_ordered': new_sites_ordered,
                        'coords_ordered': new_coords_ordered,
                        'stats': new_stats,
                        'segments_summary': new_segments
                    })
                    
                    st.success("âœ… ItinÃ©raire recalculÃ© avec le nouvel ordre!")
                    st.rerun()
            
            with col2:
                if st.button("â†©ï¸ Restaurer l'ordre original", use_container_width=True):
                    st.session_state.manual_order = original_order.copy()
                    st.session_state.manual_itinerary = None
                    st.success("Ordre original restaurÃ©!")
                    st.rerun()
            
            with col3:
                if st.button("ðŸŽ¯ Optimiser automatiquement", use_container_width=True):
                    # RÃ©optimiser avec IA Adja
                    try:
                        ai_order, ai_success, ai_message = optimize_route_with_ai(sites_ordered, coords_ordered, base_location, deepseek_api_key)
                        if ai_success and isinstance(ai_order, list):
                            st.session_state.manual_order = ai_order
                            st.success(f"Ordre optimisÃ© automatiquement par IA Adja! {ai_message}")
                        else:
                            # Fallback vers TSP si l'IA Adja Ã©choue ou rÃ©ponse invalide
                            optimized_order = solve_tsp_fixed_start_end(durations_matrix)
                            st.session_state.manual_order = optimized_order
                            st.warning(f"IA Adja indisponible ou rÃ©ponse invalide, optimisation TSP utilisÃ©e. {ai_message if not ai_success else ''}")
                    except Exception as e:
                        # Fallback vers TSP en cas d'erreur
                        optimized_order = solve_tsp_fixed_start_end(durations_matrix)
                        st.session_state.manual_order = optimized_order
                        st.warning(f"Erreur IA Adja ({str(e)[:50]}...), optimisation TSP utilisÃ©e.")
                    st.rerun()

            with col4:
                if st.button("âš™ï¸ Optimiser (OR-Tools)", use_container_width=True):
                    try:
                        # Construire les durÃ©es de service Ã  partir des donnÃ©es des sites (DurÃ©e (h))
                        service_times_sec = []
                        for i in range(len(sites_ordered)):
                            dur_h = sites_ordered[i].get("DurÃ©e (h)", 0)
                            try:
                                service_times_sec.append(int(float(dur_h or 0) * 3600))
                            except Exception:
                                service_times_sec.append(0)

                        optimized_path = solve_tsp_ortools_fixed_start_end(durations_matrix, service_times_sec, time_limit_s=5)
                        st.session_state.manual_order = optimized_path
                        st.success("Ordre optimisÃ© par OR-Tools (matrice OSRM/Maps).")
                    except Exception as e:
                        st.warning(f"OR-Tools indisponible ou erreur: {str(e)[:80]}... Fallback TSP.")
                        st.session_state.manual_order = solve_tsp_fixed_start_end(durations_matrix)
                    st.rerun()
    
    with tab_map:
        st.subheader("Carte de l'itinÃ©raire")
        
        if coords_ordered:
            center_lat = sum(c[1] for c in coords_ordered) / len(coords_ordered)
            center_lon = sum(c[0] for c in coords_ordered) / len(coords_ordered)
            
            m = folium.Map(location=[center_lat, center_lon], zoom_start=7)
            
            # TracÃ© de l'itinÃ©raire : tentative de rÃ©cupÃ©ration de la route rÃ©elle via OSRM, sinon fallback sur la ligne droite
            try:
                coord_str = ";".join([f"{c[0]},{c[1]}" for c in coords_ordered])
                url = f"{osrm_base_url.rstrip('/')}/route/v1/driving/{coord_str}?overview=full&geometries=geojson"
                resp = requests.get(url, timeout=10)
                route_pts = None
                if resp.status_code == 200:
                    data = resp.json()
                    if data.get('routes'):
                        geom = data['routes'][0].get('geometry')
                        if isinstance(geom, dict) and geom.get('coordinates'):
                            route_pts = [[lat, lon] for lon, lat in geom['coordinates']]
                if not route_pts:
                    route_pts = [[c[1], c[0]] for c in coords_ordered]
            except Exception:
                route_pts = [[c[1], c[0]] for c in coords_ordered]
            folium.PolyLine(locations=route_pts, color="blue", weight=3, opacity=0.7).add_to(m)
            
            # Export Google Maps (ouvrir et copier)
            try:
                gmaps_pairs = [f"{c[1]},{c[0]}" for c in coords_ordered]
                gmaps_url = "https://www.google.com/maps/dir/" + "/".join(gmaps_pairs) + "/?hl=fr"
                st.markdown(f"[ðŸ“ Ouvrir dans Google Maps]({gmaps_url})")
                st.text_input("Lien Google Maps", value=gmaps_url, help="Copiez ce lien pour partager ou ouvrir l'itinÃ©raire dans Google Maps.", label_visibility="collapsed")
            except Exception:
                st.info("Lien Google Maps non disponible")
            
            # PrÃ©parer affichage spÃ©cial si dÃ©part et arrivÃ©e sont au mÃªme endroit
            n_steps = len(sites_ordered)
            start_end_same = False
            if n_steps >= 2:
                lat0, lon0 = coords_ordered[0][1], coords_ordered[0][0]
                latN, lonN = coords_ordered[-1][1], coords_ordered[-1][0]
                start_end_same = abs(lat0 - latN) < 1e-4 and abs(lon0 - lonN) < 1e-4

            for i, site in enumerate(sites_ordered):
                # Si le dÃ©part et l'arrivÃ©e sont identiques, afficher un double numÃ©ro sur le point de dÃ©part et ne pas dupliquer le dernier point
                if i == 0 and start_end_same:
                    bg_color_left = '#2ecc71'  # Vert pour dÃ©part
                    bg_color_right = '#e74c3c'  # Rouge pour arrivÃ©e
                    html = f"""
<div style=\"display:flex; align-items:center; gap:4px;\">
  <div style=\"background-color:{bg_color_left}; color:white; border-radius:50%; width:28px; height:28px; text-align:center; font-size:14px; font-weight:bold; line-height:28px; border:2px solid white; box-shadow:0 0 3px rgba(0,0,0,0.5);\">1</div>
  <div style=\"background-color:{bg_color_right}; color:white; border-radius:50%; width:28px; height:28px; text-align:center; font-size:14px; font-weight:bold; line-height:28px; border:2px solid white; box-shadow:0 0 3px rgba(0,0,0,0.5);\">{n_steps}</div>
</div>
"""
                    folium.Marker(
                        location=[coords_ordered[i][1], coords_ordered[i][0]],
                        popup=f"Ã‰tapes 1 et {n_steps}: {site['Ville']}<br>{site.get('Type', '-')}",
                        tooltip=f"Ã‰tapes 1 et {n_steps}: {site['Ville']}",
                        icon=folium.DivIcon(
                            icon_size=(36, 28),
                            icon_anchor=(18, 14),
                            html=html
                        )
                    ).add_to(m)
                    continue
                if start_end_same and i == n_steps - 1:
                    # Ne pas dupliquer l'arrivÃ©e si elle est au mÃªme endroit que le dÃ©part
                    continue

                color = 'green' if i == 0 else 'red' if i == len(sites_ordered)-1 else 'blue'
                icon = 'play' if i == 0 else 'stop' if i == len(sites_ordered)-1 else 'info-sign'
                
                # IcÃ´ne numÃ©rotÃ©e via DivIcon (couleur selon Ã©tape)
                bg_color = '#2ecc71' if i == 0 else '#e74c3c' if i == len(sites_ordered)-1 else '#3498db'
                folium.Marker(
                    location=[coords_ordered[i][1], coords_ordered[i][0]],
                    popup=f"Ã‰tape {i+1}: {site['Ville']}<br>{site.get('Type', '-')}",
                    tooltip=f"Ã‰tape {i+1}: {site['Ville']}",
                    icon=folium.DivIcon(
                        icon_size=(28, 28),
                        icon_anchor=(14, 14),
                        html=f"""
<div style=\"background-color:{bg_color}; color:white; border-radius:50%; width:28px; height:28px; text-align:center; font-size:14px; font-weight:bold; line-height:28px; border:2px solid white; box-shadow:0 0 3px rgba(0,0,0,0.5);\">{i+1}</div>
"""
                    )
                ).add_to(m)
            
            st_folium(m, width=None, height=500, use_container_width=True)
            
            # TÃ©lÃ©chargements KML/KMZ de l'itinÃ©raire (points + trace)
            try:
                import zipfile
                from io import BytesIO
                
                # Construire le KML: Placemarks pour chaque Ã©tape + LineString pour la route
                doc_name = mission_title if mission_title else "ItinÃ©raire"
                kml_parts = [
                    "<?xml version=\"1.0\" encoding=\"UTF-8\"?>",
                    "<kml xmlns=\"http://www.opengis.net/kml/2.2\">",
                    "  <Document>",
                    f"    <name>{doc_name}</name>"
                ]
                
                # Placemarks pour les points (Ã©viter la duplication si dÃ©part=arrivÃ©e)
                for i, site in enumerate(sites_ordered):
                    if start_end_same and i == n_steps - 1:
                        continue
                    lon, lat = coords_ordered[i][0], coords_ordered[i][1]
                    site_name = site.get('Ville', '')
                    site_type = site.get('Type', '-')
                    placemark = f"""
    <Placemark>
      <name>Ã‰tape {i+1}: {site_name}</name>
      <description>{site_type}</description>
      <Point>
        <coordinates>{lon},{lat},0</coordinates>
      </Point>
    </Placemark>"""
                    kml_parts.append(placemark)
                
                # LineString pour la trace (OSRM si disponible, sinon ligne droite)
                line_coords = "\n".join([f"{pt[1]},{pt[0]},0" for pt in route_pts])
                linestring = f"""
    <Placemark>
      <name>Route</name>
      <Style>
        <LineStyle>
          <color>ff0000ff</color>
          <width>3</width>
        </LineStyle>
      </Style>
      <LineString>
        <tessellate>1</tessellate>
        <coordinates>
{line_coords}
        </coordinates>
      </LineString>
    </Placemark>"""
                kml_parts.append(linestring)
                kml_parts.append("  </Document>")
                kml_parts.append("</kml>")
                kml_content = "\n".join(kml_parts)
                
                # PrÃ©parer tÃ©lÃ©chargement KML
                kml_bytes = kml_content.encode('utf-8')
                
                # PrÃ©parer tÃ©lÃ©chargement KMZ (doc.kml dans une archive ZIP)
                kmz_buffer = BytesIO()
                with zipfile.ZipFile(kmz_buffer, 'w', compression=zipfile.ZIP_DEFLATED) as zf:
                    zf.writestr('doc.kml', kml_content)
                kmz_bytes = kmz_buffer.getvalue()
                
                col_kml, col_kmz = st.columns(2)
                with col_kml:
                    st.download_button(
                        label="ðŸ“¥ TÃ©lÃ©charger KML",
                        data=kml_bytes,
                        file_name=f"itineraire_{datetime.now().strftime('%Y%m%d')}.kml",
                        mime="application/vnd.google-earth.kml+xml",
                        use_container_width=True
                    )
                with col_kmz:
                    st.download_button(
                        label="ðŸ“¥ TÃ©lÃ©charger KMZ (Google Earth)",
                        data=kmz_bytes,
                        file_name=f"itineraire_{datetime.now().strftime('%Y%m%d')}.kmz",
                        mime="application/vnd.google-earth.kmz",
                        use_container_width=True
                    )
            except Exception as e:
                st.warning(f"Impossible de prÃ©parer l'export KML/KMZ: {str(e)[:80]}â€¦")
    
    with tab_export:
        st.subheader("Export")
        
        current_itinerary = st.session_state.manual_itinerary if st.session_state.manual_itinerary else itinerary
        
        excel_data = []
        for day, sdt, edt, desc in current_itinerary:
            excel_data.append({
                "Jour": day,
                "Date": (start_date + timedelta(days=day-1)).strftime("%d/%m/%Y"),
                "DÃ©but": sdt.strftime("%H:%M"),
                "Fin": edt.strftime("%H:%M"),
                "DurÃ©e (min)": int((edt - sdt).total_seconds() / 60),
                "ActivitÃ©": desc
            })
        
        df_export = pd.DataFrame(excel_data)
        
        from io import BytesIO
        output = BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df_export.to_excel(writer, sheet_name='Planning', index=False)
            pd.DataFrame(sites_ordered).to_excel(writer, sheet_name='Sites', index=False)
        
        col_excel, col_html, col_ics = st.columns(3)
        
        with col_excel:
            st.download_button(
                label="ðŸ“¥ TÃ©lÃ©charger Excel",
                data=output.getvalue(),
                file_name=f"mission_{datetime.now().strftime('%Y%m%d')}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True
            )
        
        with col_html:
            html_export = build_professional_html(
                current_itinerary,
                start_date,
                stats,
                sites_ordered,
                segments_summary,
                default_speed_kmh,
                mission_title,
                coords_ordered,
                include_map=st.session_state.get("include_map_prof_html", False),
                lunch_start_time=st.session_state.get("lunch_start_time"),
                lunch_end_time=st.session_state.get("lunch_end_time"),
                lunch_duration_min=st.session_state.get("lunch_duration_min", 60),
                prayer_start_time=st.session_state.get("prayer_start_time"),
                prayer_duration_min=st.session_state.get("prayer_duration_min", 20),
                include_details=st.session_state.get("include_prof_details", True)
            )
            st.download_button(
                label="ðŸ“¥ TÃ©lÃ©charger HTML",
                data=html_export,
                file_name=f"mission_{datetime.now().strftime('%Y%m%d')}.html",
                mime="text/html",
                use_container_width=True
            )
        with col_ics:
            ics_export = build_ics_from_itinerary(current_itinerary, start_date, mission_title)
            st.download_button(
                label="ðŸ“¥ TÃ©lÃ©charger ICS",
                data=ics_export,
                file_name=f"mission_{datetime.now().strftime('%Y%m%d')}.ics",
                mime="text/calendar",
                use_container_width=True
            )

    with tab_report:
        st.subheader("ðŸ“‹ GÃ©nÃ©ration de rapport de mission")
        
        with st.expander("ðŸ¤– GÃ©nÃ©rer un rapport complet", expanded=False):
            st.markdown("**Utilisez l'IA Adja pour gÃ©nÃ©rer un rapport professionnel orientÃ© activitÃ©s**")
            
            # Onglets pour organiser l'interface
            tab_basic, tab_details, tab_questions, tab_construction, tab_generate = st.tabs([
                "ðŸ“ Rapport basique", "ðŸ“‹ DÃ©tails mission", "ðŸ¤– Questions IA Adja", "ðŸ—ï¸ ProcÃ¨s-verbal", "ðŸš€ GÃ©nÃ©ration"
            ])
            
            with tab_basic:
                st.markdown("### ðŸ“„ Rapport rapide (version simplifiÃ©e)")
                
                # Options de rapport basique
                col1, col2 = st.columns(2)
                
                with col1:
                    report_type = st.selectbox(
                        "Type de rapport",
                        ["Rapport complet", "RÃ©sumÃ© exÃ©cutif", "Rapport technique", "Rapport financier", "ProcÃ¨s-verbal professionnel"],
                        help="Choisissez le type de rapport Ã  gÃ©nÃ©rer"
                    )
                
                with col2:
                    report_tone = st.selectbox(
                        "Ton du rapport",
                        ["Professionnel", "Formel", "DÃ©contractÃ©", "Technique"],
                        help="DÃ©finissez le ton du rapport"
                    )
                
                # Options avancÃ©es (sans expander imbriquÃ©)
                st.markdown("**Options avancÃ©es**")
                
                col_opt1, col_opt2 = st.columns(2)
                
                with col_opt1:
                    include_recommendations = st.checkbox("Inclure des recommandations", value=True)
                    include_risks = st.checkbox("Analyser les risques", value=False)
                
                with col_opt2:
                    include_costs = st.checkbox("Estimation des coÃ»ts", value=False)
                    include_timeline = st.checkbox("Planning dÃ©taillÃ©", value=True)
                
                custom_context = st.text_area(
                    "Contexte supplÃ©mentaire (optionnel)",
                    placeholder="Ajoutez des informations spÃ©cifiques Ã  votre mission...",
                    height=100
                )
                
                if st.button("ðŸš€ GÃ©nÃ©rer rapport basique", type="primary", use_container_width=True):
                    if st.session_state.planning_results:
                        # Animation d'attente amÃ©liorÃ©e avec barre de progression
                        progress_bar = st.progress(0)
                        status_text = st.empty()
                        
                        try:
                            status_text.text("ðŸ”„ Collecte des donnÃ©es de mission...")
                            progress_bar.progress(20)
                            mission_data = collect_mission_data_for_ai()
                                
                            status_text.text("ðŸ“ Construction du prompt...")
                            progress_bar.progress(40)
                            prompt = build_report_prompt(
                                mission_data, report_type, report_tone,
                                include_recommendations, include_risks, 
                                include_costs, include_timeline, custom_context
                            )
                            
                            status_text.text("ðŸ¤– GÃ©nÃ©ration du rapport par l'IA Adja...")
                            progress_bar.progress(60)
                            response = requests.post(
                                "https://api.deepseek.com/v1/chat/completions",
                                headers={
                                    "Authorization": f"Bearer {deepseek_api_key}",
                                    "Content-Type": "application/json"
                                },
                                json={
                                    "model": "deepseek-chat",
                                    "messages": [{"role": "user", "content": prompt}],
                                    "temperature": 0.7,
                                    "max_tokens": 4000
                                }
                            )
                            
                            status_text.text("âœ… Finalisation du rapport...")
                            progress_bar.progress(100)
                            
                            if response.status_code == 200:
                                report_content = response.json()["choices"][0]["message"]["content"]
                                
                                # Nettoyer les Ã©lÃ©ments d'animation
                                progress_bar.empty()
                                status_text.empty()
                                
                                st.success("âœ… Rapport gÃ©nÃ©rÃ© avec succÃ¨s!")
                                
                                # Affichage du rapport
                                st.markdown("### ðŸ“„ Votre rapport")
                                st.markdown(report_content)
                                
                                # Boutons de tÃ©lÃ©chargement
                                col_txt, col_md, col_html = st.columns(3)
                                
                                with col_txt:
                                    st.download_button(
                                        label="ðŸ“„ TXT",
                                        data=report_content,
                                        file_name=f"rapport_{datetime.now().strftime('%Y%m%d_%H%M')}.txt",
                                        mime="text/plain",
                                        use_container_width=True
                                    )
                                
                                with col_md:
                                    st.download_button(
                                        label="ðŸ“ MD",
                                        data=report_content,
                                        file_name=f"rapport_{datetime.now().strftime('%Y%m%d_%H%M')}.md",
                                        mime="text/markdown",
                                        use_container_width=True
                                    )
                                
                                with col_html:
                                    html_content = f"""
                                    <!DOCTYPE html>
                                    <html>
                                    <head>
                                        <meta charset="UTF-8">
                                        <title>Rapport de Mission</title>
                                        <style>
                                            body {{ font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; margin: 40px; line-height: 1.6; color: #333; }}
                                            h1, h2, h3 {{ color: #2c3e50; }}
                                            h1 {{ border-bottom: 3px solid #3498db; padding-bottom: 10px; }}
                                            h2 {{ border-left: 4px solid #3498db; padding-left: 15px; }}
                                            .header {{ text-align: center; margin-bottom: 30px; background: #f8f9fa; padding: 20px; border-radius: 10px; }}
                                            .footer {{ margin-top: 30px; text-align: center; font-size: 0.9em; color: #666; }}
                                            ul, ol {{ margin-left: 20px; }}
                                            strong {{ color: #2c3e50; }}
                                        </style>
                                    </head>
                                    <body>
                                        <div class="header">
                                            <h1>Rapport de Mission</h1>
                                            <p><strong>GÃ©nÃ©rÃ© le {datetime.now().strftime('%d/%m/%Y Ã  %H:%M')}</strong></p>
                                            <p>Type: {report_type} | Ton: {report_tone}</p>
                                        </div>
                                    {report_content.replace(chr(10), '<br>')}
                                        <div class="footer">
                                            <p>Rapport gÃ©nÃ©rÃ© automatiquement par l'IA Adja DeepSeek</p>
                                        </div>
                                    </body>
                                    </html>
                                    """
                                    st.download_button(
                                        label="ðŸŒ HTML",
                                        data=html_content,
                                        file_name=f"rapport_{datetime.now().strftime('%Y%m%d_%H%M')}.html",
                                        mime="text/html",
                                        use_container_width=True
                                    )
                            
                            else:
                                st.error(f"âŒ Erreur API: {response.status_code}")
                                st.error(response.text)
                        
                        except Exception as e:
                            st.error(f"âŒ Erreur lors de la gÃ©nÃ©ration: {str(e)}")
                else:
                    st.warning("âš ï¸ Aucun planning disponible. Veuillez d'abord optimiser votre itinÃ©raire.")
            
            with tab_details:
                st.markdown("### ðŸ“‹ Informations dÃ©taillÃ©es de la mission")
                st.info("ðŸ’¡ Remplissez ces informations pour enrichir votre rapport")
                
                mission_data = collect_enhanced_mission_data()
            
            with tab_questions:
                st.markdown("### ðŸ¤– Questions interactives pour personnaliser le rapport")
                st.info("ðŸ’¡ RÃ©pondez Ã  ces questions pour obtenir un rapport sur mesure")
                
                questions_data = ask_interactive_questions()
            
            with tab_construction:
                st.markdown("### ðŸ—ï¸ GÃ©nÃ©ration de procÃ¨s-verbal de chantier")
                st.info("ðŸ’¡ CrÃ©ez un procÃ¨s-verbal professionnel pour vos visites de chantier")
                
                construction_data = collect_construction_report_data()
                
                if st.button("ðŸ“‹ GÃ©nÃ©rer le procÃ¨s-verbal", type="primary", use_container_width=True):
                    if st.session_state.planning_results:
                        # Animation d'attente amÃ©liorÃ©e
                        progress_bar = st.progress(0)
                        status_text = st.empty()
                        
                        try:
                            # Ã‰tape 1: Collecte des donnÃ©es
                            status_text.text("ðŸ“‹ Collecte des donnÃ©es de mission...")
                            progress_bar.progress(20)
                            time.sleep(0.3)
                            
                            mission_data = collect_mission_data_for_ai()
                            questions_data = construction_data
                            
                            # Ã‰tape 2: PrÃ©paration du rapport
                            status_text.text("ðŸ”§ PrÃ©paration du procÃ¨s-verbal...")
                            progress_bar.progress(40)
                            time.sleep(0.3)
                            
                            # Ã‰tape 3: GÃ©nÃ©ration IA Adja Adja
                            status_text.text("ðŸ¤– GÃ©nÃ©ration par l'IA Adja...")
                            progress_bar.progress(70)
                            
                            pv_result = generate_pv_report(mission_data, questions_data, deepseek_api_key)
                            
                            # Ã‰tape 4: Finalisation
                            status_text.text("âœ¨ Finalisation du procÃ¨s-verbal...")
                            progress_bar.progress(100)
                            time.sleep(0.3)
                            
                            # Nettoyage de l'animation
                            progress_bar.empty()
                            status_text.empty()
                            
                            if pv_result["success"]:
                                st.success("âœ… ProcÃ¨s-verbal gÃ©nÃ©rÃ© avec succÃ¨s!")
                                
                                # Affichage du PV
                                st.markdown("### ðŸ“‹ Votre procÃ¨s-verbal")
                                st.markdown(pv_result["content"])
                                
                                # Informations du PV pour les tÃ©lÃ©chargements
                                pv_structure = construction_data.get('pv_structure', 'Structure non spÃ©cifiÃ©e')
                                pv_date = construction_data.get('pv_date', datetime.now().date())
                                pv_site = construction_data.get('pv_site', 'Site non spÃ©cifiÃ©')
                                pv_zone = construction_data.get('pv_zone', 'Zone non spÃ©cifiÃ©e')
                                pv_mission_type = construction_data.get('pv_mission_type', 'Mission non spÃ©cifiÃ©e')
                                pv_responsable = construction_data.get('pv_responsable', 'Responsable non spÃ©cifiÃ©')
                                pv_fonction = construction_data.get('pv_fonction', 'Fonction non spÃ©cifiÃ©e')
                                pv_content = pv_result["content"]
                                
                                # Boutons de tÃ©lÃ©chargement
                                col_pv_txt, col_pv_html, col_pv_pdf, col_pv_rtf = st.columns(4)
                                
                                with col_pv_txt:
                                    pv_txt_content = f"""
PROCÃˆS-VERBAL DE VISITE DE CHANTIER

Structure: {pv_structure}
Date: {pv_date.strftime('%d/%m/%Y')}
Site: {pv_site}
Zone: {pv_zone}

{pv_content}

Fait Ã  Dakar, le {datetime.now().strftime('%d/%m/%Y')}

{pv_responsable}
{pv_fonction}
                                    """
                                    st.download_button(
                                        label="ðŸ“„ TXT",
                                        data=pv_txt_content.strip(),
                                        file_name=f"PV_chantier_{pv_site.replace(' ', '_')}_{pv_date.strftime('%Y%m%d')}.txt",
                                        mime="text/plain",
                                        use_container_width=True
                                    )
                                
                                with col_pv_html:
                                    pv_html_content = f"""
                                        <!DOCTYPE html>
                                        <html>
                                        <head>
                                            <meta charset="UTF-8">
                                            <title>ProcÃ¨s-verbal de visite de chantier</title>
                                            <style>
                                                body {{ font-family: 'Arial', sans-serif; margin: 40px; line-height: 1.6; color: #333; }}
                                                .header {{ text-align: center; margin-bottom: 30px; }}
                                                .header h1 {{ color: #2c3e50; border-bottom: 2px solid #3498db; padding-bottom: 10px; }}
                                                .info-table {{ width: 100%; border-collapse: collapse; margin: 20px 0; }}
                                                .info-table td {{ padding: 8px; border: 1px solid #ddd; }}
                                                .info-table .label {{ background-color: #f8f9fa; font-weight: bold; width: 120px; }}
                                                .signature {{ margin-top: 50px; }}
                                                .signature-line {{ border-top: 1px solid #333; width: 200px; margin: 20px 0; }}
                                            </style>
                                        </head>
                                        <body>
                                            <div class="header">
                                                <h1>ProcÃ¨s-verbal de visite de chantier</h1>
                                                <p><strong>{pv_structure}</strong></p>
                                                <p>Travaux d'extension PA DAL zone {pv_zone}</p>
                                            </div>

                                            <table class="info-table">
                                                <tr>
                                                    <td class="label">DATE:</td>
                                                    <td>{pv_date.strftime('%d/%m/%Y')}</td>
                                                    <td class="label">SITE:</td>
                                                    <td>{pv_site}</td>
                                                </tr>
                                                <tr>
                                                    <td class="label">MISSION:</td>
                                                    <td>{pv_mission_type}</td>
                                                    <td class="label">ZONE:</td>
                                                    <td>{pv_zone}</td>
                                                </tr>
                                                <tr>
                                                    <td class="label">RESPONSABLE:</td>
                                                    <td>{pv_responsable}</td>
                                                    <td class="label">FONCTION:</td>
                                                    <td>{pv_fonction}</td>
                                                </tr>
                                            </table>

                                            {pv_content.replace(chr(10), '<br>')}

                                        <div class="signature">
                                            <p>Fait Ã  Dakar, le {datetime.now().strftime('%d/%m/%Y')}</p>
                                            <div class="signature-line"></div>
                                            <p><strong>{pv_responsable}</strong></p>
                                        </div>
                                    </body>
                                    </html>
                                    """
                                    st.download_button(
                                        label="ðŸŒ HTML",
                                        data=pv_html_content,
                                        file_name=f"PV_chantier_{pv_site.replace(' ', '_')}_{pv_date.strftime('%Y%m%d')}.html",
                                        mime="text/html",
                                        use_container_width=True
                                    )
                                
                                with col_pv_pdf:
                                    if PDF_AVAILABLE:
                                        pv_full_content = f"""
PROCÃˆS-VERBAL DE VISITE DE CHANTIER

Structure: {pv_structure}
Date: {pv_date.strftime('%d/%m/%Y')}
Site: {pv_site}
Zone: {pv_zone}
Mission: {pv_mission_type}
Responsable: {pv_responsable}
Fonction: {pv_fonction}

{pv_content}"""
                                        pdf_data = create_pv_pdf(
                                            content=pv_full_content,
                                            title="ProcÃ¨s-verbal de visite de chantier",
                                            author=pv_responsable
                                        )
                                        st.download_button(
                                            label="ðŸ“„ PDF",
                                            data=pdf_data,
                                            file_name=f"PV_chantier_{pv_site.replace(' ', '_')}_{pv_date.strftime('%Y%m%d')}.pdf",
                                            mime="application/pdf",
                                            use_container_width=True
                                        )
                                    else:
                                        st.info("PDF non disponible")
                                
                                # Suppression de l'export RTF au profit du DOCX
                            
                            else:
                                st.error(f"âŒ Erreur: {pv_result['error']}")
                        
                        except Exception as e:
                            st.error("âŒ Erreur lors de la gÃ©nÃ©ration du procÃ¨s-verbal")

            with tab_generate:
                st.markdown("### ðŸš€ GÃ©nÃ©ration du rapport amÃ©liorÃ©")
                st.info("ðŸ’¡ Utilisez cette section aprÃ¨s avoir rempli les dÃ©tails et rÃ©pondu aux questions")
                
                # VÃ©rification des prÃ©requis
                has_details = hasattr(st.session_state, 'mission_context') and st.session_state.mission_context.get('objective')
                has_questions = 'report_focus' in st.session_state

                if has_details:
                    st.success("âœ… DonnÃ©es dÃ©taillÃ©es collectÃ©es")
                else:
                    st.warning("âš ï¸ Aucune donnÃ©e dÃ©taillÃ©e - Allez dans l'onglet 'DÃ©tails mission'")

                if has_questions:
                    st.success("âœ… Questions rÃ©pondues")
                else:
                    st.warning("âš ï¸ Questions non rÃ©pondues - Allez dans l'onglet 'Questions IA Adja'")
                
                # AperÃ§u des paramÃ¨tres
                if has_questions:
                    st.markdown("**ParamÃ¨tres du rapport :**")
                    col_preview1, col_preview2 = st.columns(2)

                    with col_preview1:
                        if 'report_focus' in st.session_state:
                            st.write(f"ðŸŽ¯ **Focus :** {', '.join(st.session_state.report_focus)}")
                        if 'target_audience' in st.session_state:
                            st.write(f"ðŸ‘¥ **Public :** {st.session_state.target_audience}")

                    with col_preview2:
                        if 'report_length' in st.session_state:
                            st.write(f"ðŸ“„ **Longueur :** {st.session_state.report_length}")
                        if 'specific_request' in st.session_state and st.session_state.specific_request:
                            st.write(f"âœ¨ **Demande spÃ©ciale :** Oui")
                
                # Boutons d'action
                col_gen1, col_gen2 = st.columns([2, 1])

                with col_gen1:
                    generate_enhanced = st.button(
                        "ðŸš€ GÃ©nÃ©rer le rapport amÃ©liorÃ©", 
                        type="primary", 
                        use_container_width=True,
                        disabled=not (has_details or has_questions)
                    )

                with col_gen2:
                    if st.button("ðŸ”„ RÃ©initialiser", use_container_width=True):
                        # Supprimer toutes les donnÃ©es de session liÃ©es au rapport
                        for key in list(st.session_state.keys()):
                            if key.startswith(('mission_', 'activity_', 'report_', 'target_', 'specific_', 'notes_', 'success_', 'contacts_', 'outcomes_', 'follow_up_', 'challenges', 'lessons_', 'recommendations', 'overall_', 'highlight_', 'discuss_', 'future_', 'cost_', 'time_', 'stakeholder_', 'include_')):
                                del st.session_state[key]
                        st.rerun()

                if generate_enhanced:
                    if st.session_state.planning_results:
                        # Animation d'attente amÃ©liorÃ©e avec barre de progression
                        progress_bar = st.progress(0)
                        status_text = st.empty()
                        
                        try:
                            status_text.text("ðŸ”„ Collecte des donnÃ©es de mission...")
                            progress_bar.progress(15)
                            mission_data = collect_mission_data_for_ai()
                            
                            status_text.text("ðŸ“‹ PrÃ©paration des questions...")
                            progress_bar.progress(30)
                            time.sleep(0.5)
                            
                            questions_data = {
                                'report_focus': st.session_state.get('report_focus', []),
                                'target_audience': st.session_state.get('target_audience', 'Direction gÃ©nÃ©rale'),
                                'report_length': st.session_state.get('report_length', 'Moyen (3-5 pages)'),
                                'include_metrics': st.session_state.get('include_metrics', True),
                                'highlight_successes': st.session_state.get('highlight_successes', True),
                                'discuss_challenges': st.session_state.get('discuss_challenges', True),
                                'future_planning': st.session_state.get('future_planning', True),
                                'cost_analysis': st.session_state.get('cost_analysis', False),
                                'time_efficiency': st.session_state.get('time_efficiency', True),
                                'stakeholder_feedback': st.session_state.get('stakeholder_feedback', False),
                                'specific_request': st.session_state.get('specific_request', '')
                            }
                            
                            # Ã‰tape 3: Construction du prompt
                            status_text.text("ðŸ”§ Construction du prompt personnalisÃ©...")
                            progress_bar.progress(50)
                            time.sleep(0.5)
                            
                            # Ã‰tape 4: GÃ©nÃ©ration IA Adja
                            status_text.text("ðŸ¤– GÃ©nÃ©ration du rapport par l'IA Adja...")
                            progress_bar.progress(70)
                            
                            # GÃ©nÃ©ration du rapport
                            report_result = generate_enhanced_ai_report(
                                mission_data, 
                                questions_data,
                                deepseek_api_key
                            )
                            
                            # Ã‰tape 5: Finalisation
                            status_text.text("âœ¨ Finalisation du rapport...")
                            progress_bar.progress(100)
                            time.sleep(0.3)
                            
                            # Nettoyage de l'animation
                            progress_bar.empty()
                            status_text.empty()
                            
                            if report_result["success"]:
                                st.success("âœ… Rapport amÃ©liorÃ© gÃ©nÃ©rÃ© avec succÃ¨s!")
                                
                                # Affichage du rapport
                                st.markdown("### ðŸ“„ Votre rapport amÃ©liorÃ©")
                                report_content = report_result["content"]
                                st.markdown(report_content)
                                
                                # Boutons de tÃ©lÃ©chargement
                                col_txt, col_md, col_html, col_copy = st.columns(4)
                                
                                with col_txt:
                                    st.download_button(
                                        label="ðŸ“„ TXT",
                                        data=report_content,
                                        file_name=f"rapport_ameliore_{datetime.now().strftime('%Y%m%d_%H%M')}.txt",
                                        mime="text/plain",
                                        use_container_width=True
                                    )
                                
                                with col_md:
                                    st.download_button(
                                        label="ðŸ“ MD",
                                        data=report_content,
                                        file_name=f"rapport_ameliore_{datetime.now().strftime('%Y%m%d_%H%M')}.md",
                                        mime="text/markdown",
                                        use_container_width=True
                                    )
                                
                                with col_html:
                                    html_content = f"""
                                    <!DOCTYPE html>
                                    <html>
                                    <head>
                                        <meta charset="UTF-8">
                                        <title>Rapport de Mission AmÃ©liorÃ©</title>
                                        <style>
                                            body {{ font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; margin: 40px; line-height: 1.6; color: #333; }}
                                            h1, h2, h3 {{ color: #2c3e50; }}
                                            h1 {{ border-bottom: 3px solid #3498db; padding-bottom: 10px; }}
                                            h2 {{ border-left: 4px solid #3498db; padding-left: 15px; }}
                                            .header {{ text-align: center; margin-bottom: 30px; background: #f8f9fa; padding: 20px; border-radius: 10px; }}
                                            .footer {{ margin-top: 30px; text-align: center; font-size: 0.9em; color: #666; }}
                                            ul, ol {{ margin-left: 20px; }}
                                            strong {{ color: #2c3e50; }}
                                        </style>
                                    </head>
                                    <body>
                                        <div class="header">
                                            <h1>Rapport de Mission AmÃ©liorÃ©</h1>
                                            <p><strong>GÃ©nÃ©rÃ© le {datetime.now().strftime('%d/%m/%Y Ã  %H:%M')}</strong></p>
                                            <p>Public cible: {questions_data.get('target_audience', 'Non spÃ©cifiÃ©')}</p>
                                        </div>
                                        {report_content.replace(chr(10), '<br>')}
                                        <div class="footer">
                                            <p>Rapport gÃ©nÃ©rÃ© automatiquement par l'IA Adja DeepSeek</p>
                                        </div>
                                    </body>
                                    </html>
                                    """
                                    st.download_button(
                                        label="ðŸŒ HTML",
                                        data=html_content,
                                        file_name=f"rapport_ameliore_{datetime.now().strftime('%Y%m%d_%H%M')}.html",
                                        mime="text/html",
                                        use_container_width=True
                                    )
                                
                                with col_copy:
                                    st.code(report_content, language=None)
                            
                            else:
                                st.error(f"âŒ Erreur: {report_result['error']}")
                        
                        except Exception as e:
                            st.error(f"âŒ Erreur lors de la gÃ©nÃ©ration: {str(e)}")
                    else:
                        st.warning("âš ï¸ Aucun planning disponible. Veuillez d'abord optimiser votre itinÃ©raire.")

# --------------------------
# MODULE RAPPORT IA ADJA AMÃ‰LIORÃ‰ (ANCIEN - Ã€ SUPPRIMER)
# --------------------------
if False and st.session_state.planning_results:
    st.markdown("---")
    st.header("ðŸ“‹ GÃ©nÃ©ration de rapport de mission")
    
    with st.expander("ðŸ¤– GÃ©nÃ©rer un rapport complet", expanded=False):
        st.markdown("**Utilisez l'IA Adja pour gÃ©nÃ©rer un rapport professionnel orientÃ© activitÃ©s**")
        
        # Onglets pour organiser l'interface
        tab_basic, tab_details, tab_questions, tab_construction, tab_generate = st.tabs([
            "ðŸ“ Rapport basique", "ðŸ“‹ DÃ©tails mission", "ðŸ¤– Questions IA Adja", "ðŸ—ï¸ ProcÃ¨s-verbal", "ðŸš€ GÃ©nÃ©ration"
        ])
        
        with tab_basic:
            st.markdown("### ðŸ“„ Rapport rapide (version simplifiÃ©e)")
            
            # Options de rapport basique
            col1, col2 = st.columns(2)
            
            with col1:
                report_type = st.selectbox(
                    "Type de rapport",
                    ["Rapport complet", "RÃ©sumÃ© exÃ©cutif", "Rapport technique", "Rapport financier", "ProcÃ¨s-verbal professionnel"],
                    help="Choisissez le type de rapport Ã  gÃ©nÃ©rer"
                )
            
            with col2:
                report_tone = st.selectbox(
                    "Ton du rapport",
                    ["Professionnel", "Formel", "DÃ©contractÃ©", "Technique"],
                    help="DÃ©finissez le ton du rapport"
                )
            
            # Options avancÃ©es (sans expander imbriquÃ©)
            st.markdown("**Options avancÃ©es**")
            col3, col4 = st.columns(2)
            
            with col3:
                include_recommendations = st.checkbox("Inclure des recommandations", value=True)
                include_risks = st.checkbox("Inclure l'analyse des risques", value=True)
            
            with col4:
                include_costs = st.checkbox("Inclure l'analyse des coÃ»ts", value=True)
                include_timeline = st.checkbox("Inclure la timeline dÃ©taillÃ©e", value=True)
            
            custom_context = st.text_area(
                "Contexte supplÃ©mentaire (optionnel)",
                placeholder="Ajoutez des informations spÃ©cifiques sur votre mission, objectifs, contraintes...",
                height=100
            )
            
            # Bouton de gÃ©nÃ©ration basique
            if st.button("ðŸš€ GÃ©nÃ©rer le rapport basique", type="secondary", use_container_width=True):
                if not deepseek_api_key:
                    st.error("âŒ ClÃ© API DeepSeek manquante")
                else:
                    # Animation amÃ©liorÃ©e avec barre de progression
                    progress_bar = st.progress(0)
                    status_text = st.empty()
                    
                    # Ã‰tape 1: Collecte des donnÃ©es
                    status_text.text("ðŸ“‹ Collecte des donnÃ©es de mission...")
                    progress_bar.progress(20)
                    mission_data = collect_mission_data_for_ai()
                    
                    # Ã‰tape 2: PrÃ©paration du prompt
                    status_text.text("ðŸ¤– Construction du prompt IA Adja...")
                    progress_bar.progress(40)
                    time.sleep(0.5)
                    
                    # Ã‰tape 3: GÃ©nÃ©ration IA Adja
                    status_text.text("ðŸ¤– GÃ©nÃ©ration du rapport par l'IA Adja...")
                    progress_bar.progress(70)
                    time.sleep(0.3)
                    
                    # GÃ©nÃ©ration selon le type de rapport sÃ©lectionnÃ©
                    if report_type == "ProcÃ¨s-verbal professionnel":
                            # GÃ©nÃ©ration du procÃ¨s-verbal avec l'IA Adja
                            questions_data_pv = {
                                'context': custom_context,
                                'observations': 'Observations dÃ©taillÃ©es de la mission',
                                'issues': 'ProblÃ¨mes identifiÃ©s lors de la mission',
                                'actions': 'Actions rÃ©alisÃ©es pendant la mission',
                                'recommendations': 'Recommandations pour la suite'
                            }
                            
                            report_content, error = generate_pv_report(
                                mission_data, 
                                questions_data_pv,
                                deepseek_api_key
                            )
                            
                            if error:
                                st.error(f"âŒ Erreur lors de la gÃ©nÃ©ration du PV: {error}")
                            else:
                                st.success("âœ… ProcÃ¨s-verbal gÃ©nÃ©rÃ© avec succÃ¨s!")
                                
                                # Affichage du PV
                                st.markdown("### ðŸ“‹ ProcÃ¨s-verbal gÃ©nÃ©rÃ©")
                                st.markdown(report_content)
                                
                                # Options d'export spÃ©cialisÃ©es pour le PV
                                st.markdown("### ðŸ’¾ Export du procÃ¨s-verbal")
                                col_txt, col_html, col_pdf = st.columns(3)
                                
                                with col_txt:
                                    st.download_button(
                                        label="ðŸ“„ TÃ©lÃ©charger TXT",
                                        data=report_content,
                                        file_name=f"pv_mission_{datetime.now().strftime('%Y%m%d_%H%M')}.txt",
                                        mime="text/plain",
                                        use_container_width=True
                                    )
                                
                                with col_html:
                                    # HTML formatÃ© pour le PV
                                    html_pv = f"""
                                    <!DOCTYPE html>
                                    <html>
                                    <head>
                                        <meta charset="UTF-8">
                                        <title>ProcÃ¨s-verbal de Mission</title>
                                        <style>
                                            @page {{ margin: 2cm; }}
                                            body {{ 
                                                font-family: 'Times New Roman', serif; 
                                                font-size: 12pt; 
                                                line-height: 1.4; 
                                                color: #000; 
                                                margin: 0;
                                            }}
                                            .header {{ 
                                                text-align: center; 
                                                margin-bottom: 30px; 
                                                border-bottom: 2px solid #000;
                                                padding-bottom: 15px;
                                            }}
                                            .header h1 {{ 
                                                font-size: 18pt; 
                                                margin: 0; 
                                                text-transform: uppercase;
                                                font-weight: bold;
                                            }}
                                            h2 {{ 
                                                font-size: 14pt; 
                                                margin: 25px 0 10px 0; 
                                                text-decoration: underline;
                                                font-weight: bold;
                                            }}
                                            h3 {{ 
                                                font-size: 12pt; 
                                                margin: 20px 0 8px 0; 
                                                font-weight: bold;
                                            }}
                                            .signature {{ 
                                                margin-top: 40px; 
                                                text-align: right;
                                            }}
                                            .signature-line {{ 
                                                border-top: 1px solid #000; 
                                                width: 200px; 
                                                margin: 30px 0 5px auto;
                                            }}
                                            ul {{ margin-left: 20px; }}
                                            li {{ margin-bottom: 5px; }}
                                        </style>
                                    </head>
                                    <body>
                                        <div class="header">
                                            <h1>ProcÃ¨s-verbal de Mission</h1>
                                            <p><strong>GÃ©nÃ©rÃ© le {datetime.now().strftime('%d/%m/%Y Ã  %H:%M')}</strong></p>
                                        </div>
                                        {report_content.replace(chr(10), '<br>')}
                                        <div class="signature">
                                            <p>Fait Ã  Dakar, le {datetime.now().strftime('%d/%m/%Y')}</p>
                                            <div class="signature-line"></div>
                                            <p><strong>Responsable Mission</strong></p>
                                        </div>
                                    </body>
                                    </html>
                                    """
                                    
                                    st.download_button(
                                        label="ðŸŒ TÃ©lÃ©charger HTML",
                                        data=html_pv,
                                        file_name=f"pv_mission_{datetime.now().strftime('%Y%m%d_%H%M')}.html",
                                        mime="text/html",
                                        use_container_width=True
                                    )
                                
                                with col_pdf:
                                    st.info("ðŸ’¡ Ouvrez le fichier HTML dans votre navigateur et utilisez 'Imprimer > Enregistrer au format PDF' pour obtenir un PDF professionnel.")
                    else:
                        # GÃ©nÃ©ration du rapport basique (utilisation de l'ancienne fonction)
                        # Pour le rapport basique, on utilise une version simplifiÃ©e
                        questions_data_simple = {
                                'report_focus': report_type,
                                'target_audience': 'Ã‰quipe',
                                'report_length': 'Moyen',
                                'include_successes': include_recommendations,
                                'include_challenges': include_risks,
                                'include_costs': include_costs,
                                'include_planning': include_timeline,
                                'custom_requests': custom_context
                            }
                        
                        report_content = generate_enhanced_ai_report(
                            mission_data_simple, 
                            questions_data_simple,
                            deepseek_api_key
                        )
                        
                        # Ã‰tape 4: Finalisation
                        status_text.text("âœ… Finalisation du rapport...")
                        progress_bar.progress(100)
                        time.sleep(0.3)
                        
                        # Nettoyage des Ã©lÃ©ments d'animation
                        progress_bar.empty()
                        status_text.empty()
                        
                        if report_content:
                            st.success("âœ… Rapport gÃ©nÃ©rÃ© avec succÃ¨s!")
                            
                            # Affichage du rapport
                            st.markdown("### ðŸ“„ Rapport gÃ©nÃ©rÃ©")
                            st.markdown(report_content)
                            
                            # Options d'export
                            st.markdown("### ðŸ’¾ Export du rapport")
                            
                            # PremiÃ¨re ligne : formats de base
                            col_txt, col_md, col_html = st.columns(3)
                            
                            with col_txt:
                                st.download_button(
                                    label="ðŸ“„ TÃ©lÃ©charger TXT",
                                    data=report_content,
                                    file_name=f"rapport_mission_{datetime.now().strftime('%Y%m%d_%H%M')}.txt",
                                    mime="text/plain",
                                    use_container_width=True
                                )
                            
                            with col_md:
                                st.download_button(
                                    label="ðŸ“ TÃ©lÃ©charger MD",
                                    data=report_content,
                                    file_name=f"rapport_mission_{datetime.now().strftime('%Y%m%d_%H%M')}.md",
                                    mime="text/markdown",
                                    use_container_width=True
                                )
                            
                            with col_html:
                                # Conversion HTML pour PDF
                                html_report = f"""
                                <!DOCTYPE html>
                                <html>
                                <head>
                                    <meta charset="UTF-8">
                                    <title>Rapport de Mission</title>
                                    <style>
                                        body {{ font-family: Arial, sans-serif; margin: 40px; line-height: 1.6; }}
                                        h1, h2, h3 {{ color: #2c3e50; }}
                                        .header {{ text-align: center; margin-bottom: 30px; }}
                                    </style>
                                </head>
                                <body>
                                    <div class="header">
                                        <h1>Rapport de Mission</h1>
                                        <p>GÃ©nÃ©rÃ© le {datetime.now().strftime('%d/%m/%Y Ã  %H:%M')}</p>
                                    </div>
                                    {report_content.replace(chr(10), '<br>')}
                                </body>
                                </html>
                                """
                                
                                st.download_button(
                                    label="ðŸŒ TÃ©lÃ©charger HTML",
                                    data=html_report,
                                    file_name=f"rapport_mission_{datetime.now().strftime('%Y%m%d_%H%M')}.html",
                                    mime="text/html",
                                    use_container_width=True
                                )
                            
                            # DeuxiÃ¨me ligne : formats professionnels (PDF et Word)
                            if PDF_AVAILABLE:
                                st.markdown("#### ðŸ“‹ Formats professionnels")
                                col_pdf, col_word_rtf, col_word_docx = st.columns(3)
                                
                                with col_pdf:
                                    try:
                                        pdf_data = create_pv_pdf(
                                            content=report_content,
                                            title="Rapport de Mission",
                                            author="Responsable Mission"
                                        )
                                        st.download_button(
                                            label="ðŸ“„ TÃ©lÃ©charger PDF",
                                            data=pdf_data,
                                            file_name=f"rapport_mission_{datetime.now().strftime('%Y%m%d_%H%M')}.pdf",
                                            mime="application/pdf",
                                            use_container_width=True
                                        )
                                    except Exception as e:
                                        st.error(f"Erreur gÃ©nÃ©ration PDF: {str(e)}")
                                
                                # Suppression de l'export Word RTF, on conserve uniquement DOCX
                                
                                with col_word_docx:
                                    try:
                                        docx_data = create_docx_document(
                                            content=report_content,
                                            title="Rapport de Mission"
                                        )
                                        st.download_button(
                                            label="ðŸ“„ Word (.docx)",
                                            data=docx_data,
                                            file_name=f"rapport_mission_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                            use_container_width=True
                                        )
                                    except Exception as e:
                                        st.error(f"Erreur gÃ©nÃ©ration Word DOCX: {str(e)}")
                            if not PDF_AVAILABLE:
                                st.info("ðŸ’¡ Installez reportlab pour activer l'export PDF et Word professionnel.")
                        else:
                            st.error("âŒ Erreur lors de la gÃ©nÃ©ration du rapport")
        
        with tab_details:
            st.markdown("### ðŸ“‹ Collecte de donnÃ©es dÃ©taillÃ©es")
            st.info("ðŸ’¡ Remplissez ces informations pour obtenir un rapport plus riche et personnalisÃ©")
            
            # Interface de collecte de donnÃ©es enrichies
            collect_enhanced_mission_data()
        
        with tab_questions:
            st.markdown("### ðŸ¤– Questions pour personnaliser le rapport")
            st.info("ðŸ’¡ RÃ©pondez Ã  ces questions pour que l'IA Adja gÃ©nÃ¨re un rapport adaptÃ© Ã  vos besoins")
            
            # Interface de questions interactives
            questions_data = ask_interactive_questions()
        
        with tab_construction:
            st.markdown("### ðŸ—ï¸ ProcÃ¨s-verbal de visite de chantier")
            st.info("ðŸ’¡ GÃ©nÃ©rez un procÃ¨s-verbal professionnel au format officiel")
            
            # Formulaire pour procÃ¨s-verbal de chantier
            st.markdown("#### ðŸ“‹ Informations gÃ©nÃ©rales")
            
            col_pv1, col_pv2 = st.columns(2)
            
            with col_pv1:
                pv_date = st.date_input("ðŸ“… Date de visite", value=datetime.now().date())
                pv_site = st.text_input("ðŸ—ï¸ Site/Chantier", placeholder="Ex: Villengara et Kolda")
                pv_structure = st.text_input("ðŸ¢ Structure", placeholder="Ex: DAL/GPR/ESP")
                pv_zone = st.text_input("ðŸ—ºï¸ Titre projet", placeholder="Ex: PA DAL zone SUD")
            
            with col_pv2:
                pv_mission_type = st.selectbox(
                    "ðŸ“ Type de mission",
                    ["Visite de chantier", "Inspection technique", "Suivi de travaux", "RÃ©ception de travaux", "Autre"]
                )
                pv_responsable = st.text_input("ðŸ‘¤ Responsable mission", placeholder="Ex: Moctar TALL")
                pv_fonction = st.text_input("ðŸ’¼ Fonction", placeholder="Ex: IngÃ©nieur")
                pv_contact = st.text_input("ðŸ“ž Contact", placeholder="Ex: +221 XX XXX XX XX")
            
            st.markdown("#### ðŸŽ¯ Objectifs de la mission")
            pv_objectifs = st.text_area(
                "DÃ©crivez les objectifs principaux",
                placeholder="Ex: ContrÃ´ler l'avancement des travaux, vÃ©rifier la conformitÃ©, identifier les problÃ¨mes...",
                height=100
            )
            
            st.markdown("#### ðŸ“Š Observations et constats")
            
            # Sections d'observations
            col_obs1, col_obs2 = st.columns(2)
            
            with col_obs1:
                st.markdown("**ðŸ” Constats positifs**")
                pv_positifs = st.text_area(
                    "Points positifs observÃ©s",
                    placeholder="Ex: Respect des dÃ©lais, qualitÃ© des matÃ©riaux, sÃ©curitÃ©...",
                    height=120,
                    key="pv_positifs"
                )
                
                st.markdown("**âš ï¸ Points d'attention**")
                pv_attention = st.text_area(
                    "Points nÃ©cessitant une attention",
                    placeholder="Ex: Retards mineurs, ajustements nÃ©cessaires...",
                    height=120,
                    key="pv_attention"
                )
            
            with col_obs2:
                st.markdown("**âŒ ProblÃ¨mes identifiÃ©s**")
                pv_problemes = st.text_area(
                    "ProblÃ¨mes et non-conformitÃ©s",
                    placeholder="Ex: DÃ©fauts de construction, non-respect des normes...",
                    height=120,
                    key="pv_problemes"
                )
                
                st.markdown("**ðŸ’¡ Recommandations**")
                pv_recommandations = st.text_area(
                    "Actions recommandÃ©es",
                    placeholder="Ex: Corrections Ã  apporter, amÃ©liorations suggÃ©rÃ©es...",
                    height=120,
                    key="pv_recommandations"
                )
            
            st.markdown("#### ðŸ“ˆ Avancement et planning")
            col_plan1, col_plan2 = st.columns(2)
            
            with col_plan1:
                pv_avancement = st.slider("ðŸ“Š Avancement global (%)", 0, 100, 50)
                pv_respect_delais = st.selectbox("â° Respect des dÃ©lais", ["Conforme", "LÃ©ger retard", "Retard important"])
            
            with col_plan2:
                pv_prochaine_visite = st.date_input("ðŸ“… Prochaine visite prÃ©vue", value=datetime.now().date() + timedelta(days=30))
                pv_urgence = st.selectbox("ðŸš¨ Niveau d'urgence", ["Faible", "Moyen", "Ã‰levÃ©", "Critique"])
            
            st.markdown("#### ðŸ‘¥ Participants et contacts")
            pv_participants = st.text_area(
                "Liste des participants Ã  la visite",
                placeholder="Ex: Moctar TALL (IngÃ©nieur), Jean DUPONT (Chef de chantier), Marie MARTIN (Architecte)...",
                height=80
            )
            
            # GÃ©nÃ©ration du procÃ¨s-verbal
            if st.button("ðŸ“‹ GÃ©nÃ©rer le procÃ¨s-verbal", type="primary", use_container_width=True):
                if not deepseek_api_key:
                    st.error("âŒ ClÃ© API DeepSeek manquante")
                elif not pv_site or not pv_objectifs:
                    st.error("âŒ Veuillez remplir au minimum le site et les objectifs")
                else:
                    # Animation d'attente amÃ©liorÃ©e
                    progress_bar = st.progress(0)
                    status_text = st.empty()
                    
                    try:
                        # Ã‰tape 1: Collecte des donnÃ©es
                        status_text.text("ðŸ“‹ Collecte des informations du chantier...")
                        progress_bar.progress(15)
                        time.sleep(0.3)
                        
                        # DonnÃ©es pour le procÃ¨s-verbal
                        pv_data = {
                            'date': pv_date.strftime('%d/%m/%Y'),
                            'site': pv_site,
                            'structure': pv_structure,
                            'zone': pv_zone,
                            'mission_type': pv_mission_type,
                            'responsable': pv_responsable,
                            'fonction': pv_fonction,
                            'contact': pv_contact,
                            'objectifs': pv_objectifs,
                            'positifs': pv_positifs,
                            'attention': pv_attention,
                            'problemes': pv_problemes,
                            'recommandations': pv_recommandations,
                            'avancement': pv_avancement,
                            'respect_delais': pv_respect_delais,
                            'prochaine_visite': pv_prochaine_visite.strftime('%d/%m/%Y'),
                            'urgence': pv_urgence,
                            'participants': pv_participants
                        }
                        
                        # Mise Ã  jour de l'animation - PrÃ©paration du rapport
                        progress_bar.progress(45)
                        status_text.text("ðŸ“ PrÃ©paration du rapport de chantier...")
                        time.sleep(0.5)
                        
                        # Mise Ã  jour de l'animation - GÃ©nÃ©ration IA Adja
                        progress_bar.progress(70)
                        status_text.text("ðŸ¤– GÃ©nÃ©ration du rapport avec l'IA Adja...")
                        
                        # GÃ©nÃ©ration avec l'IA Adja
                        pv_content = generate_construction_report(pv_data, deepseek_api_key)
                        
                        # Mise Ã  jour de l'animation - Finalisation
                        progress_bar.progress(100)
                        status_text.text("âœ… Rapport gÃ©nÃ©rÃ© avec succÃ¨s!")
                        time.sleep(0.5)
                        
                        # Nettoyage de l'animation
                        progress_bar.empty()
                        status_text.empty()
                        
                        if pv_content:
                            st.success("âœ… ProcÃ¨s-verbal gÃ©nÃ©rÃ© avec succÃ¨s!")
                            
                            # Affichage du procÃ¨s-verbal
                            st.markdown("### ðŸ“„ ProcÃ¨s-verbal gÃ©nÃ©rÃ©")
                            st.markdown(pv_content)
                            
                            # Options d'export spÃ©cialisÃ©es
                            st.markdown("### ðŸ’¾ Export du procÃ¨s-verbal")
                            col_pv_txt, col_pv_pdf, col_pv_word = st.columns(3)
                            
                            with col_pv_txt:
                                st.download_button(
                                    label="ðŸ“„ Format TXT",
                                    data=pv_content,
                                    file_name=f"PV_chantier_{pv_site.replace(' ', '_')}_{pv_date.strftime('%Y%m%d')}.txt",
                                    mime="text/plain",
                                    use_container_width=True
                                )
                            
                            with col_pv_pdf:
                                # HTML formatÃ© pour impression PDF
                                html_pv = f"""
                                <!DOCTYPE html>
                                <html>
                                <head>
                                    <meta charset="UTF-8">
                                    <title>ProcÃ¨s-verbal de visite de chantier</title>
                                    <style>
                                        @page {{ margin: 2cm; }}
                                        body {{ 
                                            font-family: 'Times New Roman', serif; 
                                            font-size: 12pt; 
                                            line-height: 1.4; 
                                            color: #000; 
                                            margin: 0;
                                        }}
                                        .header {{ 
                                            text-align: center; 
                                            margin-bottom: 30px; 
                                            border-bottom: 2px solid #000;
                                            padding-bottom: 15px;
                                        }}
                                        .header h1 {{ 
                                            font-size: 18pt; 
                                            margin: 0; 
                                            text-transform: uppercase;
                                            font-weight: bold;
                                        }}
                                        .info-table {{ 
                                            width: 100%; 
                                            border-collapse: collapse; 
                                            margin: 20px 0;
                                        }}
                                        .info-table td {{ 
                                            border: 1px solid #000; 
                                            padding: 8px; 
                                            vertical-align: top;
                                        }}
                                        .info-table .label {{ 
                                            background-color: #f0f0f0; 
                                            font-weight: bold; 
                                            width: 30%;
                                        }}
                                        h2 {{ 
                                            font-size: 14pt; 
                                            margin: 25px 0 10px 0; 
                                            text-decoration: underline;
                                            font-weight: bold;
                                        }}
                                        h3 {{ 
                                            font-size: 12pt; 
                                            margin: 20px 0 8px 0; 
                                            font-weight: bold;
                                        }}
                                        .signature {{ 
                                            margin-top: 40px; 
                                            text-align: right;
                                        }}
                                        .signature-line {{ 
                                            border-top: 1px solid #000; 
                                            width: 200px; 
                                            margin: 30px 0 5px auto;
                                        }}
                                        ul {{ margin-left: 20px; }}
                                        li {{ margin-bottom: 5px; }}
                                    </style>
                                </head>
                                <body>
                                    <div class="header">
                                        <h1>ProcÃ¨s-verbal de visite de chantier</h1>
                                        <p><strong>{pv_structure}</strong></p>
                                        <p>Travaux d'extension PA DAL zone {pv_zone}</p>
                                    </div>
                                    
                                    <table class="info-table">
                                        <tr>
                                            <td class="label">DATE:</td>
                                            <td>{pv_date.strftime('%d/%m/%Y')}</td>
                                            <td class="label">SITE:</td>
                                            <td>{pv_site}</td>
                                        </tr>
                                        <tr>
                                            <td class="label">MISSION:</td>
                                            <td>{pv_mission_type}</td>
                                            <td class="label">ZONE:</td>
                                            <td>{pv_zone}</td>
                                        </tr>
                                        <tr>
                                            <td class="label">RESPONSABLE:</td>
                                            <td>{pv_responsable}</td>
                                            <td class="label">FONCTION:</td>
                                            <td>{pv_fonction}</td>
                                        </tr>
                                    </table>
                                    
                                    {pv_content.replace(chr(10), '<br>')}
                                    
                                    <div class="signature">
                                        <p>Fait Ã  Dakar, le {datetime.now().strftime('%d/%m/%Y')}</p>
                                        <div class="signature-line"></div>
                                        <p><strong>{pv_responsable}</strong></p>
                                    </div>
                                </body>
                                </html>
                                """
                                
                                st.download_button(
                                    label="ðŸ“‹ Format HTML",
                                    data=html_pv,
                                    file_name=f"PV_chantier_{pv_site.replace(' ', '_')}_{pv_date.strftime('%Y%m%d')}.html",
                                    mime="text/html",
                                    use_container_width=True
                                )
                            
                            with col_pv_word:
                                # Format Word-compatible
                                word_content = f"""
                                PROCÃˆS-VERBAL DE VISITE DE CHANTIER
                                
                                Structure: {pv_structure}
                                Date: {pv_date.strftime('%d/%m/%Y')}
                                Site: {pv_site}
                                Zone: {pv_zone}
                                
                                {pv_content}
                                
                                Fait Ã  Dakar, le {datetime.now().strftime('%d/%m/%Y')}
                                
                                {pv_responsable}
                                {pv_fonction}
                                """
                                
                                st.download_button(
                                    label="ðŸ“ Format TXT",
                                    data=word_content,
                                    file_name=f"PV_chantier_{pv_site.replace(' ', '_')}_{pv_date.strftime('%Y%m%d')}.txt",
                                    mime="text/plain",
                                    use_container_width=True
                                )
                            
                            # DeuxiÃ¨me ligne : formats professionnels (PDF et Word)
                            if PDF_AVAILABLE:
                                st.markdown("#### ðŸ“‹ Formats professionnels")
                                col_pv_pdf, col_pv_rtf, col_pv_docx = st.columns(3)
                                
                                with col_pv_pdf:
                                    try:
                                        # Contenu formatÃ© pour le PV
                                        pv_full_content = f"""Structure: {pv_structure}
Date: {pv_date.strftime('%d/%m/%Y')}
Site: {pv_site}
Zone: {pv_zone}
Mission: {pv_mission_type}
Responsable: {pv_responsable}
Fonction: {pv_fonction}

{pv_content}"""
                                        
                                        pdf_data = create_pv_pdf(
                                            content=pv_full_content,
                                            title="ProcÃ¨s-verbal de visite de chantier",
                                            author=pv_responsable
                                        )
                                        st.download_button(
                                            label="ðŸ“„ PDF",
                                            data=pdf_data,
                                            file_name=f"PV_chantier_{pv_site.replace(' ', '_')}_{pv_date.strftime('%Y%m%d')}.pdf",
                                            mime="application/pdf",
                                            use_container_width=True
                                        )
                                    except Exception as e:
                                        st.error(f"Erreur gÃ©nÃ©ration PDF: {str(e)}")
                                
                                # Suppression de l'export Word RTF, on conserve uniquement DOCX
                                
                                with col_pv_docx:
                                    try:
                                        docx_data = create_docx_document(
                                            content=pv_full_content,
                                            title="ProcÃ¨s-verbal de visite de chantier"
                                        )
                                        st.download_button(
                                            label="ðŸ“„ Word (.docx)",
                                            data=docx_data,
                                            file_name=f"PV_chantier_{pv_site.replace(' ', '_')}_{pv_date.strftime('%Y%m%d')}.docx",
                                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                            use_container_width=True
                                        )
                                    except Exception as e:
                                        st.error(f"Erreur gÃ©nÃ©ration Word DOCX: {str(e)}")
                            else:
                                st.info("ðŸ’¡ Installez reportlab pour activer l'export PDF et Word professionnel.")

                    except Exception as e:
                        st.error(f"âŒ Erreur lors de la gÃ©nÃ©ration du procÃ¨s-verbal: {str(e)}")
                        try:
                            progress_bar.empty()
                            status_text.empty()
                        except Exception:
                            pass

        with tab_generate:
            st.markdown("### ðŸš€ GÃ©nÃ©ration du rapport amÃ©liorÃ©")
            st.info("ðŸ’¡ Utilisez cette section aprÃ¨s avoir rempli les dÃ©tails et rÃ©pondu aux questions")
            
            # VÃ©rification des donnÃ©es disponibles
            has_details = hasattr(st.session_state, 'mission_context') and st.session_state.mission_context.get('objective')
            has_questions = 'report_focus' in st.session_state
            
            if has_details:
                st.success("âœ… DonnÃ©es dÃ©taillÃ©es collectÃ©es")
            else:
                st.warning("âš ï¸ Aucune donnÃ©e dÃ©taillÃ©e - Allez dans l'onglet 'DÃ©tails mission'")
            
            if has_questions:
                st.success("âœ… Questions rÃ©pondues")
            else:
                st.warning("âš ï¸ Questions non rÃ©pondues - Allez dans l'onglet 'Questions IA Adja'")
            
            # AperÃ§u des paramÃ¨tres
            if has_questions:
                st.markdown("**ParamÃ¨tres du rapport :**")
                col_preview1, col_preview2 = st.columns(2)
                
                with col_preview1:
                    if 'report_focus' in st.session_state:
                        st.write(f"ðŸŽ¯ **Focus :** {', '.join(st.session_state.report_focus)}")
                    if 'target_audience' in st.session_state:
                        st.write(f"ðŸ‘¥ **Public :** {st.session_state.target_audience}")
                
                with col_preview2:
                    if 'report_length' in st.session_state:
                        st.write(f"ðŸ“„ **Longueur :** {st.session_state.report_length}")
                    if 'specific_request' in st.session_state and st.session_state.specific_request:
                        st.write(f"âœ¨ **Demande spÃ©ciale :** Oui")
            
            # Bouton de gÃ©nÃ©ration amÃ©liorÃ©e
            col_gen1, col_gen2 = st.columns([2, 1])
            
            with col_gen1:
                generate_enhanced = st.button(
                    "ðŸš€ GÃ©nÃ©rer le rapport amÃ©liorÃ©", 
                    type="primary", 
                    use_container_width=True,
                    disabled=not (has_details or has_questions)
                )
            
            with col_gen2:
                if st.button("ðŸ”„ RÃ©initialiser", use_container_width=True):
                    # RÃ©initialiser les donnÃ©es
                    for key in list(st.session_state.keys()):
                        if key.startswith(('mission_', 'activity_', 'report_', 'target_', 'specific_', 'notes_', 'success_', 'contacts_', 'outcomes_', 'follow_up_', 'challenges', 'lessons_', 'recommendations', 'overall_', 'highlight_', 'discuss_', 'future_', 'cost_', 'time_', 'stakeholder_', 'include_')):
                            del st.session_state[key]
                    st.rerun()
            
            if generate_enhanced:
                if not deepseek_api_key:
                    st.error("âŒ ClÃ© API DeepSeek manquante")
                else:
                    # Initialisation de l'animation d'attente
                    progress_bar = st.progress(0)
                    status_text = st.empty()
                    
                    # Ã‰tape 1: Collecte des donnÃ©es de mission
                    status_text.text("ðŸ“‹ Collecte des donnÃ©es de mission...")
                    progress_bar.progress(20)
                    time.sleep(0.5)
                    mission_data = collect_mission_data_for_ai()
                    
                    # Ã‰tape 2: Collecte des rÃ©ponses aux questions
                    status_text.text("â“ Collecte des rÃ©ponses aux questions...")
                    progress_bar.progress(40)
                    time.sleep(0.5)
                    questions_data = {
                        'report_focus': st.session_state.get('report_focus', []),
                        'target_audience': st.session_state.get('target_audience', 'Direction gÃ©nÃ©rale'),
                        'report_length': st.session_state.get('report_length', 'Moyen (3-5 pages)'),
                        'include_metrics': st.session_state.get('include_metrics', True),
                        'highlight_successes': st.session_state.get('highlight_successes', True),
                        'discuss_challenges': st.session_state.get('discuss_challenges', True),
                        'future_planning': st.session_state.get('future_planning', True),
                        'cost_analysis': st.session_state.get('cost_analysis', False),
                        'time_efficiency': st.session_state.get('time_efficiency', True),
                        'stakeholder_feedback': st.session_state.get('stakeholder_feedback', False),
                        'specific_request': st.session_state.get('specific_request', '')
                    }
                    
                    # Ã‰tape 3: Construction du prompt
                    status_text.text("ðŸ”§ Construction du prompt personnalisÃ©...")
                    progress_bar.progress(60)
                    time.sleep(0.5)
                    
                    # Ã‰tape 4: GÃ©nÃ©ration du rapport amÃ©liorÃ©
                    status_text.text("ðŸ¤– GÃ©nÃ©ration du rapport amÃ©liorÃ© par l'IA Adja...")
                    progress_bar.progress(80)
                    time.sleep(0.5)
                    report_content = generate_enhanced_ai_report(
                        mission_data, 
                        questions_data,
                        deepseek_api_key
                        )
                        
                    # Ã‰tape 5: Finalisation
                    status_text.text("âœ… Finalisation du rapport...")
                    progress_bar.progress(100)
                    time.sleep(0.5)
                    
                    # Nettoyage de l'animation
                    progress_bar.empty()
                    status_text.empty()
                        
                    if report_content:
                            st.success("âœ… Rapport amÃ©liorÃ© gÃ©nÃ©rÃ© avec succÃ¨s!")
                            
                            # Affichage du rapport
                            st.markdown("### ðŸ“„ Rapport gÃ©nÃ©rÃ©")
                            st.markdown(report_content)
                            
                            # Options d'export amÃ©liorÃ©es
                            st.markdown("### ðŸ’¾ Export du rapport")
                            col_txt, col_md, col_html, col_copy = st.columns(4)
                            
                            with col_txt:
                                st.download_button(
                                    label="ðŸ“„ TXT",
                                    data=report_content,
                                    file_name=f"rapport_ameliore_{datetime.now().strftime('%Y%m%d_%H%M')}.txt",
                                    mime="text/plain",
                                    use_container_width=True
                                )
                            
                            with col_md:
                                st.download_button(
                                    label="ðŸ“ MD",
                                    data=report_content,
                                    file_name=f"rapport_ameliore_{datetime.now().strftime('%Y%m%d_%H%M')}.md",
                                    mime="text/markdown",
                                    use_container_width=True
                                )
                            
                            with col_html:
                                # Conversion HTML amÃ©liorÃ©e
                                html_report = f"""
                                <!DOCTYPE html>
                                <html>
                                <head>
                                    <meta charset="UTF-8">
                                    <title>Rapport de Mission AmÃ©liorÃ©</title>
                                    <style>
                                        body {{ font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; margin: 40px; line-height: 1.6; color: #333; }}
                                        h1, h2, h3 {{ color: #2c3e50; }}
                                        h1 {{ border-bottom: 3px solid #3498db; padding-bottom: 10px; }}
                                        h2 {{ border-left: 4px solid #3498db; padding-left: 15px; }}
                                        .header {{ text-align: center; margin-bottom: 30px; background: #f8f9fa; padding: 20px; border-radius: 10px; }}
                                        .footer {{ margin-top: 30px; text-align: center; font-size: 0.9em; color: #666; }}
                                        ul, ol {{ margin-left: 20px; }}
                                        strong {{ color: #2c3e50; }}
                                    </style>
                                </head>
                                <body>
                                    <div class="header">
                                        <h1>Rapport de Mission AmÃ©liorÃ©</h1>
                                        <p><strong>GÃ©nÃ©rÃ© le {datetime.now().strftime('%d/%m/%Y Ã  %H:%M')}</strong></p>
                                        <p>Public cible: {questions_data.get('target_audience', 'Non spÃ©cifiÃ©')}</p>
                                    </div>
                                    {report_content.replace(chr(10), '<br>')}
                                    <div class="footer">
                                        <p>Rapport gÃ©nÃ©rÃ© automatiquement par l'IA Adja DeepSeek</p>
                                    </div>
                                </body>
                                </html>
                                """
                                
                                st.download_button(
                                    label="ðŸŒ HTML",
                                    data=html_report,
                                    file_name=f"rapport_ameliore_{datetime.now().strftime('%Y%m%d_%H%M')}.html",
                                    mime="text/html",
                                    use_container_width=True
                                )
                            
                            with col_copy:
                                if st.button("ðŸ“‹ Copier", use_container_width=True):
                                    st.write("ðŸ“‹ Contenu copiÃ© dans le presse-papiers!")
                                    st.code(report_content, language=None)
                    else:
                        st.error("âŒ Erreur lors de la gÃ©nÃ©ration du rapport")

st.markdown("---")
st.caption("ðŸš€ Planificateur de Mission v2.4")
st.caption("ðŸ’» Developed by @Moctar All rights reserved")
